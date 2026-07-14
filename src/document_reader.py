from __future__ import annotations

from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path

import pdfplumber
from docx import Document


@dataclass
class DocumentReadResult:
    filename: str
    extension: str
    text: str
    page_count: int | None = None
    warnings: list[str] = field(default_factory=list)


class DocumentReadError(Exception):
    """Error controlado durante la lectura de un documento."""


def read_txt(content: bytes) -> str:
    """Lee un TXT probando codificaciones habituales."""

    encodings = ["utf-8-sig", "utf-8", "latin-1"]

    for encoding in encodings:
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue

    raise DocumentReadError(
        "No fue posible identificar la codificación del archivo TXT."
    )


def read_docx(content: bytes) -> str:
    """Extrae párrafos y tablas de un archivo DOCX."""

    try:
        document = Document(BytesIO(content))
    except Exception as error:
        raise DocumentReadError(
            "El archivo DOCX está dañado o no tiene una estructura válida."
        ) from error

    extracted_parts: list[str] = []

    # Párrafos normales
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            extracted_parts.append(text)

    # Contenido de tablas
    for table_number, table in enumerate(document.tables, start=1):
        extracted_parts.append(f"\n[TABLA {table_number}]")

        for row in table.rows:
            cells = [
                cell.text.strip().replace("\n", " ")
                for cell in row.cells
            ]

            if any(cells):
                extracted_parts.append(" | ".join(cells))

    return "\n".join(extracted_parts)


def read_pdf(content: bytes) -> tuple[str, int]:
    """Extrae el texto disponible en un archivo PDF."""

    page_texts: list[str] = []

    try:
        with pdfplumber.open(BytesIO(content)) as pdf:
            page_count = len(pdf.pages)

            for page_number, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                text = text.strip()

                if text:
                    page_texts.append(
                        f"\n[PÁGINA {page_number}]\n{text}"
                    )

    except Exception as error:
        raise DocumentReadError(
            "El archivo PDF está dañado, protegido o no puede leerse."
        ) from error

    return "\n".join(page_texts).strip(), page_count


def read_document(
    filename: str,
    content: bytes,
) -> DocumentReadResult:
    """
    Identifica el tipo de archivo y extrae su contenido textual.
    """

    safe_filename = Path(filename).name
    extension = Path(safe_filename).suffix.lower()

    if not content:
        raise DocumentReadError("El archivo está vacío.")

    warnings: list[str] = []
    page_count: int | None = None

    if extension == ".txt":
        text = read_txt(content)

    elif extension == ".docx":
        text = read_docx(content)

    elif extension == ".pdf":
        text, page_count = read_pdf(content)

        if len(text.strip()) < 30:
            warnings.append(
                "El PDF contiene muy poco texto extraíble. "
                "Puede ser un documento escaneado o contener información "
                "principalmente visual."
            )

    else:
        raise DocumentReadError(
            f"El formato {extension or 'desconocido'} no está soportado."
        )

    if not text.strip():
        warnings.append(
            "No se encontró contenido textual dentro del documento."
        )

    return DocumentReadResult(
        filename=safe_filename,
        extension=extension.lstrip("."),
        text=text.strip(),
        page_count=page_count,
        warnings=warnings,
    )