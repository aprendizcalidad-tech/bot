from __future__ import annotations

import base64
from copy import deepcopy
import hashlib
import html
import json
import os
import re
import random
import shutil
import subprocess
import tempfile
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from io import BytesIO
from pathlib import Path
from typing import Literal
from urllib.parse import parse_qs, urlparse
from xml.sax.saxutils import escape

import pdfplumber
import streamlit as st
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from google import genai
from google.genai import types
from google.oauth2 import service_account
from googleapiclient.discovery import build as build_drive_api
from googleapiclient.errors import HttpError
from googleapiclient.http import MediaIoBaseDownload, MediaIoBaseUpload
from pydantic import BaseModel, ConfigDict, Field
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfgen import canvas as pdf_canvas
from reportlab.platypus import (
    KeepTogether,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


# =========================================================
# CONFIGURACIÓN GENERAL
# =========================================================

GUIDE_EXTENSIONS = ["docx", "pdf", "txt"]
VIDEO_EXTENSIONS = [
    "mp4",
    "mpeg",
    "mov",
    "avi",
    "flv",
    "mpg",
    "webm",
    "wmv",
    "3gp",
    "3gpp",
]
SOURCE_EXTENSIONS = GUIDE_EXTENSIONS + VIDEO_EXTENSIONS
ALLOWED_EXTENSIONS = GUIDE_EXTENSIONS

MODEL_DEFAULT = "gemini-flash-latest"
VIDEO_MODEL_DEFAULT = "gemini-3.5-flash"
MAX_FILE_SIZE_MB = 20
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024
# Los videos grandes deben entrar por Google Drive para no llenar la memoria
# de Streamlit. El cargador directo queda limitado a archivos pequeños.
MAX_VIDEO_SIZE_MB = 250
MAX_VIDEO_SIZE_BYTES = MAX_VIDEO_SIZE_MB * 1024 * 1024
MAX_DRIVE_VIDEO_SIZE_MB = 1900
MAX_DRIVE_VIDEO_SIZE_BYTES = MAX_DRIVE_VIDEO_SIZE_MB * 1024 * 1024
VIDEO_PROCESSING_TIMEOUT_SECONDS = 3600
# Se inicia directamente en cinco minutos. La versión anterior intentaba diez
# minutos y, al fallar, repetía el mismo contenido en dos tramos de cinco.
VIDEO_CHUNK_SECONDS = 300
VIDEO_MIN_CHUNK_SECONDS = 120  # respaldo mínimo de 2 minutos
VIDEO_LONG_THRESHOLD_SECONDS = 1  # todo video con duración conocida usa tramos
VIDEO_MAX_OUTPUT_TOKENS = 32768
VIDEO_ACTION_AUDIT_ENABLED = True
VIDEO_ACTION_AUDIT_MODE = "selective"
VIDEO_MAX_PARALLEL_CHUNKS = 2
VIDEO_CHECKPOINT_ENABLED = True
VIDEO_CHECKPOINT_VERSION = "video-v5-5min-selective-resumable"
VIDEO_CHECKPOINT_LOCAL_DIR = ".video_checkpoints"
VIDEO_POLL_INTERVAL_SECONDS = 5
DRIVE_DOWNLOAD_TIMEOUT_SECONDS = 3600
DRIVE_DOWNLOAD_CHUNK_MB = 8
DRIVE_DOWNLOAD_CHUNK_BYTES = DRIVE_DOWNLOAD_CHUNK_MB * 1024 * 1024
DRIVE_READONLY_SCOPE = "https://www.googleapis.com/auth/drive.readonly"
DRIVE_FILE_SCOPE = "https://www.googleapis.com/auth/drive.file"
MAX_GUIDES = 8
MAX_TEXT_CHARS_PER_FILE = 80_000
MAX_VIDEO_SOURCE_CHARS = 1_200_000
MAX_TOTAL_TEXT_CHARS = 1_400_000
MIN_ACTION_TEXT_CHARS = 12
MAX_TOTAL_UPLOAD_MB = 300
MAX_TOTAL_UPLOAD_BYTES = MAX_TOTAL_UPLOAD_MB * 1024 * 1024

STATUS_LABELS = {
    "completo": "Completo",
    "parcial": "Parcial",
    "faltante": "Faltante",
}

MIME_TYPES = {
    "pdf": "application/pdf",
    "docx": (
        "application/vnd.openxmlformats-officedocument."
        "wordprocessingml.document"
    ),
    "txt": "text/plain",
    "mp4": "video/mp4",
    "mpeg": "video/mpeg",
    "mov": "video/quicktime",
    "avi": "video/x-msvideo",
    "flv": "video/x-flv",
    "mpg": "video/mpeg",
    "webm": "video/webm",
    "wmv": "video/wmv",
    "3gp": "video/3gpp",
    "3gpp": "video/3gpp",
}


# =========================================================
# MODELOS DE DATOS PARA RESPUESTAS ESTRUCTURADAS DE GEMINI
# =========================================================


class MissingQuestion(BaseModel):
    model_config = ConfigDict(extra="forbid")

    category: str = Field(
        default="Información del proceso",
        description=(
            "Categoría breve para agrupar la pregunta, por ejemplo: "
            "Datos operativos, Responsables, Fechas y periodos o Control documental."
        ),
    )
    question: str = Field(
        description=(
            "Pregunta específica que debe responder el usuario. "
            "No debe ser genérica ni repetir otra pregunta."
        )
    )
    why_needed: str = Field(
        description="Razón breve por la cual el dato es necesario."
    )
    required: bool = Field(
        description="Indica si la respuesta es indispensable para generar el documento."
    )


class SectionAssessment(BaseModel):
    model_config = ConfigDict(extra="forbid")

    order: int = Field(ge=1)
    title: str
    guide_instruction: str = Field(
        description=(
            "Instrucción normativa interpretada desde la guía. "
            "No es contenido final del documento."
        )
    )
    criteria: list[str]
    required: bool
    status: Literal["completo", "parcial", "faltante"]
    evidence: list[str] = Field(
        description=(
            "Fragmentos o referencias breves del documento de origen "
            "que sustentan la sección."
        )
    )
    draft_content: str = Field(
        description=(
            "Borrador construido únicamente con información sustentada. "
            "Debe quedar vacío cuando no sea posible redactar sin inventar."
        )
    )
    output_format: Literal["parrafo", "lista", "tabla", "mixto"]
    missing_questions: list[MissingQuestion]


class GuideAnalysis(BaseModel):
    model_config = ConfigDict(extra="forbid")

    detected_process: str
    selected_guide_index: int = Field(ge=0)
    selected_guide_name: str
    supporting_guide_indices: list[int]
    selection_reason: str
    proposed_document_title: str
    general_requirements: list[str]
    sections: list[SectionAssessment]
    warnings: list[str]


class UserAnswer(BaseModel):
    model_config = ConfigDict(extra="forbid")

    section_title: str
    question: str
    answer: str


class FinalTable(BaseModel):
    model_config = ConfigDict(extra="forbid")

    title: str
    headers: list[str]
    rows: list[list[str]]


class FinalSection(BaseModel):
    model_config = ConfigDict(extra="forbid")

    order: int = Field(ge=1)
    title: str
    paragraphs: list[str]
    bullets: list[str]
    numbered_items: list[str] = Field(
        default_factory=list,
        description=(
            "Pasos o actividades que deben presentarse como lista numerada. "
            "No deben incluir viñetas ni el número dentro del texto."
        ),
    )
    tables: list[FinalTable]
    source_basis: list[str]


class ValidationItem(BaseModel):
    model_config = ConfigDict(extra="forbid")

    section_title: str
    criterion: str
    status: Literal["cumple", "parcial", "no_aplica"]
    note: str


class FinalDocument(BaseModel):
    model_config = ConfigDict(extra="forbid")

    title: str
    subtitle: str
    introductory_note: str
    sections: list[FinalSection]
    validation: list[ValidationItem]
    warnings: list[str]


class AuditReport(BaseModel):
    model_config = ConfigDict(extra="forbid")

    validation: list[ValidationItem]
    warnings: list[str]
    editorial_summary: str = Field(
        description=(
            "Resumen breve de la calidad de redacción, coherencia, trazabilidad "
            "y cumplimiento del documento revisado."
        )
    )




class VideoAction(BaseModel):
    """Acción operativa atómica identificada en el documento o video."""

    model_config = ConfigDict(extra="ignore")

    action_id: str = Field(
        default="",
        description="Identificador secuencial. Se normaliza localmente como ACC-0001.",
    )
    timestamp_start: str = Field(
        default="",
        description="Hora de inicio absoluta en formato HH:MM:SS.",
    )
    timestamp_end: str = Field(
        default="",
        description="Hora de finalización absoluta cuando pueda determinarse.",
    )
    actor: str = Field(
        default="",
        description="Cargo, rol o hablante que ejecuta la acción.",
    )
    system: str = Field(
        default="",
        description="Aplicación, sistema o plataforma utilizada.",
    )
    location_path: str = Field(
        default="",
        description="Módulo, menú, pestaña, pantalla o ruta de navegación.",
    )
    action: str = Field(
        description=(
            "Una sola acción operativa, concreta y verificable. No debe contener "
            "una secuencia completa ni fusionar acciones diferentes."
        )
    )
    interface_element: str = Field(
        default="",
        description="Botón, enlace, campo, filtro, lista, casilla o archivo utilizado.",
    )
    data_handled: str = Field(
        default="",
        description="Dato consultado, registrado, modificado, cargado o descargado.",
    )
    validation: str = Field(
        default="",
        description="Validación, condición o comprobación realizada antes de continuar.",
    )
    result: str = Field(
        default="",
        description="Mensaje, estado, registro o resultado que confirma la acción.",
    )
    evidence: str = Field(
        default="",
        description="Evidencia visual o auditiva que sustenta la acción.",
    )
    uncertainty: str = Field(
        default="",
        description="Dato dudoso o no legible que requiere confirmación.",
    )


class VideoActionBatch(BaseModel):
    """Inventario de acciones de un tramo de video."""

    model_config = ConfigDict(extra="ignore")

    actions: list[VideoAction] = Field(
        default_factory=list,
        description=(
            "Todas las acciones operativas atómicas del tramo, en orden "
            "cronológico y sin fusionar acciones diferentes."
        ),
    )


class VideoTranscript(BaseModel):
    """Transcripción estructurada producida a partir de un video."""

    model_config = ConfigDict(extra="ignore")

    detected_language: str = Field(
        default="",
        description="Idioma principal detectado en el video.",
    )
    duration_estimate: str = Field(
        default="",
        description=(
            "Duración aproximada observada, por ejemplo 08:35. "
            "Déjala vacía si no puede determinarse con seguridad."
        ),
    )
    speakers: list[str] = Field(
        default_factory=list,
        description=(
            "Etiquetas de hablantes detectados, por ejemplo Hablante 1, "
            "Hablante 2 o nombres propios solo cuando sean explícitos."
        ),
    )
    full_transcript: str = Field(
        default="",
        description=(
            "Transcripción completa y cronológica. Debe incluir marcas de "
            "tiempo y no resumir intervenciones relevantes."
        ),
    )
    actions: list[VideoAction] = Field(
        default_factory=list,
        description=(
            "Inventario exhaustivo de acciones operativas atómicas. Cada clic, "
            "selección, navegación, diligenciamiento, carga, validación o resultado "
            "diferenciable debe aparecer como un elemento independiente."
        ),
    )
    visual_evidence: list[str] = Field(
        default_factory=list,
        description=(
            "Información visual relevante con marca de tiempo: textos, tablas, "
            "formularios, sistemas, acciones, rutas y evidencias."
        ),
    )
    key_facts: list[str] = Field(
        default_factory=list,
        description="Hechos verificables mencionados o mostrados en el video.",
    )
    uncertainties: list[str] = Field(
        default_factory=list,
        description=(
            "Fragmentos inaudibles, nombres dudosos, textos ilegibles o "
            "información que requiere confirmación humana."
        ),
    )

class AppError(Exception):
    """Error controlado que puede mostrarse al usuario."""


# =========================================================
# UTILIDADES DE CONFIGURACIÓN Y ARCHIVOS
# =========================================================


def read_secret(name: str, default: str = "") -> str:
    """Lee un secreto desde Streamlit o una variable de entorno."""

    try:
        value = st.secrets.get(name, default)
        return str(value) if value is not None else default
    except Exception:
        return os.getenv(name, default)


def read_drive_service_account_info() -> dict:
    """Lee la cuenta de servicio de Drive desde los Secrets.

    Formatos aceptados:
    1. DRIVE_SERVICE_ACCOUNT_JSON como cadena JSON completa.
    2. Tabla TOML [drive_service_account] con los campos del JSON.
    """

    info: dict = {}

    try:
        nested = st.secrets.get("drive_service_account")
        if nested:
            info = dict(nested)
    except Exception:
        info = {}

    if not info:
        raw = read_secret("DRIVE_SERVICE_ACCOUNT_JSON", "").strip()
        if raw:
            try:
                parsed = json.loads(raw)
            except json.JSONDecodeError as error:
                raise AppError(
                    "DRIVE_SERVICE_ACCOUNT_JSON no contiene un JSON válido. "
                    "Pega el contenido completo del archivo de cuenta de servicio."
                ) from error
            if not isinstance(parsed, dict):
                raise AppError(
                    "DRIVE_SERVICE_ACCOUNT_JSON debe contener un objeto JSON."
                )
            info = parsed

    if info.get("private_key"):
        info["private_key"] = str(info["private_key"]).replace("\\n", "\n")

    return info


def drive_service_account_email() -> str:
    """Devuelve el correo que debe recibir acceso al archivo o carpeta."""

    try:
        return str(read_drive_service_account_info().get("client_email", "")).strip()
    except AppError:
        return ""


def build_drive_service(writable: bool = False):
    """Crea un cliente de Google Drive autenticado con cuenta de servicio.

    El modo normal conserva el acceso de solo lectura usado para descargar
    videos. El modo writable añade ``drive.file`` exclusivamente para guardar
    checkpoints dentro de una carpeta compartida con la cuenta de servicio.
    """

    info = read_drive_service_account_info()
    required = {"type", "client_email", "private_key", "token_uri"}
    missing = sorted(field for field in required if not info.get(field))
    if missing:
        raise AppError(
            "Falta configurar la cuenta de servicio de Google Drive. "
            "Campos ausentes: " + ", ".join(missing) + "."
        )

    scopes = [DRIVE_READONLY_SCOPE]
    if writable:
        scopes.append(DRIVE_FILE_SCOPE)

    try:
        credentials = service_account.Credentials.from_service_account_info(
            info,
            scopes=scopes,
        )
        return build_drive_api(
            "drive",
            "v3",
            credentials=credentials,
            cache_discovery=False,
        )
    except Exception as error:
        raise AppError(
            "No fue posible autenticar la cuenta de servicio de Google Drive. "
            "Revisa el JSON guardado en Secrets."
        ) from error

def extract_drive_file_id(reference: str) -> str:
    """Extrae el ID desde un enlace de Drive o acepta directamente el ID."""

    value = str(reference or "").strip()
    if not value:
        raise AppError("Debes pegar el enlace o el ID del video de Google Drive.")

    if re.fullmatch(r"[A-Za-z0-9_-]{15,}", value):
        return value

    try:
        parsed = urlparse(value)
        query = parse_qs(parsed.query)
        for key in ("id", "file_id"):
            candidate = (query.get(key) or [""])[0].strip()
            if re.fullmatch(r"[A-Za-z0-9_-]{15,}", candidate):
                return candidate

        patterns = (
            r"/file/d/([A-Za-z0-9_-]{15,})",
            r"/d/([A-Za-z0-9_-]{15,})",
            r"/uc/([A-Za-z0-9_-]{15,})",
        )
        for pattern in patterns:
            match = re.search(pattern, parsed.path)
            if match:
                return match.group(1)
    except Exception:
        pass

    raise AppError(
        "No se pudo identificar el archivo de Drive. Usa un enlace de archivo "
        "como https://drive.google.com/file/d/ID/view o pega únicamente el ID."
    )


def safe_filename(filename: str) -> str:
    """Evita utilizar rutas incluidas en el nombre del archivo."""

    return Path(filename).name


def file_extension(filename: str) -> str:
    """Devuelve la extensión sin punto y en minúsculas."""

    return Path(filename).suffix.lower().lstrip(".")


def format_file_size(size_bytes: int) -> str:
    if size_bytes < 1024:
        return f"{size_bytes} B"
    if size_bytes < 1024 * 1024:
        return f"{size_bytes / 1024:.1f} KB"
    return f"{size_bytes / (1024 * 1024):.1f} MB"


def is_video_extension(extension: str) -> bool:
    return extension.lower() in VIDEO_EXTENSIONS


def uploaded_file_fingerprint(filename: str, content: bytes) -> str:
    """Identifica un archivo sin guardar su contenido en disco permanentemente."""

    digest = hashlib.sha256(content).hexdigest()
    return f"{safe_filename(filename)}:{len(content)}:{digest}"


def validate_uploaded_file(
    filename: str,
    content: bytes,
    allowed_extensions: list[str] | None = None,
) -> None:
    """Valida extensión, contenido y tamaño según el tipo de archivo."""

    extension = file_extension(filename)
    allowed = allowed_extensions or ALLOWED_EXTENSIONS

    if extension not in allowed:
        raise AppError(
            f"{filename}: el formato .{extension or 'desconocido'} no está permitido."
        )

    if not content:
        raise AppError(f"{filename}: el archivo está vacío.")

    if is_video_extension(extension):
        if len(content) > MAX_VIDEO_SIZE_BYTES:
            raise AppError(
                f"{filename}: supera el límite de {MAX_VIDEO_SIZE_MB} MB para video."
            )
    elif len(content) > MAX_FILE_SIZE_BYTES:
        raise AppError(
            f"{filename}: supera el límite de {MAX_FILE_SIZE_MB} MB."
        )


def decode_txt(content: bytes) -> str:
    """Lee TXT probando codificaciones frecuentes."""

    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue

    raise AppError("No fue posible determinar la codificación del archivo TXT.")


def extract_docx_text(content: bytes) -> str:
    """Extrae títulos, párrafos, tablas, encabezados y pies de página."""

    try:
        document = Document(BytesIO(content))
    except Exception as error:
        raise AppError("El archivo DOCX está dañado o no es válido.") from error

    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        style_name = (paragraph.style.name or "").lower()
        if "heading" in style_name or "título" in style_name or "titulo" in style_name:
            parts.append(f"[TÍTULO] {text}")
        else:
            parts.append(text)

    for table_number, table in enumerate(document.tables, start=1):
        parts.append(f"[TABLA {table_number}]")
        for row in table.rows:
            cells = [
                re.sub(r"\s+", " ", cell.text).strip()
                for cell in row.cells
            ]
            if any(cells):
                parts.append(" | ".join(cells))

    seen_headers: set[str] = set()
    seen_footers: set[str] = set()

    for section in document.sections:
        header_text = "\n".join(
            paragraph.text.strip()
            for paragraph in section.header.paragraphs
            if paragraph.text.strip()
        )
        if header_text and header_text not in seen_headers:
            parts.append(f"[ENCABEZADO]\n{header_text}")
            seen_headers.add(header_text)

        footer_text = "\n".join(
            paragraph.text.strip()
            for paragraph in section.footer.paragraphs
            if paragraph.text.strip()
        )
        if footer_text and footer_text not in seen_footers:
            parts.append(f"[PIE DE PÁGINA]\n{footer_text}")
            seen_footers.add(footer_text)

    return "\n".join(parts).strip()


def extract_pdf_text(content: bytes) -> tuple[str, int]:
    """Extrae texto local de PDF para vista previa; Gemini recibe el PDF nativo."""

    pages: list[str] = []

    try:
        with pdfplumber.open(BytesIO(content)) as pdf:
            page_count = len(pdf.pages)
            for page_number, page in enumerate(pdf.pages, start=1):
                text = (page.extract_text() or "").strip()
                if text:
                    pages.append(f"[PÁGINA {page_number}]\n{text}")
    except Exception as error:
        raise AppError("El PDF está dañado, protegido o no puede leerse.") from error

    return "\n\n".join(pages).strip(), page_count


def extract_text(filename: str, content: bytes) -> tuple[str, int | None]:
    """Extrae texto para DOCX, PDF o TXT."""

    extension = file_extension(filename)

    if extension == "txt":
        return decode_txt(content).strip(), None

    if extension == "docx":
        return extract_docx_text(content), None

    if extension == "pdf":
        try:
            return extract_pdf_text(content)
        except AppError:
            # La extracción local es solo para vista previa. Gemini recibirá
            # el PDF completo de forma nativa y hará el análisis principal.
            return "", None

    if is_video_extension(extension):
        # El video se transcribe mediante Gemini Files API. No intentamos
        # extraerlo localmente para evitar depender de FFmpeg u OCR.
        return "", None

    raise AppError(f"Formato no soportado: .{extension}")


def build_file_record(
    uploaded_file,
    role: str,
    index: int,
    allowed_extensions: list[str] | None = None,
) -> dict:
    """Convierte UploadedFile en una estructura persistente en sesión."""

    name = safe_filename(uploaded_file.name)
    content = uploaded_file.getvalue()
    validate_uploaded_file(name, content, allowed_extensions=allowed_extensions)
    text, page_count = extract_text(name, content)

    warnings: list[str] = []
    extension = file_extension(name)

    if extension == "pdf" and len(text) < 30:
        warnings.append(
            "El PDF tiene poco texto extraíble localmente. "
            "Gemini lo analizará visualmente como PDF nativo."
        )

    if extension in {"docx", "txt"} and not text:
        warnings.append("No se encontró texto legible en el archivo.")

    if is_video_extension(extension):
        warnings.append(
            "El archivo de origen es un video. Debe transcribirse y revisarse "
            "antes de comparar su información con las guías."
        )

    return {
        "role": role,
        "index": index,
        "name": name,
        "extension": extension,
        "mime_type": MIME_TYPES[extension],
        "content": content,
        "size_bytes": len(content),
        "text": text,
        "page_count": page_count,
        "warnings": warnings,
        "fingerprint": uploaded_file_fingerprint(name, content),
        "is_video": is_video_extension(extension),
    }


def clipped_text(text: str, limit: int = MAX_TEXT_CHARS_PER_FILE) -> str:
    """Limita texto para evitar solicitudes excesivas."""

    if len(text) <= limit:
        return text

    return (
        text[:limit]
        + "\n\n[CONTENIDO RECORTADO POR LÍMITE TÉCNICO DE LA APLICACIÓN]"
    )


def pdf_document_part(record: dict) -> types.Part:
    """Construye una entrada PDF nativa para generateContent."""

    return types.Part.from_bytes(
        data=record["content"],
        mime_type="application/pdf",
    )


def append_record_to_gemini_input(
    parts: list[types.Part],
    record: dict,
    label: str,
) -> None:
    """Agrega un archivo al input multimodal como objetos Part válidos.

    Las transcripciones de video no se recortan al límite general de 80.000
    caracteres. El inventario de acciones y la transcripción aprobada deben llegar
    completos a las fases de selección de guía, redacción y auditoría.
    """

    parts.append(
        types.Part.from_text(
            text=(
                f"\n===== INICIO {label}: {record['name']} =====\n"
                f"Rol del archivo: {record['role']}\n"
            )
        )
    )

    if record["extension"] == "pdf":
        parts.append(pdf_document_part(record))
    else:
        raw_text = str(record.get("text") or "").strip()
        is_video_source = bool(
            record.get("is_video")
            or record.get("video_actions")
            or str(record.get("origin_kind") or "").lower()
            in {"google_drive", "direct_upload"}
        )

        if is_video_source:
            extracted_text = raw_text[:MAX_VIDEO_SOURCE_CHARS]
            if len(raw_text) > MAX_VIDEO_SOURCE_CHARS:
                extracted_text += (
                    "\n\n[ADVERTENCIA: EL ORIGEN DE VIDEO SUPERÓ EL LÍMITE DE "
                    "SEGURIDAD. EL INVENTARIO DE ACCIONES DEBE USARSE COMO FUENTE "
                    "PRINCIPAL Y NO PUEDE RESUMIRSE.]"
                )
        else:
            extracted_text = clipped_text(raw_text).strip()

        if not extracted_text:
            extracted_text = "[ARCHIVO SIN TEXTO EXTRAÍBLE]"
        parts.append(types.Part.from_text(text=extracted_text))

    parts.append(
        types.Part.from_text(
            text=f"\n===== FIN {label}: {record['name']} =====\n"
        )
    )

# =========================================================
# PROMPTS Y LLAMADAS A GEMINI
# =========================================================


def analysis_prompt(guide_names: list[str], source_name: str) -> str:
    catalog = "\n".join(
        f"- Índice {index}: {name}"
        for index, name in enumerate(guide_names)
    )

    return f"""
Actúa como analista documental, profesional de calidad y redactor técnico institucional.

OBJETIVO
Comparar las GUÍAS cargadas por el usuario con el DOCUMENTO DE ORIGEN,
seleccionar la guía principal aplicable, interpretar sus instrucciones y
determinar qué contenido puede redactarse con evidencia verificable.

CATÁLOGO DE GUÍAS
{catalog}

DOCUMENTO DE ORIGEN
{source_name}

MÉTODO OBLIGATORIO EN CUATRO ETAPAS
ETAPA 1 — INTERPRETAR LA GUÍA
- Identifica el tipo de documento, las secciones, el orden, los criterios,
  las tablas, las listas y las reglas de presentación.
- Las guías son instrucciones de construcción; no copies sus explicaciones
  como contenido factual del documento final.
- FIDELIDAD ESTRUCTURAL OBLIGATORIA: identifica todos los numerales, pasos,
  actividades, controles, requisitos y subapartados de la guía. No los resumas,
  no los agrupes y no los omitas. Cada elemento enumerado debe conservarse como
  un elemento independiente. Si la guía contiene 8 pasos, el análisis debe
  reconocer los 8 y el documento final debe contener los 8.

ETAPA 2 — EXTRAER EVIDENCIA
- Localiza hechos explícitos en el documento de origen.
- Si el origen proviene de un video, utiliza el bloque INVENTARIO DE ACCIONES OPERATIVAS como fuente prioritaria y conserva todas las acciones ACC-XXXX.
- Una sección PASO A PASO o PROCEDIMIENTO sin una cantidad fija en la guía debe desarrollarse con una actividad independiente por cada acción operativa.
- Distingue entre información explícita, información derivable sin agregar
  hechos nuevos e información ausente.
- Conserva literalmente nombres, cargos, radicados, fechas, cifras,
  plataformas y denominaciones oficiales.

ETAPA 3 — PROPONER REDACCIÓN
- Redacta draft_content con lenguaje institucional claro, directo y preciso.
- No inventes datos, responsables, decisiones, resultados, periodos ni
  controles.
- Evita repeticiones, muletillas, frases ambiguas y contenido genérico.

ETAPA 4 — IDENTIFICAR BRECHAS
- Formula únicamente preguntas críticas que no puedan resolverse con el
  origen.
- Cada pregunta debe ser específica, breve y responder a un solo dato.
- Agrupa cada pregunta en una categoría útil: Datos operativos,
  Responsables, Fechas y periodos, Control documental u otra equivalente.
- No preguntes por datos de encabezado que la guía ya deja como XX o
  pendiente de asignación, salvo que sean indispensables para el contenido.

INTERPRETACIÓN
Ejemplo: si la guía dice “OBJETIVO: describir la intención del documento en
términos del qué y el para qué”, debes redactar un objetivo que explique QUÉ
se hará y PARA QUÉ se hará. No copies esa instrucción.

REGLAS
1. Selecciona una guía principal usando su índice real.
2. Incluye guías complementarias solo si aportan reglas compatibles.
3. Ordena las secciones según la guía principal.
4. evidence debe referirse al documento de origen, no a la guía.
5. draft_content debe ser contenido final propuesto.
6. Evita preguntas duplicadas; una respuesta puede servir a varias secciones.
7. Para procedimientos secuenciales usa output_format “lista” o “mixto”.
8. selected_guide_name debe coincidir exactamente con el archivo del catálogo.
9. supporting_guide_indices solo puede contener índices válidos y distintos del principal.
10. La concisión aplica únicamente a la redacción de cada elemento; nunca autoriza
    reducir la cantidad de pasos, requisitos, actividades o controles exigidos.
11. Cuando la guía enumere N elementos, conserva N elementos independientes y en
    el mismo orden. No combines dos pasos en uno aunque sean similares.
12. Cuando la guía ordene DETALLAR un paso a paso sin fijar una cantidad, no interpretes la instrucción como un único paso. Debes reconocer todas las acciones operativas del origen y mantener una acción diferenciable por paso.
13. Está prohibido convertir un inventario de múltiples acciones ACC-XXXX en uno o pocos pasos generales.

ESTADOS
- completo: información suficiente para cumplir todos los criterios.
- parcial: existe un borrador sustentado, pero falta información crítica.
- faltante: no es posible redactar sin inventar.

Devuelve únicamente el objeto estructurado solicitado.
""".strip()


def generation_prompt(
    analysis: GuideAnalysis,
    answers: list[UserAnswer],
) -> str:
    return f"""
Actúa como redactor técnico institucional y auditor de cumplimiento documental.

Debes producir un DOCUMENTO FINAL profesional utilizando:
1. Las guías aplicables adjuntas.
2. El documento de origen.
3. El análisis estructurado previo.
4. Las respuestas verificables del usuario.

ANÁLISIS PREVIO
{analysis.model_dump_json(indent=2)}

RESPUESTAS DEL USUARIO
{json.dumps([answer.model_dump() for answer in answers], ensure_ascii=False, indent=2)}

PROCESO OBLIGATORIO
1. INTERPRETAR: confirma la finalidad y criterios de cada sección.
2. REDACTAR: construye el contenido con trazabilidad y lenguaje institucional.
3. NORMALIZAR: corrige ortografía, concordancia, puntuación, repeticiones,
   mayúsculas, fechas y denominaciones.
4. AUDITAR: evalúa cada criterio y registra el resultado en validation.

REGLAS DE REDACCIÓN
- Usa un estilo formal, preciso, coherente y eficiente.
- Prioriza oraciones claras y directas; evita párrafos inflados.
- La eficiencia de redacción NO permite resumir ni ampliar la estructura exigida: conserva
  todos los pasos, actividades, controles, requisitos y subapartados de la guía.
- Mantén una idea principal por oración cuando sea posible.
- Usa voz institucional o impersonal.
- Para procedimientos, usa verbos de acción y orden lógico.
- Si el origen contiene un INVENTARIO DE ACCIONES OPERATIVAS, cada ACC-XXXX debe convertirse en un paso o fila independiente, en el mismo orden cronológico.
- No redactes un resumen del procedimiento. Debes explicar qué se hace, dónde, con qué elemento, qué dato se gestiona y cómo se verifica, usando la evidencia disponible para cada acción.
- No inventes hechos, responsables, fechas, resultados, decisiones ni controles.
- Conserva literalmente datos sensibles y denominaciones oficiales.
- No copies las instrucciones de la guía como contenido factual.
- El objetivo debe integrar claramente el QUÉ y el PARA QUÉ cuando la guía lo exija.
- El alcance debe indicar inicio, límite y finalización solo cuando estén sustentados.
- Las definiciones deben presentarse en orden alfabético cuando la guía lo solicite.
- No incluyas una sección de datos faltantes en el documento final.

FIDELIDAD ESTRUCTURAL OBLIGATORIA
- Devuelve exactamente una sección por cada título de la guía principal.
- Copia cada título de sección literalmente, sin renombrarlo, resumirlo, combinarlo
  con otro ni sustituirlo por un sinónimo.
- Conserva exactamente el orden de aparición de los títulos en la guía principal.
- Reproduce todos los elementos enumerados por la guía, uno a uno y en el mismo
  orden. No resumas, fusiones, agrupes, sustituyas ni omitas elementos.
- Si la guía exige 8 pasos fijos, numbered_items debe contener exactamente 8 pasos independientes: no 7, no 9.
- Si la guía no fija una cantidad y ordena detallar el PASO A PASO o el PROCEDIMIENTO, la cantidad se obtiene del inventario de acciones del origen. En ese caso no existe un máximo artificial y no deben eliminarse pasos.
- Cada paso puede mejorarse lingüísticamente, pero debe conservar su acción,
  responsable, condición, control y resultado cuando estos estén sustentados.
- Antes de responder, cuenta los elementos de la guía y compáralos con los del
  documento final. Corrige cualquier diferencia antes de devolver el JSON.

ESTRUCTURA DE LISTAS
- Usa numbered_items exclusivamente para pasos, actividades secuenciales,
  procedimientos, etapas o instrucciones ordenadas.
- Cada elemento de numbered_items debe ir sin “1.”, “2.” ni viñeta inicial.
- Usa bullets para responsabilidades, definiciones, condiciones, referencias
  y listados sin secuencia.
- No mezcles una viñeta con un número, por ejemplo “• 1.”.

TABLAS
- Usa tablas solo cuando la guía las solicite o mejoren claramente la lectura.
- En control de cambios conserva las columnas de la guía.
- Evita textos con espacios repetidos o saltos manuales innecesarios.

VALIDACIÓN
- validation debe revisar cada criterio de la guía.
- Marca cumple, parcial o no_aplica y explica brevemente.
- source_basis debe indicar el sustento de cada sección.
- Registra en warnings cualquier limitación que no pueda resolverse sin inventar.

Devuelve únicamente el objeto estructurado solicitado.
""".strip()


def audit_prompt(
    analysis: GuideAnalysis,
    final_document: FinalDocument,
) -> str:
    return f"""
Actúa como auditor documental independiente.

Compara el borrador final contra las instrucciones y criterios identificados
en la guía. Evalúa contenido, redacción y trazabilidad.

ANÁLISIS DE LA GUÍA
{analysis.model_dump_json(indent=2)}

BORRADOR A REVISAR
{final_document.model_dump_json(indent=2)}

CRITERIOS DE AUDITORÍA
1. Cada sección cumple su finalidad y criterios.
2. La redacción es formal, clara, precisa y sin repeticiones.
3. No existen hechos inventados ni afirmaciones sin sustento.
4. Nombres, cargos, fechas, cifras y sistemas son consistentes.
5. Los pasos están ordenados y no mezclan viñetas con numeración.
6. Para listas fijas, la cantidad coincide exactamente con la guía. Para procedimientos dinámicos, la cantidad coincide con el inventario de acciones del origen y ninguna ACC-XXXX fue resumida, fusionada u omitida.
7. Las tablas son coherentes y no contienen espacios artificiales.
8. El documento mantiene trazabilidad con el origen y las respuestas.
9. Las limitaciones reales se registran como advertencias.

Devuelve validation para todos los criterios relevantes, warnings concretos
y un editorial_summary breve. No reescribas el documento.
""".strip()


def section_regeneration_prompt(
    analysis: GuideAnalysis,
    answers: list[UserAnswer],
    final_document: FinalDocument,
    section: FinalSection,
) -> str:
    return f"""
Actúa como redactor técnico institucional.

Regenera ÚNICAMENTE la sección indicada, manteniendo el orden y la finalidad
definidos por la guía. Usa exclusivamente hechos sustentados en los archivos,
el análisis y las respuestas del usuario.

ANÁLISIS
{analysis.model_dump_json(indent=2)}

RESPUESTAS
{json.dumps([answer.model_dump() for answer in answers], ensure_ascii=False, indent=2)}

CONTEXTO DEL DOCUMENTO
Título: {final_document.title}
Subtítulo: {final_document.subtitle}

SECCIÓN A REGENERAR
{section.model_dump_json(indent=2)}

REGLAS
- Mejora claridad, precisión, cohesión y eficiencia.
- No inventes información.
- Conserva el número de orden de la sección.
- Usa numbered_items para pasos secuenciales y bullets para listas no ordenadas.
- No incluyas números dentro del texto de numbered_items.
- Conserva todos los pasos exigidos por la guía. Si la sección es dinámica, conserva una acción del inventario por paso o fila; no reduzcas, combines ni omitas acciones durante la regeneración.
- Mantén las tablas solo si son necesarias o exigidas por la guía.
- Devuelve únicamente la estructura FinalSection.
""".strip()



NUMBERED_GUIDE_LINE_RE = re.compile(
    r"^\s*(?:[•●▪◦\-–—]\s*)?(?P<number>\d{1,3})\s*[.)\-:]\s+(?P<text>.+?)\s*$"
)
GUIDE_BULLET_LINE_RE = re.compile(
    r"^\s*[•●▪◦\-–—]\s*(?P<text>.+?)\s*$"
)
GUIDE_INLINE_HEADING_RE = re.compile(
    r"^\s*(?P<title>[A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ0-9 /()_\-]{2,90})\s*:\s*(?P<body>.*)$"
)

KNOWN_GUIDE_SECTION_TITLES = {
    "OBJETIVO",
    "ALCANCE",
    "RESPONSABLES",
    "DEFINICIONES",
    "CONDICIONES GENERALES",
    "CONDICIONES ESPECIALES",
    "POLÍTICAS",
    "POLITICAS",
    "LINEAMIENTOS",
    "DESARROLLO",
    "METODOLOGÍA",
    "METODOLOGIA",
    "PROCEDIMIENTO",
    "PASO A PASO",
    "ACTIVIDADES",
    "DOCUMENTOS DE REFERENCIA",
    "REFERENCIAS",
    "ANEXOS",
    "REGISTROS",
    "CONTROL DE CAMBIOS",
    "INDICADORES",
    "RIESGOS",
    "CONTROLES",
}

PROCEDURAL_TITLE_TERMS = (
    "paso",
    "procedimiento",
    "actividad",
    "etapa",
    "metodolog",
    "desarrollo",
    "instruccion",
)


def _clean_guide_line(line: str) -> str:
    """Limpia etiquetas de extracción sin alterar el título institucional."""

    cleaned = re.sub(r"\s+", " ", str(line)).strip()
    cleaned = re.sub(r"^\[(?:PÁGINA|PAGINA)\s+\d+\]\s*", "", cleaned, flags=re.I)
    cleaned = re.sub(r"^\[TABLA\s+\d+\]\s*", "", cleaned, flags=re.I)
    return cleaned.strip()


def _is_page_artifact_line(line: str) -> bool:
    """Ignora encabezados, pies y metadatos repetidos entre páginas."""

    normalized = re.sub(r"\s+", " ", line).strip().lower()
    if not normalized:
        return False

    artifact_terms = (
        "copia no controlada",
        "antes de imprimir",
        "piense en el medio ambiente",
        "plantilla de ",
    )
    if any(term in normalized for term in artifact_terms):
        return True

    if re.match(r"^(código|codigo|versión|version|fecha|página|pagina)\s*:", normalized):
        return True
    if re.fullmatch(r"x{2}(?:-x{2}){1,4}", normalized):
        return True
    return False


def _guide_heading_from_line(line: str) -> tuple[str, str] | None:
    """Devuelve (título exacto, contenido en la misma línea) cuando es encabezado."""

    original = _clean_guide_line(line)
    if not original or _is_page_artifact_line(original):
        return None

    if original.upper().startswith("[ENCABEZADO]") or original.upper().startswith("[PIE DE PÁGINA]"):
        return None

    if original.upper().startswith("[TÍTULO]") or original.upper().startswith("[TITULO]"):
        title = re.sub(r"^\[(?:TÍTULO|TITULO)\]\s*", "", original, flags=re.I).strip().strip(":")
        return (title, "") if title else None

    inline = GUIDE_INLINE_HEADING_RE.match(original)
    if inline:
        title = inline.group("title").strip()
        body = inline.group("body").strip()
        if title.upper() in KNOWN_GUIDE_SECTION_TITLES or len(title.split()) <= 12:
            return title, body

    normalized_title = original.strip().strip(":")
    if normalized_title.upper() in KNOWN_GUIDE_SECTION_TITLES:
        return normalized_title, ""

    if NUMBERED_GUIDE_LINE_RE.match(normalized_title) or GUIDE_BULLET_LINE_RE.match(normalized_title):
        return None
    if len(normalized_title) > 100 or len(normalized_title.split()) > 12:
        return None
    if normalized_title.endswith((".", ";", ",")):
        return None
    if "|" in normalized_title:
        return None

    letters = [char for char in normalized_title if char.isalpha()]
    if not letters:
        return None
    uppercase_ratio = sum(char.isupper() for char in letters) / len(letters)
    table_header_terms = sum(
        term in normalized_title.upper()
        for term in ("FECHA", "VERSIÓN", "DESCRIPCIÓN", "RESPONSABLE", "CÓDIGO")
    )
    if table_header_terms >= 2:
        return None

    # Encabezados institucionales o el nombre de la plantilla no son secciones.
    lowered = normalized_title.lower()
    if lowered.startswith(("plantilla de ", "formato de ")):
        return None

    if uppercase_ratio >= 0.82:
        return normalized_title, ""
    return None


def _extract_steps_from_section(title: str, body_lines: list[str]) -> list[str]:
    """Extrae únicamente los pasos que la guía define de forma explícita.

    Reglas:
    - Si existen numerales (1., 2., 3...), esos numerales son la única fuente
      válida para determinar la cantidad de pasos.
    - Las viñetas sin número que aparezcan dentro de un paso se consideran
      aclaraciones o subpuntos del paso anterior; nunca crean un paso adicional.
    - Si no hay numerales y la sección es claramente procedimental, cada viñeta
      independiente puede representar un paso.
    - Se eliminan duplicados causados por saltos de página o encabezados repetidos.
    """

    procedural = any(term in title.lower() for term in PROCEDURAL_TITLE_TERMS)
    numbered_items: list[dict] = []
    bullet_items: list[str] = []
    explicit_numbering_found = False

    for raw_line in body_lines:
        line = _clean_guide_line(raw_line)
        if not line or _is_page_artifact_line(line):
            continue
        if line.upper().startswith(("[ENCABEZADO]", "[PIE DE PÁGINA]", "[PIE DE PAGINA]")):
            continue
        if _guide_heading_from_line(line):
            continue

        numbered = NUMBERED_GUIDE_LINE_RE.match(line)
        if numbered:
            explicit_numbering_found = True
            step_number = int(numbered.group("number"))
            step_text = numbered.group("text").strip()

            existing = next(
                (item for item in numbered_items if item["number"] == step_number),
                None,
            )
            if existing is None:
                numbered_items.append({"number": step_number, "text": step_text})
            elif step_text and step_text.casefold() not in existing["text"].casefold():
                existing["text"] = f"{existing['text'].rstrip()} {step_text}".strip()
            continue

        bullet = GUIDE_BULLET_LINE_RE.match(line)
        if bullet:
            bullet_text = bullet.group("text").strip()
            nested_number = NUMBERED_GUIDE_LINE_RE.match(bullet_text)
            if nested_number:
                explicit_numbering_found = True
                step_number = int(nested_number.group("number"))
                step_text = nested_number.group("text").strip()
                existing = next(
                    (item for item in numbered_items if item["number"] == step_number),
                    None,
                )
                if existing is None:
                    numbered_items.append({"number": step_number, "text": step_text})
                elif step_text and step_text.casefold() not in existing["text"].casefold():
                    existing["text"] = f"{existing['text'].rstrip()} {step_text}".strip()
                continue

            if explicit_numbering_found and numbered_items:
                numbered_items[-1]["text"] = (
                    f"{numbered_items[-1]['text'].rstrip()} {bullet_text}"
                ).strip()
            elif procedural:
                bullet_items.append(bullet_text)
            continue

        if explicit_numbering_found and numbered_items:
            numbered_items[-1]["text"] = (
                f"{numbered_items[-1]['text'].rstrip()} {line}"
            ).strip()
        elif procedural and bullet_items:
            bullet_items[-1] = f"{bullet_items[-1].rstrip()} {line}".strip()

    if explicit_numbering_found:
        numbered_items.sort(key=lambda item: item["number"])
        result: list[str] = []
        seen_numbers: set[int] = set()
        for item in numbered_items:
            if item["number"] in seen_numbers:
                continue
            seen_numbers.add(item["number"])
            value = _strip_list_prefix(item["text"])
            if value:
                result.append(value)
        return result

    if procedural:
        return [
            value
            for value in (_strip_list_prefix(item) for item in bullet_items)
            if value
        ]

    return []


def _has_explicit_numbered_steps(body_lines: list[str]) -> bool:
    return any(
        NUMBERED_GUIDE_LINE_RE.match(_clean_guide_line(line))
        for line in body_lines
        if _clean_guide_line(line)
    )


def _looks_like_dynamic_procedure(title: str, instruction: str) -> bool:
    """Detecta instrucciones que piden construir el procedimiento desde el origen."""

    combined = f"{title} {instruction}".lower()
    procedural_title = any(term in title.lower() for term in PROCEDURAL_TITLE_TERMS)
    dynamic_phrases = (
        "detallar el paso a paso",
        "todas las acciones",
        "cantidad de pasos no es fija",
        "documento o video",
        "video de origen",
        "sin omitir actividades",
        "describir detalladamente",
        "actividad mencionada",
        "flujo de actividades",
    )
    return procedural_title and any(phrase in combined for phrase in dynamic_phrases)


def _procedure_step_mode(
    title: str,
    instruction: str,
    numbered_items: list[str],
    explicit_numbering: bool,
) -> str:
    """Clasifica la sección como lista fija, procedimiento dinámico o no procedimental."""

    title_key = title.lower()
    instruction_key = instruction.lower()
    procedural = any(term in title_key for term in PROCEDURAL_TITLE_TERMS)

    if explicit_numbering and numbered_items:
        return "fixed"

    if _looks_like_dynamic_procedure(title, instruction) or procedural:
        table_terms = all(
            term in instruction_key
            for term in ("actividad", "descripción", "responsable", "evidencia")
        )
        if "procedimiento" in title_key and table_terms:
            return "dynamic_table"
        return "dynamic_list"

    return "none"


def extract_exact_guide_structure(text: str) -> list[dict]:
    """Extrae títulos, instrucciones y tipo de secuencia de la guía.

    Una guía puede contener pasos fijos explícitamente numerados o una instrucción
    dinámica que ordena construir el paso a paso a partir del documento/video. Las
    viñetas explicativas de una instrucción dinámica son criterios, no pasos fijos.
    """

    raw_lines = str(text or "").splitlines()
    sections: list[dict] = []
    current_title: str | None = None
    current_body: list[str] = []

    def flush() -> None:
        nonlocal current_title, current_body
        if not current_title:
            current_body = []
            return

        body_lines = [
            _clean_guide_line(line)
            for line in current_body
            if _clean_guide_line(line)
            and not _is_page_artifact_line(_clean_guide_line(line))
        ]
        instruction = re.sub(r"\s+", " ", " ".join(body_lines)).strip()
        explicit_numbering = _has_explicit_numbered_steps(body_lines)
        numbered_items = _extract_steps_from_section(current_title, body_lines)

        # Las viñetas de una instrucción dinámica describen qué debe contener cada
        # paso; no representan una cantidad fija de pasos.
        if _looks_like_dynamic_procedure(current_title, instruction) and not explicit_numbering:
            numbered_items = []

        step_mode = _procedure_step_mode(
            current_title,
            instruction,
            numbered_items,
            explicit_numbering,
        )

        sections.append(
            {
                "order": len(sections) + 1,
                "title": current_title.strip(),
                "instruction": instruction,
                "numbered_items": numbered_items,
                "step_mode": step_mode,
            }
        )
        current_title = None
        current_body = []

    for raw_line in raw_lines:
        line = _clean_guide_line(raw_line)
        if not line or _is_page_artifact_line(line):
            continue

        heading = _guide_heading_from_line(line)
        if heading:
            title, same_line_body = heading
            flush()
            current_title = title.strip()
            current_body = [same_line_body] if same_line_body else []
            continue

        if current_title:
            current_body.append(line)

    flush()

    unique: list[dict] = []
    seen_titles: set[str] = set()
    for section in sections:
        key = re.sub(r"\s+", " ", section["title"]).strip().casefold()
        if not key or key in seen_titles:
            continue
        seen_titles.add(key)
        section["order"] = len(unique) + 1
        unique.append(section)
    return unique

def build_guide_structure_manifest(guides: list[dict]) -> list[dict]:
    manifest: list[dict] = []
    for guide_index, guide in enumerate(guides):
        sections = extract_exact_guide_structure(guide.get("text", ""))
        manifest.append(
            {
                "guide_index": guide_index,
                "guide_name": guide.get("name", f"Guía {guide_index}"),
                "sections": sections,
            }
        )
    return manifest


def build_guide_fidelity_manifest(guides: list[dict]) -> list[dict]:
    """Mantiene compatibilidad con la interfaz de control de pasos."""

    result: list[dict] = []
    for guide in build_guide_structure_manifest(guides):
        for section in guide["sections"]:
            if section["numbered_items"]:
                result.append(
                    {
                        "guide_index": guide["guide_index"],
                        "guide_name": guide["guide_name"],
                        "section_title": section["title"],
                        "items": list(section["numbered_items"]),
                        "expected_count": len(section["numbered_items"]),
                    }
                )
    return result


def render_fidelity_manifest(manifest: list[dict]) -> str:
    if not manifest:
        return "No se detectaron listas numeradas obligatorias mediante extracción local."

    blocks = [
        "ESTRUCTURA BLOQUEADA DE LA GUÍA — CUMPLIMIENTO OBLIGATORIO",
        "Los títulos deben copiarse exactamente y los pasos no pueden fusionarse ni omitirse.",
    ]
    for sequence_index, sequence in enumerate(manifest, start=1):
        blocks.append(
            f"\nSECUENCIA {sequence_index} | TÍTULO EXACTO: {sequence['section_title']} | "
            f"CANTIDAD EXACTA: {sequence['expected_count']}"
        )
        for item_index, item in enumerate(sequence["items"], start=1):
            blocks.append(f"{item_index}. {item}")
    return "\n".join(blocks)


def render_exact_guide_structure(guides: list[dict], selected_guide_index: int) -> str:
    manifest = build_guide_structure_manifest(guides)
    selected = next(
        (item for item in manifest if item["guide_index"] == selected_guide_index),
        None,
    )
    if not selected or not selected["sections"]:
        return "No se pudo extraer una estructura local exacta de la guía seleccionada."

    lines = [
        "TÍTULOS Y ORDEN BLOQUEADOS POR LA GUÍA PRINCIPAL",
        "Debes devolver exactamente estas secciones, con estos títulos y en este orden.",
    ]
    for section in selected["sections"]:
        lines.append(f"\n{section['order']}. TÍTULO EXACTO: {section['title']}")
        mode = section.get("step_mode", "none")
        if mode == "fixed":
            lines.append(f"   PASOS FIJOS OBLIGATORIOS: {len(section['numbered_items'])}")
            for index, item in enumerate(section["numbered_items"], start=1):
                lines.append(f"   {index}. {item}")
        elif mode == "dynamic_list":
            lines.append(
                "   PROCEDIMIENTO DINÁMICO: crear un paso independiente por cada "
                "acción ACC-XXXX del origen; no resumir ni limitar la cantidad."
            )
        elif mode == "dynamic_table":
            lines.append(
                "   PROCEDIMIENTO DINÁMICO EN TABLA: crear una fila independiente "
                "por cada acción ACC-XXXX del origen."
            )
    return "\n".join(lines)

def _normalized_title_tokens(text: str) -> set[str]:
    normalized = re.sub(r"[^a-záéíóúñ0-9]+", " ", str(text).lower())
    ignored = {"de", "del", "la", "las", "el", "los", "y", "para", "a", "en"}
    return {token for token in normalized.split() if token and token not in ignored}


def _title_similarity(expected: str, actual: str) -> float:
    expected_normal = re.sub(r"\s+", " ", expected).strip().casefold()
    actual_normal = re.sub(r"\s+", " ", actual).strip().casefold()
    if expected_normal == actual_normal:
        return 2.0
    if expected_normal in actual_normal or actual_normal in expected_normal:
        return 1.3
    expected_tokens = _normalized_title_tokens(expected)
    actual_tokens = _normalized_title_tokens(actual)
    union = expected_tokens | actual_tokens
    return len(expected_tokens & actual_tokens) / len(union) if union else 0.0


def _best_section_index(expected_title: str, sections: list, used: set[int]) -> int | None:
    best_index: int | None = None
    best_score = 0.0
    for index, section in enumerate(sections):
        if index in used:
            continue
        score = _title_similarity(expected_title, getattr(section, "title", ""))
        if score > best_score:
            best_score = score
            best_index = index
    return best_index if best_score >= 0.18 else None


def lock_analysis_to_selected_guide(
    analysis: GuideAnalysis,
    guides: list[dict],
) -> GuideAnalysis:
    """Obliga al análisis a utilizar todos los títulos exactos de la guía."""

    manifest = build_guide_structure_manifest(guides)
    selected = next(
        (item for item in manifest if item["guide_index"] == analysis.selected_guide_index),
        None,
    )
    if not selected or not selected["sections"]:
        return analysis

    original_sections = list(analysis.sections)
    used: set[int] = set()
    locked_sections: list[SectionAssessment] = []

    for expected in selected["sections"]:
        match_index = _best_section_index(expected["title"], original_sections, used)
        matched = original_sections[match_index] if match_index is not None else None
        if match_index is not None:
            used.add(match_index)

        criteria = list(matched.criteria) if matched else []
        title_rule = f"Conservar exactamente el título '{expected['title']}'."
        if title_rule not in criteria:
            criteria.insert(0, title_rule)

        mode = expected.get("step_mode", "none")
        if mode == "fixed":
            count_rule = (
                f"Incluir exactamente {len(expected['numbered_items'])} pasos "
                "independientes en el mismo orden de la guía."
            )
            if count_rule not in criteria:
                criteria.append(count_rule)
        elif mode in {"dynamic_list", "dynamic_table"}:
            dynamic_rule = (
                "Construir el procedimiento completo con una actividad independiente "
                "por cada acción operativa verificable del origen; no resumir, fusionar "
                "ni limitar artificialmente la cantidad."
            )
            if dynamic_rule not in criteria:
                criteria.append(dynamic_rule)

        if mode == "dynamic_table":
            output_format = "tabla"
        elif mode in {"fixed", "dynamic_list"}:
            output_format = "lista"
        else:
            output_format = matched.output_format if matched else "parrafo"

        locked_sections.append(
            SectionAssessment(
                order=expected["order"],
                title=expected["title"],
                guide_instruction=(
                    expected["instruction"]
                    or (matched.guide_instruction if matched else "")
                ),
                criteria=criteria,
                required=matched.required if matched else True,
                status=matched.status if matched else "faltante",
                evidence=list(matched.evidence) if matched else [],
                draft_content=matched.draft_content if matched else "",
                output_format=output_format,
                missing_questions=list(matched.missing_questions) if matched else [],
            )
        )

    return analysis.model_copy(update={"sections": locked_sections})

def _collect_generated_numbered_items(section: FinalSection | None) -> list[str]:
    if section is None:
        return []

    candidates: list[str] = []
    for item in section.numbered_items:
        value = _strip_list_prefix(item)
        if value:
            candidates.append(value)

    for item in section.bullets:
        if NUMBERED_GUIDE_LINE_RE.match(item) or re.match(r"^\s*\d+[.)\-:]", item):
            value = _strip_list_prefix(item)
            if value:
                candidates.append(value)

    for paragraph in section.paragraphs:
        for line in str(paragraph).splitlines():
            if NUMBERED_GUIDE_LINE_RE.match(line):
                value = _strip_list_prefix(line)
                if value:
                    candidates.append(value)

    return list(dict.fromkeys(candidates))


def _remove_embedded_numbered_lines(paragraphs: list[str]) -> list[str]:
    cleaned: list[str] = []
    for paragraph in paragraphs:
        lines = [line for line in str(paragraph).splitlines() if line.strip()]
        remaining = [
            line for line in lines
            if not NUMBERED_GUIDE_LINE_RE.match(line.strip())
            and not re.match(r"^\s*\d+[.)\-:]", line.strip())
        ]
        if remaining:
            cleaned.append("\n".join(remaining))
    return cleaned


def lock_final_document_to_selected_guide(
    final_document: FinalDocument,
    analysis: GuideAnalysis,
    guides: list[dict],
) -> FinalDocument:
    """Impone títulos y orden sin borrar procedimientos dinámicos.

    Los pasos fijos se limitan a la cantidad explícita de la guía. Las secciones
    dinámicas PASO A PASO/PROCEDIMIENTO conservan su contenido y posteriormente
    se reconstruyen de forma determinista desde el inventario de acciones.
    """

    manifest = build_guide_structure_manifest(guides)
    selected = next(
        (item for item in manifest if item["guide_index"] == analysis.selected_guide_index),
        None,
    )
    if not selected or not selected["sections"]:
        return normalize_final_document_structure(final_document)

    generated_sections = list(final_document.sections)
    analysis_sections = list(analysis.sections)
    used_generated: set[int] = set()
    used_analysis: set[int] = set()
    locked_sections: list[FinalSection] = []
    warnings = list(final_document.warnings)

    for expected in selected["sections"]:
        generated_index = _best_section_index(
            expected["title"], generated_sections, used_generated
        )
        generated = generated_sections[generated_index] if generated_index is not None else None
        if generated_index is not None:
            used_generated.add(generated_index)

        analysis_index = _best_section_index(
            expected["title"], analysis_sections, used_analysis
        )
        assessed = analysis_sections[analysis_index] if analysis_index is not None else None
        if analysis_index is not None:
            used_analysis.add(analysis_index)

        paragraphs = list(generated.paragraphs) if generated else []
        bullets = list(generated.bullets) if generated else []
        tables = list(generated.tables) if generated else []
        source_basis = list(generated.source_basis) if generated else []

        if not paragraphs and assessed and assessed.draft_content.strip():
            paragraphs = [assessed.draft_content.strip()]
        if not source_basis and assessed:
            source_basis = list(assessed.evidence)

        expected_items = list(expected["numbered_items"])
        generated_items = _collect_generated_numbered_items(generated)
        mode = expected.get("step_mode", "none")

        if mode == "fixed":
            expected_count = len(expected_items)
            numbered_items = [
                generated_items[index] if index < len(generated_items) and generated_items[index]
                else expected_items[index]
                for index in range(expected_count)
            ]
            if len(generated_items) > expected_count:
                warnings.append(
                    f"Se eliminaron {len(generated_items) - expected_count} paso(s) "
                    f"adicional(es) de '{expected['title']}' porque la guía fija "
                    f"exactamente {expected_count}."
                )
            elif len(generated_items) < expected_count:
                warnings.append(
                    f"La sección '{expected['title']}' se completó hasta los "
                    f"{expected_count} pasos fijos exigidos por la guía."
                )
            bullets = [
                item for item in bullets
                if not NUMBERED_GUIDE_LINE_RE.match(item)
                and not re.match(r"^\s*\d+[.)\-:]", item)
            ]
            paragraphs = _remove_embedded_numbered_lines(paragraphs)

        elif mode == "dynamic_list":
            # Nunca eliminar los pasos de una sección dinámica. La cobertura final
            # se aplicará usando source["video_actions"].
            numbered_items = generated_items
            bullets = [
                item for item in bullets
                if not NUMBERED_GUIDE_LINE_RE.match(item)
                and not re.match(r"^\s*\d+[.)\-:]", item)
            ]
            paragraphs = _remove_embedded_numbered_lines(paragraphs)

        elif mode == "dynamic_table":
            numbered_items = []

        else:
            numbered_items = []
            if generated_items:
                warnings.append(
                    f"Se eliminaron {len(generated_items)} paso(s) no autorizados "
                    f"de la sección no procedimental '{expected['title']}'."
                )
            bullets = [
                item for item in bullets
                if not NUMBERED_GUIDE_LINE_RE.match(item)
                and not re.match(r"^\s*\d+[.)\-:]", item)
            ]
            paragraphs = _remove_embedded_numbered_lines(paragraphs)

        locked_sections.append(
            FinalSection(
                order=expected["order"],
                title=expected["title"],
                paragraphs=paragraphs,
                bullets=bullets,
                numbered_items=numbered_items,
                tables=tables,
                source_basis=source_basis,
            )
        )

    exact_titles = [section["title"] for section in selected["sections"]]
    remapped_validation: list[ValidationItem] = []
    for item in final_document.validation:
        best_title = max(
            exact_titles,
            key=lambda title: _title_similarity(title, item.section_title),
            default=item.section_title,
        )
        if _title_similarity(best_title, item.section_title) >= 0.18:
            remapped_validation.append(item.model_copy(update={"section_title": best_title}))

    locked = final_document.model_copy(
        update={
            "sections": locked_sections,
            "validation": remapped_validation,
            "warnings": list(dict.fromkeys(warnings)),
        }
    )
    return normalize_final_document_structure(locked)

def document_fidelity_issues(
    final_document: FinalDocument,
    manifest: list[dict],
    selected_guide_index: int,
) -> list[str]:
    """Valida títulos exactos, orden y cantidad exacta de pasos."""

    full_manifest = build_guide_structure_manifest([
        {"name": "selected", "text": ""}
    ]) if False else None  # mantiene la firma sin recalcular guías aquí
    issues: list[str] = []
    relevant_sequences = [
        item for item in manifest if item["guide_index"] == selected_guide_index
    ]
    for sequence in relevant_sequences:
        matching = next(
            (
                section for section in final_document.sections
                if re.sub(r"\s+", " ", section.title).strip().casefold()
                == re.sub(r"\s+", " ", sequence["section_title"]).strip().casefold()
            ),
            None,
        )
        if matching is None:
            issues.append(f"Falta el título exacto '{sequence['section_title']}'.")
            continue
        if len(matching.numbered_items) != sequence["expected_count"]:
            issues.append(
                f"La sección '{sequence['section_title']}' exige exactamente "
                f"{sequence['expected_count']} pasos y contiene {len(matching.numbered_items)}."
            )
    return issues


MISSING_ACTION_VALUES = {
    "",
    "null",
    "none",
    "n a",
    "na",
    "s d",
    "sd",
    "no aplica",
    "no disponible",
    "sin dato",
    "sin datos",
    "sin informacion",
    "sin información",
    "no identificado",
    "no identificada",
    "no especificado",
    "no especificada",
    "no determinado",
    "no determinada",
    "desconocido",
    "desconocida",
    "por confirmar",
    "pendiente",
    "vacio",
    "vacío",
    "no se evidencia",
    "no se observa",
    "ninguna",
    "ninguno",
    "sin validacion",
    "sin validación",
    "sin resultado",
    "no requerida",
    "no requerido",
}


def _clean_action_value(value: object) -> str:
    """Normaliza campos del inventario y elimina marcadores sin información."""

    if value is None:
        return ""

    cleaned = re.sub(r"\s+", " ", str(value)).strip()
    if not cleaned:
        return ""

    normalized = re.sub(
        r"""[\s"'`´“”‘’.,;:()\[\]{}<>_/\\-]+""",
        " ",
        cleaned.casefold(),
    ).strip()

    if normalized in MISSING_ACTION_VALUES:
        return ""

    placeholder_tokens = {"null", "none", "na", "n", "a", "sd"}
    tokens = normalized.split()
    if tokens and all(token in placeholder_tokens for token in tokens):
        return ""

    return cleaned


def _video_action_from_mapping(payload: dict) -> VideoAction | None:
    try:
        action_text = _clean_action_value(payload.get("action"))
        if len(action_text) < MIN_ACTION_TEXT_CHARS:
            return None
        return VideoAction(
            action_id=_clean_action_value(payload.get("action_id")),
            timestamp_start=_clean_action_value(payload.get("timestamp_start")),
            timestamp_end=_clean_action_value(payload.get("timestamp_end")),
            actor=_clean_action_value(payload.get("actor")),
            system=_clean_action_value(payload.get("system")),
            location_path=_clean_action_value(payload.get("location_path")),
            action=action_text,
            interface_element=_clean_action_value(payload.get("interface_element")),
            data_handled=_clean_action_value(payload.get("data_handled")),
            validation=_clean_action_value(payload.get("validation")),
            result=_clean_action_value(payload.get("result")),
            evidence=_clean_action_value(payload.get("evidence")),
            uncertainty=_clean_action_value(payload.get("uncertainty")),
        )
    except Exception:
        return None


def _source_video_actions(source: dict) -> list[VideoAction]:
    raw_actions = source.get("video_actions") or []
    if not raw_actions:
        transcript_payload = source.get("video_transcript") or {}
        if isinstance(transcript_payload, dict):
            raw_actions = transcript_payload.get("actions") or []

    actions: list[VideoAction] = []
    for payload in raw_actions:
        if isinstance(payload, VideoAction):
            action = payload
        elif isinstance(payload, dict):
            action = _video_action_from_mapping(payload)
        else:
            action = None
        if action is not None:
            actions.append(action)

    if not actions:
        actions = _parse_action_inventory_from_source_text(str(source.get("text") or ""))

    normalized: list[VideoAction] = []
    seen: set[str] = set()
    for index, action in enumerate(actions, start=1):
        key = re.sub(
            r"[^a-záéíóúñ0-9]+",
            " ",
            f"{action.timestamp_start} {action.action} {action.interface_element}".lower(),
        ).strip()
        if not key or key in seen:
            continue
        seen.add(key)
        normalized.append(
            action.model_copy(update={"action_id": f"ACC-{len(normalized)+1:04d}"})
        )
    return normalized


def _sentence(value: str) -> str:
    cleaned = _clean_action_value(value)
    if not cleaned:
        return ""
    return cleaned if cleaned.endswith((".", "!", "?")) else cleaned + "."


def _lower_first(value: str) -> str:
    cleaned = _clean_action_value(value)
    if not cleaned:
        return ""
    return cleaned[0].lower() + cleaned[1:]


def _capitalize_first(value: str) -> str:
    cleaned = _clean_action_value(value)
    if not cleaned:
        return ""
    return cleaned[0].upper() + cleaned[1:]


def _normalized_action_key(value: str) -> str:
    return re.sub(
        r"[^a-záéíóúüñ0-9]+",
        " ",
        _clean_action_value(value).casefold(),
    ).strip()


def _strip_actor_from_action(action_text: str, actor: str) -> str:
    """Evita repetir el rol dentro de la acción operativa."""

    action = _clean_action_value(action_text)
    actor_value = _clean_action_value(actor)
    if not action or not actor_value:
        return action

    prefixes = [
        actor_value,
        f"el {actor_value}",
        f"la {actor_value}",
        f"rol {actor_value}",
        f"el rol {actor_value}",
        f"la persona responsable del rol {actor_value}",
    ]
    lowered = action.casefold()
    for prefix in prefixes:
        if lowered.startswith(prefix.casefold() + " "):
            return action[len(prefix):].strip(" :-")
    return action


def _remove_leading_preposition(value: str) -> str:
    cleaned = _clean_action_value(value)
    return re.sub(r"^(?:de|del|en|la|el)\s+", "", cleaned, flags=re.I).strip()


def _folder_name(location_path: str) -> str:
    value = _clean_action_value(location_path)
    if not value:
        return ""
    value = re.sub(r"^carpeta\s+", "", value, flags=re.I).strip(" :-")
    value = re.sub(r"^(?:de|del)\s+", "", value, flags=re.I).strip()
    return _capitalize_first(value)


def _with_article(value: str) -> str:
    """Agrega un artículo cuando mejora la naturalidad del título."""

    cleaned = _clean_action_value(value)
    if not cleaned:
        return ""
    if re.match(r"^(?:el|la|los|las|un|una|unos|unas)\s+", cleaned, flags=re.I):
        return cleaned

    first_word = re.split(r"\s+", cleaned, maxsplit=1)[0].casefold()
    feminine_singular = {
        "acción", "actividad", "aplicación", "carpeta", "celda", "clave",
        "contraseña", "cuenta", "evidencia", "factura", "fecha", "fila",
        "información", "lista", "opción", "pantalla", "plataforma", "ruta",
        "solicitud", "tabla", "ventana",
    }
    if first_word.endswith("as"):
        article = "las"
    elif first_word.endswith(("os", "es")) or first_word in {
        "datos", "totales", "registros", "archivos", "campos", "giros",
    }:
        article = "los"
    elif first_word in feminine_singular or first_word.endswith(("ción", "sión", "dad")):
        article = "la"
    else:
        article = "el"
    return f"{article} {cleaned}"


TITLE_VERB_RULES = (
    (("comparar", "validar", "verificar", "cotejar", "conciliar"), "Validar"),
    (("guardar", "almacenar"), "Guardar"),
    (("descargar", "exportar"), "Descargar"),
    (("adjuntar", "cargar archivo"), "Adjuntar"),
    (("abrir",), "Abrir"),
    (("seleccionar", "elegir", "marcar", "desmarcar"), "Seleccionar"),
    (("hacer clic derecho", "clic derecho"), "Abrir las opciones"),
    (("hacer clic", "pulsar", "presionar"), "Seleccionar"),
    (("ingresar", "acceder", "iniciar sesión"), "Ingresar"),
    (("buscar", "ubicar", "localizar"), "Buscar"),
    (("diligenciar", "registrar", "escribir", "digitar", "modificar"), "Registrar"),
    (("enviar",), "Enviar"),
    (("aprobar",), "Aprobar"),
    (("crear",), "Crear"),
    (("eliminar",), "Eliminar"),
    (("filtrar",), "Filtrar"),
    (("consultar", "visualizar", "revisar"), "Consultar"),
    (("cerrar",), "Cerrar"),
)


IMPERATIVE_VERBS = {
    "abrir": "abre",
    "acceder": "accede",
    "ingresar": "ingresa",
    "iniciar": "inicia",
    "seleccionar": "selecciona",
    "elegir": "elige",
    "marcar": "marca",
    "desmarcar": "desmarca",
    "hacer": "haz",
    "pulsar": "pulsa",
    "presionar": "presiona",
    "ubicar": "ubica",
    "buscar": "busca",
    "localizar": "localiza",
    "consultar": "consulta",
    "diligenciar": "diligencia",
    "registrar": "registra",
    "escribir": "escribe",
    "digitar": "digita",
    "modificar": "modifica",
    "adjuntar": "adjunta",
    "cargar": "carga",
    "descargar": "descarga",
    "guardar": "guarda",
    "enviar": "envía",
    "aprobar": "aprueba",
    "validar": "valida",
    "verificar": "verifica",
    "confirmar": "confirma",
    "comparar": "compara",
    "cotejar": "coteja",
    "conciliar": "concilia",
    "crear": "crea",
    "eliminar": "elimina",
    "cerrar": "cierra",
    "filtrar": "filtra",
    "visualizar": "visualiza",
    "revisar": "revisa",
    "exportar": "exporta",
}


def _action_category(action_text: str) -> str:
    key = _normalized_action_key(action_text)
    for phrases, category in (
        (("guardar", "almacenar"), "save"),
        (("comparar", "validar", "verificar", "cotejar", "conciliar"), "validate"),
        (("descargar", "exportar"), "download"),
        (("adjuntar", "cargar archivo"), "attach"),
        (("abrir",), "open"),
        (("seleccionar", "elegir", "marcar", "desmarcar"), "select"),
        (("hacer clic", "clic derecho", "pulsar", "presionar"), "click"),
        (("buscar", "ubicar", "localizar", "filtrar"), "search"),
        (("diligenciar", "registrar", "escribir", "digitar", "modificar"), "register"),
    ):
        if any(phrase in key for phrase in phrases):
            return category
    return "generic"


def _title_verb(action_text: str) -> str:
    key = _normalized_action_key(action_text)
    for phrases, replacement in TITLE_VERB_RULES:
        if any(phrase in key for phrase in phrases):
            return replacement

    match = re.match(r"^[A-Za-zÁÉÍÓÚÜÑáéíóúüñ]+", _clean_action_value(action_text))
    return _capitalize_first(match.group(0)) if match else "Ejecutar"


def _build_step_title(action: VideoAction) -> str:
    base_action = _strip_actor_from_action(action.action, action.actor)
    data = _clean_action_value(action.data_handled)
    element = _clean_action_value(action.interface_element)
    system = _clean_action_value(action.system)
    category = _action_category(base_action)
    verb = _title_verb(base_action)

    if category == "validate" and data:
        title = f"{verb} {_with_article(data)}"
        if system:
            title += f" en {system}"
        return title

    if data:
        return f"{verb} {_with_article(data)}"

    if element and _normalized_action_key(element) not in {
        "ventana de guardado", "celdas de excel", "espacio en blanco",
    }:
        return f"{verb} {_with_article(element)}"

    cleaned_action = _clean_action_value(base_action).rstrip(".:")
    if cleaned_action:
        return _capitalize_first(cleaned_action)
    return "Ejecutar la actividad"


def _to_imperative(action_text: str) -> str:
    cleaned = _clean_action_value(action_text).rstrip(".:")
    if not cleaned:
        return ""

    special = (
        (r"^hacer clic derecho\b", "haz clic derecho"),
        (r"^hacer clic izquierdo\b", "haz clic izquierdo"),
        (r"^hacer clic\b", "haz clic"),
    )
    for pattern, replacement in special:
        if re.search(pattern, cleaned, flags=re.I):
            return re.sub(pattern, replacement, cleaned, count=1, flags=re.I)

    match = re.match(r"^([A-Za-zÁÉÍÓÚÜÑáéíóúüñ]+)(\b.*)$", cleaned)
    if not match:
        return _lower_first(cleaned)
    verb = match.group(1).casefold()
    remainder = match.group(2)
    imperative = IMPERATIVE_VERBS.get(verb)
    if imperative:
        return imperative + remainder
    return _lower_first(cleaned)


def _corporate_location(system: str, path: str, interface: str = "") -> str:
    system_value = _clean_action_value(system)
    path_value = _clean_action_value(path)
    interface_value = _clean_action_value(interface)

    if system_value.casefold() == "explorador de archivos":
        if "ventana de guardado" in interface_value.casefold():
            return "En la ventana de guardado del Explorador de archivos"
        if path_value:
            folder = _folder_name(path_value)
            return (
                f"En el Explorador de archivos, dentro de la carpeta «{folder}»"
                if folder else "En el Explorador de archivos"
            )
        return "En el Explorador de archivos"

    if system_value:
        if path_value and not re.search(r"^archivo\b", path_value, flags=re.I):
            return f"En {system_value}, en {_with_article(path_value)}"
        return f"En {system_value}"

    if path_value:
        return f"En {_with_article(path_value)}"
    return ""


def _quoted(value: str) -> str:
    cleaned = _clean_action_value(value)
    return f"«{cleaned}»" if cleaned else ""


def _base_instruction(action: VideoAction) -> str:
    base_action = _strip_actor_from_action(action.action, action.actor)
    category = _action_category(base_action)
    system = _clean_action_value(action.system)
    path = _clean_action_value(action.location_path)
    interface = _clean_action_value(action.interface_element)
    data = _clean_action_value(action.data_handled)
    location = _corporate_location(system, path, interface)
    folder = _folder_name(path)

    if category == "save":
        if system.casefold() == "explorador de archivos":
            if folder:
                return (
                    f"{location}, selecciona la carpeta «{folder}» y guarda el documento."
                )
            return f"{location}, selecciona la ubicación correspondiente y guarda el documento."
        object_text = _with_article(data) if data else "el archivo"
        return _sentence(f"{location}, guarda {object_text}" if location else f"Guarda {object_text}")

    if category == "validate":
        sentences: list[str] = []
        if system and path and re.search(r"^archivo\b", path, flags=re.I):
            file_name = _remove_leading_preposition(path)
            if file_name.casefold() in {"archivo", "archivo de reporte", "archivo del reporte"}:
                sentences.append(f"Abre el archivo del reporte en {system}.")
            else:
                sentences.append(f"Abre {_with_article(file_name)} en {system}.")
        elif system:
            sentences.append(f"Abre el archivo correspondiente en {system}.")

        if interface and "celda" in interface.casefold():
            if data:
                sentences.append(
                    f"Ubica las celdas correspondientes a {_with_article(_quoted(data))} y compara los datos."
                )
            else:
                sentences.append("Ubica las celdas que contienen los valores que se deben revisar y compara los datos.")
        else:
            target = _quoted(data or interface)
            if target:
                sentences.append(f"Ubica {target} y compara la información registrada.")
            else:
                sentences.append(_sentence(_to_imperative(base_action)))
        return " ".join(sentences)

    if category == "download":
        target = _with_article(data) if data else "el archivo"
        if interface:
            sentence = f"{location}, selecciona {_quoted(interface)} para descargar {target}"
        else:
            sentence = f"{location}, descarga {target}" if location else f"Descarga {target}"
        return _sentence(sentence)

    if category == "attach":
        target = _with_article(data) if data else "el archivo requerido"
        if interface:
            sentence = f"{location}, selecciona {_quoted(interface)} y adjunta {target}"
        else:
            sentence = f"{location}, adjunta {target}" if location else f"Adjunta {target}"
        return _sentence(sentence)

    if category == "register":
        if interface and data:
            return _sentence(
                f"{location}, ubica el campo {_quoted(interface)} y registra {_quoted(data)}"
                if location else
                f"Ubica el campo {_quoted(interface)} y registra {_quoted(data)}"
            )
        if data:
            return _sentence(
                f"{location}, registra {_quoted(data)}" if location else f"Registra {_quoted(data)}"
            )

    if category == "search":
        if interface and data:
            return _sentence(
                f"{location}, ingresa {_quoted(data)} en {_with_article(interface)} y ejecuta la búsqueda"
                if location else
                f"Ingresa {_quoted(data)} en {_with_article(interface)} y ejecuta la búsqueda"
            )

    imperative = _to_imperative(base_action)
    if interface and interface.casefold() not in imperative.casefold():
        if "clic" in imperative.casefold():
            imperative += f" sobre {_with_article(interface)}"
        elif category == "select":
            imperative += f" {_with_article(interface)}"
        else:
            imperative += f" mediante {_with_article(interface)}"

    if data and data.casefold() not in imperative.casefold():
        if category in {"open", "select"}:
            imperative += f" correspondiente a {_quoted(data)}"
        else:
            imperative += f" para gestionar {_quoted(data)}"

    sentence = f"{location}, {imperative}" if location else imperative
    return _sentence(_capitalize_first(sentence))


def _validation_instruction(validation: str, data: str) -> str:
    value = _clean_action_value(validation)
    if not value:
        return ""
    key = _normalized_action_key(value)
    data_value = _clean_action_value(data)

    if any(term in key for term in ("coincidencia", "cuadre", "cuadrar", "coincidan")):
        noun = "las cifras" if data_value else "los datos"
        return (
            "Verificación obligatoria: Antes de continuar, debes asegurarte de "
            f"que {noun} coincidan."
        )

    stripped = re.sub(
        r"^(?:verificación|validación|comprobación)\s+(?:de|del|de la|de los|de las)\s+",
        "",
        value,
        flags=re.I,
    ).strip()
    if stripped != value:
        return _sentence(
            "Verificación obligatoria: Antes de continuar, debes comprobar "
            + _lower_first(stripped)
        )

    imperative = _to_imperative(value)
    return _sentence(
        "Verificación obligatoria: Antes de continuar, debes "
        + re.sub(r"^(?:debes?|se debe)\s+", "", imperative, flags=re.I)
    )


def _result_instruction(result: str, data: str) -> str:
    value = _clean_action_value(result)
    if not value:
        return ""
    key = _normalized_action_key(value)

    if "archivo guardado" in key or "documento guardado" in key:
        return (
            "Este paso finaliza una vez confirmes que el archivo se ha guardado "
            "correctamente en la ruta indicada."
        )
    if "coincidencia confirmada" in key or "datos coinciden" in key:
        return (
            "El paso se considera completado cuando confirmes la coincidencia "
            "exacta de los datos."
        )
    if "registro creado" in key or "registro guardado" in key:
        return (
            "El paso se considera completado cuando confirmes que el registro "
            "se creó correctamente."
        )
    if "mensaje" in key:
        return _sentence(
            "El paso se considera completado cuando el sistema muestre "
            + _lower_first(value)
        )
    return _sentence(
        "El paso se considera completado cuando confirmes " + _lower_first(value)
    )


def _evidence_instruction(evidence: str, existing_text: str) -> str:
    value = _clean_action_value(evidence)
    if not value or value.casefold() in existing_text.casefold():
        return ""
    key = _normalized_action_key(value)
    if key.startswith(("se observa", "el video muestra", "captura de", "registro de")):
        return _sentence("Como evidencia de la ejecución, conserva " + _lower_first(value))
    return ""


def _split_step_title_body(text: str) -> tuple[str, str]:
    cleaned = _strip_list_prefix(text)
    lines = [line.strip() for line in cleaned.splitlines() if line.strip()]
    if len(lines) >= 2:
        return lines[0].rstrip(":"), " ".join(lines[1:]).strip()
    if lines and ":" in lines[0]:
        title, body = lines[0].split(":", 1)
        if 2 <= len(title.split()) <= 18:
            return title.strip(), body.strip()
    return "", cleaned


def _action_to_detailed_step(action: VideoAction) -> str:
    """Genera un título operativo y una descripción empresarial natural."""

    title = _build_step_title(action).rstrip(".:")
    body_parts = [_base_instruction(action)]

    validation = _validation_instruction(action.validation, action.data_handled)
    if validation:
        body_parts.append(validation)

    result = _result_instruction(action.result, action.data_handled)
    if result:
        body_parts.append(result)

    existing = " ".join(body_parts)
    evidence = _evidence_instruction(action.evidence, existing)
    if evidence:
        body_parts.append(evidence)

    uncertainty = _clean_action_value(action.uncertainty)
    if uncertainty:
        body_parts.append(
            _sentence(
                "Información pendiente de validación antes de aprobar el documento: "
                + uncertainty
            )
        )

    description = " ".join(part for part in body_parts if part).strip()
    return f"{title}:\n{description}" if description else f"{title}:"


def _action_to_procedure_row(action: VideoAction) -> list[str]:
    step_text = _action_to_detailed_step(action)
    activity, description = _split_step_title_body(step_text)
    responsible = _clean_action_value(action.actor) or "No identificado en el video"

    evidence_parts = [
        value
        for value in (
            _clean_action_value(action.validation),
            _clean_action_value(action.result),
            _clean_action_value(action.evidence),
        )
        if value
    ]
    evidence = "; ".join(dict.fromkeys(evidence_parts))
    if not evidence:
        evidence = "No se identifica una evidencia específica en el video."

    return [activity or _clean_action_value(action.action), description, responsible, evidence]

def _apply_video_action_coverage(
    final_document: FinalDocument,
    selected_manifest: dict,
    source: dict,
) -> FinalDocument:
    actions = _source_video_actions(source)
    if not actions:
        dynamic_sections = [
            section for section in selected_manifest.get("sections", [])
            if section.get("step_mode") in {"dynamic_list", "dynamic_table"}
        ]
        if dynamic_sections and source.get("is_video"):
            raise AppError(
                "La transcripción del video no contiene un inventario de acciones "
                "operativas. Vuelve a transcribir el video con esta versión antes "
                "de generar el documento."
            )
        return final_document

    sections_by_title = {
        re.sub(r"\s+", " ", section.title).strip().casefold(): section
        for section in final_document.sections
    }
    rebuilt_sections: list[FinalSection] = []
    warnings = list(final_document.warnings)

    for expected in selected_manifest.get("sections", []):
        key = re.sub(r"\s+", " ", expected["title"]).strip().casefold()
        section = sections_by_title.get(key)
        if section is None:
            continue

        mode = expected.get("step_mode", "none")
        if mode == "dynamic_list":
            steps = [_action_to_detailed_step(action) for action in actions]
            steps = [step for step in steps if step.strip()]
            section = section.model_copy(
                update={
                    "numbered_items": steps,
                    "bullets": [],
                    "source_basis": [
                        f"Inventario audiovisual {actions[0].action_id}–{actions[-1].action_id}; "
                        f"{len(actions)} acciones verificables."
                    ],
                }
            )
            warnings.append(
                f"La sección '{section.title}' se construyó con {len(steps)} "
                "pasos, uno por cada acción operativa identificada en el video."
            )

        elif mode == "dynamic_table":
            procedure_table = FinalTable(
                title="",
                headers=["Actividad", "Descripción", "Responsables", "Evidencia"],
                rows=[_action_to_procedure_row(action) for action in actions],
            )
            section = section.model_copy(
                update={
                    "numbered_items": [],
                    "tables": [procedure_table],
                    "source_basis": [
                        f"Inventario audiovisual {actions[0].action_id}–{actions[-1].action_id}; "
                        f"{len(actions)} acciones verificables."
                    ],
                }
            )
            warnings.append(
                f"La sección '{section.title}' se construyó con {len(actions)} "
                "filas, una por cada acción operativa identificada en el video."
            )

        rebuilt_sections.append(section)

    rebuilt_by_order = {section.order: section for section in rebuilt_sections}
    all_sections = [
        rebuilt_by_order.get(section.order, section)
        for section in final_document.sections
    ]
    return normalize_final_document_structure(
        final_document.model_copy(
            update={
                "sections": all_sections,
                "warnings": list(dict.fromkeys(warnings)),
            }
        )
    )


def enforce_fidelity_with_gemini(
    api_key: str,
    model: str,
    guides: list[dict],
    source: dict,
    analysis: GuideAnalysis,
    final_document: FinalDocument,
) -> FinalDocument:
    """Aplica títulos exactos y cobertura determinista de acciones."""

    del api_key, model
    locked = lock_final_document_to_selected_guide(
        final_document=final_document,
        analysis=analysis,
        guides=guides,
    )

    selected_manifest = next(
        (
            item for item in build_guide_structure_manifest(guides)
            if item["guide_index"] == analysis.selected_guide_index
        ),
        None,
    )
    if not selected_manifest:
        return locked

    locked = _apply_video_action_coverage(locked, selected_manifest, source)

    expected_titles = [section["title"] for section in selected_manifest["sections"]]
    actual_titles = [section.title for section in locked.sections]
    if actual_titles != expected_titles:
        raise AppError(
            "El documento no conserva exactamente los títulos y el orden de la guía."
        )

    actions = _source_video_actions(source)
    for expected, actual in zip(selected_manifest["sections"], locked.sections):
        mode = expected.get("step_mode", "none")
        if mode == "fixed":
            expected_count = len(expected["numbered_items"])
            actual_count = len(actual.numbered_items)
            if actual_count != expected_count:
                raise AppError(
                    f"La sección '{expected['title']}' debe contener exactamente "
                    f"{expected_count} pasos fijos y contiene {actual_count}."
                )
        elif mode == "dynamic_list" and actions:
            if len(actual.numbered_items) != len(actions):
                raise AppError(
                    f"Cobertura incompleta en '{expected['title']}': se detectaron "
                    f"{len(actions)} acciones y se generaron {len(actual.numbered_items)} pasos."
                )
            if any(not item.strip() for item in actual.numbered_items):
                raise AppError(
                    f"La sección '{expected['title']}' contiene pasos vacíos."
                )
        elif mode == "dynamic_table" and actions:
            row_count = sum(len(table.rows) for table in actual.tables)
            if row_count != len(actions):
                raise AppError(
                    f"Cobertura incompleta en '{expected['title']}': se detectaron "
                    f"{len(actions)} acciones y se generaron {row_count} filas."
                )
            for table in actual.tables:
                for row in table.rows:
                    if len(row) < 4 or not row[0].strip() or not row[1].strip():
                        raise AppError(
                            f"La tabla de '{expected['title']}' contiene una fila incompleta."
                        )

    return locked

def video_transcription_prompt(filename: str) -> str:
    """Instrucciones para transcripción completa e inventario atómico de acciones."""

    return f"""
Actúa como transcriptor profesional, observador de pantalla y analista de procesos.

ARCHIVO
{filename}

OBJETIVO
Convertir TODO el video en una fuente textual verificable y, adicionalmente,
construir un inventario exhaustivo de acciones operativas que permita redactar
un procedimiento sin resumir ni perder actividades.

REGLAS DE TRANSCRIPCIÓN
1. Transcribe de principio a fin todo el discurso audible y relevante.
2. Mantén el orden cronológico y usa marcas absolutas [HH:MM:SS].
3. Identifica hablantes sin inventar nombres.
4. No resumas, no fusiones intervenciones y no omitas instrucciones, decisiones,
   cifras, fechas, nombres, sistemas, responsables ni controles.
5. Registra como evidencia visual los textos, formularios, tablas, pantallas,
   módulos, rutas, botones, campos, archivos, mensajes y cambios de estado.
6. Cuando algo no sea comprensible, usa [inaudible HH:MM:SS] y regístralo en
   uncertainties. Nunca inventes.

INVENTARIO DE ACCIONES OBLIGATORIO
7. El campo actions es independiente de full_transcript y debe registrar TODAS
   las acciones operativas observadas o explicadas en el tramo.
8. Cada acción diferenciable debe ser un elemento independiente. Un clic, abrir
   un menú, seleccionar una pestaña, buscar un registro, aplicar un filtro,
   diligenciar un campo, adjuntar un archivo, guardar, enviar, validar, aprobar,
   descargar, confirmar un mensaje o revisar un resultado son acciones distintas.
9. Está prohibido escribir una sola acción como “ingresa al sistema y realiza el
   proceso”. Debes descomponerla en todas las acciones que realmente ocurren.
10. No omitas acciones por ser repetitivas, sencillas, obvias o de corta duración.
11. No combines acciones usando “y luego”, “posteriormente” o expresiones
    equivalentes. Si hay dos verbos operativos diferentes, normalmente deben ser
    dos elementos de actions.
12. Para cada acción completa, cuando la evidencia lo permita: timestamp_start,
    timestamp_end, actor, system, location_path, action, interface_element,
    data_handled, validation, result, evidence y uncertainty.
    Cuando un campo no tenga información verificable, devuelve una cadena vacía.
    Nunca escribas null, None, N/A, no aplica, por confirmar ni expresiones equivalentes.
13. action debe contener UNA sola acción concreta y verificable, preferiblemente
    iniciada con un verbo en infinitivo y con un objeto específico, por ejemplo:
    "Guardar el reporte de giros" o "Comparar los totales de giros". No incluyas
    al responsable, el sistema ni frases de relleno dentro de action.
14. Los campos validation y result deben contener hechos concretos. Cuando no
    exista una validación o un resultado visible, usa una cadena vacía; nunca
    escribas "Ninguna" ni expresiones equivalentes.
15. Si durante un intervalo no ocurre ninguna actividad operativa y solo existe
    conversación sin instrucciones o ejecución, actions puede quedar vacío para
    ese intervalo. No inventes pasos para aumentar la cantidad.
16. Antes de responder, revisa nuevamente el tramo y confirma que cada acción
    visible o mencionada tenga un registro independiente en actions.

Devuelve únicamente el objeto estructurado solicitado.
""".strip()

def _render_action_inventory(actions: list[VideoAction]) -> str:
    blocks: list[str] = []
    for index, action in enumerate(actions, start=1):
        action_id = action.action_id.strip() or f"ACC-{index:04d}"
        blocks.extend(
            [
                f"[ACCION {action_id}]",
                f"TIEMPO_INICIO: {action.timestamp_start}",
                f"TIEMPO_FIN: {action.timestamp_end}",
                f"RESPONSABLE: {action.actor}",
                f"SISTEMA: {action.system}",
                f"RUTA: {action.location_path}",
                f"ACCION: {action.action}",
                f"ELEMENTO: {action.interface_element}",
                f"DATO: {action.data_handled}",
                f"VALIDACION: {action.validation}",
                f"RESULTADO: {action.result}",
                f"EVIDENCIA: {action.evidence}",
                f"INCERTIDUMBRE: {action.uncertainty}",
                "",
            ]
        )
    return "\n".join(blocks).strip()


def _parse_action_inventory_from_source_text(text: str) -> list[VideoAction]:
    pattern = re.compile(
        r"(?ms)^\[ACCION\s+(ACC-\d+)\]\s*$\n(.*?)(?=^\[ACCION\s+ACC-\d+\]\s*$|^\[TRANSCRIPCIÓN COMPLETA\]|\Z)"
    )
    actions: list[VideoAction] = []
    field_map = {
        "TIEMPO_INICIO": "timestamp_start",
        "TIEMPO_FIN": "timestamp_end",
        "RESPONSABLE": "actor",
        "SISTEMA": "system",
        "RUTA": "location_path",
        "ACCION": "action",
        "ELEMENTO": "interface_element",
        "DATO": "data_handled",
        "VALIDACION": "validation",
        "RESULTADO": "result",
        "EVIDENCIA": "evidence",
        "INCERTIDUMBRE": "uncertainty",
    }
    for match in pattern.finditer(str(text or "")):
        payload: dict[str, str] = {"action_id": match.group(1)}
        for line in match.group(2).splitlines():
            if ":" not in line:
                continue
            key, value = line.split(":", 1)
            target = field_map.get(key.strip().upper())
            if target:
                payload[target] = value.strip()
        action = _video_action_from_mapping(payload)
        if action is not None:
            actions.append(action)
    return actions


def render_video_transcript_as_source(
    transcript: VideoTranscript,
    video_name: str,
) -> str:
    """Convierte transcripción y acciones en un origen editable y trazable."""

    blocks = [
        f"[DOCUMENTO DE ORIGEN GENERADO DESDE VIDEO: {video_name}]",
        f"[IDIOMA DETECTADO] {transcript.detected_language}",
    ]

    if transcript.duration_estimate.strip():
        blocks.append(f"[DURACIÓN APROXIMADA] {transcript.duration_estimate}")

    if transcript.speakers:
        blocks.append("[HABLANTES]\n- " + "\n- ".join(transcript.speakers))

    blocks.append(
        f"[TOTAL DE ACCIONES OPERATIVAS] {len(transcript.actions)}\n"
        "[INVENTARIO DE ACCIONES OPERATIVAS — UNA ACCIÓN POR REGISTRO]\n"
        + (_render_action_inventory(transcript.actions) or "[Sin acciones detectadas]")
    )

    blocks.append(
        "[TRANSCRIPCIÓN COMPLETA]\n"
        + (transcript.full_transcript.strip() or "[Sin transcripción]")
    )

    if transcript.visual_evidence:
        blocks.append("[EVIDENCIA VISUAL]\n- " + "\n- ".join(transcript.visual_evidence))
    if transcript.key_facts:
        blocks.append("[HECHOS CLAVE VERIFICABLES]\n- " + "\n- ".join(transcript.key_facts))
    if transcript.uncertainties:
        blocks.append(
            "[FRAGMENTOS QUE REQUIEREN CONFIRMACIÓN]\n- "
            + "\n- ".join(transcript.uncertainties)
        )

    return "\n\n".join(blocks).strip()

def _file_state_name(file_info) -> str:
    state = getattr(file_info, "state", None)
    if state is None:
        return ""
    name = getattr(state, "name", None)
    return str(name or state).upper()


def _close_drive_service(service) -> None:
    try:
        http = getattr(service, "_http", None)
        if http is not None and hasattr(http, "close"):
            http.close()
    except Exception:
        pass


def _drive_video_extension(filename: str, mime_type: str) -> str:
    extension = file_extension(filename)
    if extension in VIDEO_EXTENSIONS:
        return extension

    by_mime = {
        "video/mp4": "mp4",
        "video/mpeg": "mpeg",
        "video/quicktime": "mov",
        "video/x-msvideo": "avi",
        "video/x-flv": "flv",
        "video/webm": "webm",
        "video/x-ms-wmv": "wmv",
        "video/3gpp": "3gp",
    }
    return by_mime.get(str(mime_type).lower(), "")


def download_drive_video_to_path(
    drive_reference: str,
    destination_dir: Path,
    progress_callback=None,
) -> tuple[dict, Path]:
    """Descarga un video de Drive por bloques directamente al disco temporal.

    No guarda el video completo en RAM. El archivo debe haberse compartido como
    lector con el correo de la cuenta de servicio configurada en Secrets.
    """

    file_id = extract_drive_file_id(drive_reference)
    service = build_drive_service()

    def report(message: str) -> None:
        if progress_callback is not None:
            progress_callback(message)

    try:
        report("Consultando el archivo en Google Drive")
        metadata = (
            service.files()
            .get(
                fileId=file_id,
                fields=(
                    "id,name,size,mimeType,webViewLink,modifiedTime,"
                    "capabilities(canDownload)"
                ),
                supportsAllDrives=True,
            )
            .execute()
        )

        filename = safe_filename(str(metadata.get("name") or f"video_{file_id}"))
        mime_type = str(metadata.get("mimeType") or "application/octet-stream")
        extension = _drive_video_extension(filename, mime_type)

        if mime_type == "application/vnd.google-apps.vid":
            raise AppError(
                "El enlace corresponde a un archivo de Google Vids. Descárgalo "
                "o expórtalo como MP4 y vuelve a subirlo a Drive."
            )

        if not extension or not mime_type.lower().startswith("video/"):
            raise AppError(
                f"El archivo de Drive '{filename}' no parece ser un video compatible."
            )

        can_download = metadata.get("capabilities", {}).get("canDownload")
        if can_download is False:
            raise AppError(
                "El propietario de Drive bloqueó la descarga del video. "
                "Debe permitir la descarga para la cuenta de servicio."
            )

        size_bytes = int(metadata.get("size") or 0)
        if size_bytes <= 0:
            raise AppError(
                "Drive no informó el tamaño del archivo. Verifica que sea un video "
                "normal almacenado en Drive y no un acceso directo."
            )
        if size_bytes > MAX_DRIVE_VIDEO_SIZE_BYTES:
            raise AppError(
                f"El video pesa {format_file_size(size_bytes)} y supera el límite "
                f"seguro de {MAX_DRIVE_VIDEO_SIZE_MB} MB para Gemini."
            )

        if not filename.lower().endswith(f".{extension}"):
            filename = f"{filename}.{extension}"
        local_path = destination_dir / filename

        request = service.files().get_media(
            fileId=file_id,
            supportsAllDrives=True,
        )
        started_at = time.monotonic()

        with local_path.open("wb") as output:
            downloader = MediaIoBaseDownload(
                output,
                request,
                chunksize=DRIVE_DOWNLOAD_CHUNK_BYTES,
            )
            done = False
            while not done:
                if time.monotonic() - started_at > DRIVE_DOWNLOAD_TIMEOUT_SECONDS:
                    raise AppError(
                        "La descarga desde Google Drive superó el tiempo máximo "
                        "permitido. Prueba con una conexión más estable o un video menor."
                    )

                status, done = downloader.next_chunk(num_retries=3)
                if status is not None:
                    downloaded = int(getattr(status, "resumable_progress", 0) or 0)
                    progress = int(float(status.progress()) * 100)
                    report(
                        f"Descargando desde Drive: {progress}% "
                        f"({format_file_size(downloaded)} de {format_file_size(size_bytes)})"
                    )

        actual_size = local_path.stat().st_size if local_path.exists() else 0
        if actual_size <= 0:
            raise AppError("Google Drive no entregó contenido descargable.")
        if actual_size != size_bytes:
            report(
                "Drive informó un tamaño diferente al descargado; se continuará "
                "con el archivo recibido."
            )

        metadata.update(
            {
                "file_id": file_id,
                "name": filename,
                "size_bytes": actual_size,
                "mime_type": MIME_TYPES.get(extension, mime_type),
                "extension": extension,
                "drive_reference": drive_reference,
            }
        )
        return metadata, local_path

    except HttpError as error:
        status_code = getattr(getattr(error, "resp", None), "status", None)
        account_email = drive_service_account_email()
        if status_code in {403, 404}:
            detail = (
                f" Comparte el archivo o la carpeta como Lector con {account_email}."
                if account_email
                else " Configura primero la cuenta de servicio de Drive."
            )
            raise AppError(
                "La cuenta de servicio no puede acceder al video de Google Drive."
                + detail
            ) from error
        raise AppError(
            f"Google Drive rechazó la descarga. Código HTTP: {status_code or 'desconocido'}."
        ) from error
    finally:
        _close_drive_service(service)


def _video_model_candidates(configured_model: str) -> list[str]:
    """Modelos actuales con entrada de video, sin modelos retirados."""

    secret_model = read_secret("GEMINI_VIDEO_MODEL", VIDEO_MODEL_DEFAULT).strip()
    candidates = [
        secret_model,
        VIDEO_MODEL_DEFAULT,
        "gemini-3.1-flash-lite",
        configured_model,
    ]
    blocked = {
        "gemini-2.5-flash",
        "gemini-2.0-flash",
        "gemini-2.0-flash-lite",
    }
    result: list[str] = []
    for candidate in candidates:
        normalized = _normalize_model_name(candidate)
        if (
            normalized
            and normalized not in blocked
            and normalized not in result
        ):
            result.append(normalized)
    return result


def _interaction_video_mime(local_path: Path, reported_mime: str) -> str:
    """Normaliza el MIME a los valores aceptados por Interactions API."""

    extension = file_extension(local_path.name)
    official = {
        "mp4": "video/mp4",
        "mpeg": "video/mpeg",
        "mov": "video/mov",
        "avi": "video/avi",
        "flv": "video/x-flv",
        "mpg": "video/mpg",
        "webm": "video/webm",
        "wmv": "video/wmv",
        "3gp": "video/3gpp",
        "3gpp": "video/3gpp",
    }
    return official.get(extension, str(reported_mime).strip().lower())


def _gemini_video_duration_seconds(file_info) -> float | None:
    """Obtiene la duración reportada por Files API, por ejemplo ``3600.5s``."""

    metadata = getattr(file_info, "video_metadata", None)
    raw_duration = getattr(metadata, "video_duration", None) if metadata else None
    if raw_duration is None:
        return None

    value = str(raw_duration).strip().lower()
    if value.endswith("s"):
        value = value[:-1]
    try:
        seconds = float(value)
        return seconds if seconds > 0 else None
    except (TypeError, ValueError):
        return None


def _interaction_output_text(interaction) -> str:
    """Obtiene el texto de una respuesta de Interactions API."""

    output_text = getattr(interaction, "output_text", None)
    if output_text:
        return str(output_text).strip()

    # Respaldo defensivo para cambios menores del SDK.
    outputs = getattr(interaction, "outputs", None) or []
    fragments: list[str] = []
    for output in outputs:
        text = getattr(output, "text", None)
        if text:
            fragments.append(str(text))
        elif isinstance(output, dict) and output.get("text"):
            fragments.append(str(output["text"]))
    return "\n".join(fragments).strip()


def _video_schema_prompt(prompt: str, schema: dict) -> str:
    """Añade un esquema de respaldo cuando el endpoint rechaza response_format."""

    schema_text = json.dumps(schema, ensure_ascii=False, separators=(",", ":"))
    return (
        prompt
        + "\n\nFORMATO JSON OBLIGATORIO\n"
        + "Devuelve solo un objeto JSON válido, sin Markdown ni comentarios, "
          "que cumpla exactamente este esquema: "
        + schema_text
    )


def _call_video_interaction(
    client,
    model: str,
    file_uri: str,
    mime_type: str,
    prompt: str,
    schema_model: type[BaseModel],
) -> BaseModel:
    """Transcribe video con Interactions API y fallback sin esquema nativo."""

    schema = _inline_and_clean_json_schema(schema_model)
    video_input = {
        "type": "video",
        "uri": file_uri,
        "mime_type": mime_type,
        # Reduce el consumo de contexto en videos largos. Para video general,
        # Gemini recomienda resolución baja o media.
        "resolution": "low",
    }

    try:
        interaction = client.interactions.create(
            model=model,
            input=[
                video_input,
                {"type": "text", "text": prompt},
            ],
            response_format={
                "type": "text",
                "mime_type": "application/json",
                "schema": schema,
            },
        )
        output_text = _interaction_output_text(interaction)
        if not output_text:
            raise AppError("Gemini no devolvió texto al analizar el video.")
        return schema_model.model_validate_json(_clean_json_text(output_text))

    except AppError:
        raise
    except Exception as first_error:
        first_message = str(first_error).upper()
        schema_or_argument_error = (
            "400" in first_message
            or "INVALID_ARGUMENT" in first_message
            or "RESPONSE_FORMAT" in first_message
            or "SCHEMA" in first_message
        )
        if not schema_or_argument_error:
            raise

        # Algunos modelos/endpoints aceptan el video pero rechazan el esquema
        # nativo. Se repite la misma llamada sin response_format y se valida
        # localmente con Pydantic.
        fallback_prompt = _video_schema_prompt(prompt, schema)
        interaction = client.interactions.create(
            model=model,
            input=[
                video_input,
                {"type": "text", "text": fallback_prompt},
            ],
        )
        output_text = _interaction_output_text(interaction)
        if not output_text:
            raise AppError("Gemini no devolvió texto al analizar el video.")
        try:
            return schema_model.model_validate_json(_clean_json_text(output_text))
        except Exception as validation_error:
            raise AppError(
                "Gemini analizó el video, pero la transcripción no llegó en un "
                "JSON válido. Detalle: " + str(validation_error)[:500]
            ) from validation_error


def _probe_local_video_duration_seconds(local_path: Path) -> float | None:
    """Obtiene la duración real con ffprobe sin cargar el video en memoria."""

    ffprobe = shutil.which("ffprobe")
    if not ffprobe:
        return None

    try:
        result = subprocess.run(
            [
                ffprobe,
                "-v",
                "error",
                "-show_entries",
                "format=duration",
                "-of",
                "json",
                str(local_path),
            ],
            capture_output=True,
            text=True,
            timeout=90,
            check=False,
        )
        if result.returncode != 0 or not result.stdout.strip():
            return None

        payload = json.loads(result.stdout)
        duration = float(payload.get("format", {}).get("duration", 0))
        return duration if duration > 0 else None
    except Exception:
        return None


def _video_clock(seconds: float) -> str:
    total = max(0, int(round(seconds)))
    hours, remainder = divmod(total, 3600)
    minutes, secs = divmod(remainder, 60)
    return f"{hours:02d}:{minutes:02d}:{secs:02d}"



def _safe_checkpoint_component(value: str) -> str:
    return re.sub(r"[^A-Za-z0-9_.-]+", "_", str(value or "")).strip("._")


def _video_checkpoint_id(
    checkpoint_key: str,
    model: str,
    duration_seconds: float,
) -> str:
    """Identifica de manera estable el progreso de un video y su configuración."""

    raw = "|".join(
        (
            VIDEO_CHECKPOINT_VERSION,
            str(checkpoint_key or "").strip(),
            _normalize_model_name(model),
            str(int(round(duration_seconds))),
            str(VIDEO_CHUNK_SECONDS),
        )
    )
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()


def _checkpoint_root_dir(checkpoint_id: str) -> Path:
    configured = read_secret(
        "VIDEO_CHECKPOINT_LOCAL_DIR",
        VIDEO_CHECKPOINT_LOCAL_DIR,
    ).strip()
    root = Path(configured or VIDEO_CHECKPOINT_LOCAL_DIR)
    return root / _safe_checkpoint_component(checkpoint_id)


def _checkpoint_filename(start_seconds: float, end_seconds: float) -> str:
    return (
        f"chunk_{int(round(start_seconds)):06d}_"
        f"{int(round(end_seconds)):06d}.json"
    )


def _checkpoint_payload(
    checkpoint_id: str,
    model: str,
    start_seconds: float,
    end_seconds: float,
    transcript: VideoTranscript,
    audited: bool,
    audit_reasons: list[str],
) -> dict:
    return {
        "version": VIDEO_CHECKPOINT_VERSION,
        "checkpoint_id": checkpoint_id,
        "model": _normalize_model_name(model),
        "start_seconds": float(start_seconds),
        "end_seconds": float(end_seconds),
        "audited": bool(audited),
        "audit_reasons": list(audit_reasons),
        "saved_at_epoch": time.time(),
        "transcript": transcript.model_dump(),
    }


def _checkpoint_from_payload(payload: dict) -> tuple[float, float, VideoTranscript] | None:
    try:
        if payload.get("version") != VIDEO_CHECKPOINT_VERSION:
            return None
        start_seconds = float(payload["start_seconds"])
        end_seconds = float(payload["end_seconds"])
        if end_seconds <= start_seconds:
            return None
        transcript = VideoTranscript.model_validate(payload["transcript"])
        if not transcript.full_transcript.strip():
            return None
        return start_seconds, end_seconds, transcript
    except Exception:
        return None


def _write_json_atomic(path: Path, payload: dict) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    temporary.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    temporary.replace(path)


def _drive_checkpoint_folder_id() -> str:
    return read_secret("DRIVE_CHECKPOINT_FOLDER_ID", "").strip()


def _video_parallel_workers() -> int:
    """Permite reducir la concurrencia desde Secrets si la cuota es limitada."""

    raw = read_secret(
        "VIDEO_MAX_PARALLEL_CHUNKS",
        str(VIDEO_MAX_PARALLEL_CHUNKS),
    ).strip()
    try:
        return max(1, min(2, int(raw)))
    except (TypeError, ValueError):
        return VIDEO_MAX_PARALLEL_CHUNKS


def _drive_query_literal(value: str) -> str:
    return str(value).replace("\\", "\\\\").replace("'", "\\'")


def _find_drive_item(
    service,
    parent_id: str,
    name: str,
    mime_type: str | None = None,
) -> dict | None:
    terms = [
        f"'{_drive_query_literal(parent_id)}' in parents",
        f"name = '{_drive_query_literal(name)}'",
        "trashed = false",
    ]
    if mime_type:
        terms.append(f"mimeType = '{_drive_query_literal(mime_type)}'")
    response = (
        service.files()
        .list(
            q=" and ".join(terms),
            spaces="drive",
            fields="files(id,name,mimeType,modifiedTime)",
            pageSize=10,
            supportsAllDrives=True,
            includeItemsFromAllDrives=True,
        )
        .execute()
    )
    files = response.get("files") or []
    return files[0] if files else None


def _ensure_drive_checkpoint_folder(service, checkpoint_id: str) -> str:
    root_id = _drive_checkpoint_folder_id()
    if not root_id:
        return ""

    folder_name = f"video_checkpoint_{checkpoint_id[:20]}"
    folder_mime = "application/vnd.google-apps.folder"
    existing = _find_drive_item(service, root_id, folder_name, folder_mime)
    if existing:
        return str(existing["id"])

    created = (
        service.files()
        .create(
            body={
                "name": folder_name,
                "mimeType": folder_mime,
                "parents": [root_id],
                "appProperties": {
                    "checkpoint_id": checkpoint_id,
                    "checkpoint_version": VIDEO_CHECKPOINT_VERSION,
                },
            },
            fields="id",
            supportsAllDrives=True,
        )
        .execute()
    )
    return str(created.get("id") or "")


def _save_checkpoint_to_drive(
    checkpoint_id: str,
    filename: str,
    payload: dict,
) -> bool:
    if not _drive_checkpoint_folder_id():
        return False

    service = None
    try:
        service = build_drive_service(writable=True)
        folder_id = _ensure_drive_checkpoint_folder(service, checkpoint_id)
        if not folder_id:
            return False

        encoded = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        media = MediaIoBaseUpload(
            BytesIO(encoded),
            mimetype="application/json",
            resumable=False,
        )
        existing = _find_drive_item(
            service,
            folder_id,
            filename,
            "application/json",
        )
        if existing:
            service.files().update(
                fileId=existing["id"],
                media_body=media,
                fields="id",
                supportsAllDrives=True,
            ).execute()
        else:
            service.files().create(
                body={
                    "name": filename,
                    "parents": [folder_id],
                    "mimeType": "application/json",
                },
                media_body=media,
                fields="id",
                supportsAllDrives=True,
            ).execute()
        return True
    except Exception:
        # El checkpoint local sigue protegiendo frente a reruns y reinicios suaves.
        return False
    finally:
        if service is not None:
            _close_drive_service(service)


def _download_drive_file_bytes(service, file_id: str) -> bytes:
    request = service.files().get_media(
        fileId=file_id,
        supportsAllDrives=True,
    )
    output = BytesIO()
    downloader = MediaIoBaseDownload(output, request, chunksize=1024 * 1024)
    done = False
    while not done:
        _, done = downloader.next_chunk(num_retries=2)
    return output.getvalue()


def _load_checkpoints_from_drive(
    checkpoint_id: str,
) -> list[tuple[float, float, VideoTranscript]]:
    if not _drive_checkpoint_folder_id():
        return []

    service = None
    try:
        service = build_drive_service(writable=True)
        folder_id = _ensure_drive_checkpoint_folder(service, checkpoint_id)
        if not folder_id:
            return []

        result: list[tuple[float, float, VideoTranscript]] = []
        page_token = None
        while True:
            response = (
                service.files()
                .list(
                    q=(
                        f"'{_drive_query_literal(folder_id)}' in parents and "
                        "mimeType = 'application/json' and trashed = false"
                    ),
                    spaces="drive",
                    fields="nextPageToken,files(id,name)",
                    pageSize=100,
                    pageToken=page_token,
                    supportsAllDrives=True,
                    includeItemsFromAllDrives=True,
                )
                .execute()
            )
            for item in response.get("files") or []:
                if not str(item.get("name") or "").startswith("chunk_"):
                    continue
                try:
                    raw = _download_drive_file_bytes(service, str(item["id"]))
                    payload = json.loads(raw.decode("utf-8"))
                    parsed = _checkpoint_from_payload(payload)
                    if parsed is not None:
                        result.append(parsed)
                except Exception:
                    continue
            page_token = response.get("nextPageToken")
            if not page_token:
                break
        return result
    except Exception:
        return []
    finally:
        if service is not None:
            _close_drive_service(service)


def _delete_drive_checkpoint_folder(checkpoint_id: str) -> bool:
    root_id = _drive_checkpoint_folder_id()
    if not root_id or not checkpoint_id:
        return False

    service = None
    try:
        service = build_drive_service(writable=True)
        folder_name = f"video_checkpoint_{checkpoint_id[:20]}"
        folder = _find_drive_item(
            service,
            root_id,
            folder_name,
            "application/vnd.google-apps.folder",
        )
        if not folder:
            return False
        service.files().delete(
            fileId=folder["id"],
            supportsAllDrives=True,
        ).execute()
        return True
    except Exception:
        return False
    finally:
        if service is not None:
            _close_drive_service(service)


def _clear_video_checkpoints(checkpoint_id: str) -> None:
    """Elimina el progreso cuando el usuario solicita empezar desde cero."""

    if not checkpoint_id:
        return
    local_dir = _checkpoint_root_dir(checkpoint_id)
    try:
        if local_dir.exists():
            shutil.rmtree(local_dir)
    except Exception:
        pass
    _delete_drive_checkpoint_folder(checkpoint_id)


def _load_video_chunk_checkpoints(
    checkpoint_id: str,
    progress_callback=None,
) -> list[tuple[float, float, VideoTranscript]]:
    """Recupera primero el progreso local y después completa desde Drive."""

    if not VIDEO_CHECKPOINT_ENABLED:
        return []

    by_interval: dict[tuple[int, int], tuple[float, float, VideoTranscript]] = {}
    local_dir = _checkpoint_root_dir(checkpoint_id)
    if local_dir.exists():
        for path in sorted(local_dir.glob("chunk_*.json")):
            try:
                payload = json.loads(path.read_text(encoding="utf-8"))
                parsed = _checkpoint_from_payload(payload)
                if parsed is None:
                    continue
                key = (int(round(parsed[0])), int(round(parsed[1])))
                by_interval[key] = parsed
            except Exception:
                continue

    remote = _load_checkpoints_from_drive(checkpoint_id)
    for parsed in remote:
        key = (int(round(parsed[0])), int(round(parsed[1])))
        if key not in by_interval:
            by_interval[key] = parsed
            payload = _checkpoint_payload(
                checkpoint_id=checkpoint_id,
                model="recuperado",
                start_seconds=parsed[0],
                end_seconds=parsed[1],
                transcript=parsed[2],
                audited=False,
                audit_reasons=["Recuperado desde Google Drive"],
            )
            # Conserva la versión exacta esperada al materializar localmente.
            payload["version"] = VIDEO_CHECKPOINT_VERSION
            _write_json_atomic(
                local_dir / _checkpoint_filename(parsed[0], parsed[1]),
                payload,
            )

    chunks = sorted(by_interval.values(), key=lambda item: (item[0], item[1]))
    if chunks and progress_callback is not None:
        progress_callback(
            f"Progreso recuperado: {len(chunks)} tramo(s) ya procesado(s)."
        )
    return chunks


def _save_video_chunk_checkpoint(
    checkpoint_id: str,
    model: str,
    start_seconds: float,
    end_seconds: float,
    transcript: VideoTranscript,
    audited: bool,
    audit_reasons: list[str],
) -> None:
    if not VIDEO_CHECKPOINT_ENABLED:
        return

    payload = _checkpoint_payload(
        checkpoint_id=checkpoint_id,
        model=model,
        start_seconds=start_seconds,
        end_seconds=end_seconds,
        transcript=transcript,
        audited=audited,
        audit_reasons=audit_reasons,
    )
    filename = _checkpoint_filename(start_seconds, end_seconds)
    local_path = _checkpoint_root_dir(checkpoint_id) / filename
    _write_json_atomic(local_path, payload)
    _save_checkpoint_to_drive(checkpoint_id, filename, payload)


def _interval_is_covered(
    start_seconds: float,
    end_seconds: float,
    chunks: list[tuple[float, float, VideoTranscript]],
) -> bool:
    """Valida que uno o varios checkpoints cubran completamente un intervalo."""

    relevant = sorted(
        (
            (max(start_seconds, start), min(end_seconds, end))
            for start, end, _ in chunks
            if end > start_seconds and start < end_seconds
        ),
        key=lambda item: item[0],
    )
    cursor = start_seconds
    for start, end in relevant:
        if start > cursor + 0.5:
            return False
        cursor = max(cursor, end)
        if cursor >= end_seconds - 0.5:
            return True
    return cursor >= end_seconds - 0.5


def _action_is_too_general(action: VideoAction) -> bool:
    value = _normalized_action_key(action.action)
    generic_phrases = (
        "realizar el proceso",
        "continuar con el proceso",
        "hacer el procedimiento",
        "gestionar la informacion",
        "ejecutar la actividad",
        "realizar las acciones",
    )
    return any(phrase in value for phrase in generic_phrases)


def _chunk_needs_action_audit(
    transcript: VideoTranscript,
) -> tuple[bool, list[str]]:
    """Decide si el tramo requiere una segunda lectura audiovisual."""

    if not VIDEO_ACTION_AUDIT_ENABLED:
        return False, []
    if VIDEO_ACTION_AUDIT_MODE != "selective":
        return True, ["Auditoría completa configurada"]

    reasons: list[str] = []
    actions = list(transcript.actions)
    full_text = transcript.full_transcript.strip()

    if not actions:
        reasons.append("No se detectaron acciones en la primera lectura")
    if transcript.uncertainties:
        reasons.append("El tramo contiene incertidumbres")
    if any(not _clean_action_value(action.timestamp_start) for action in actions):
        reasons.append("Hay acciones sin marca de tiempo")
    if any(_action_is_too_general(action) for action in actions):
        reasons.append("Hay acciones redactadas de forma demasiado general")

    verb_mentions = len(ACTION_VERB_RE.findall(full_text))
    if full_text and len(full_text) >= 1800 and len(actions) < 3:
        reasons.append("La transcripción es extensa para la cantidad de acciones")
    if verb_mentions >= 8 and verb_mentions > max(5, len(actions) * 2 + 2):
        reasons.append("La transcripción menciona más operaciones que el inventario")

    return bool(reasons), reasons


def _generate_video_range_with_retries(
    client,
    model: str,
    uploaded_file,
    prompt: str,
    start_seconds: float,
    end_seconds: float,
) -> VideoTranscript:
    last_error: Exception | None = None
    for attempt in range(3):
        try:
            result = _call_video_generate_content_fallback(
                client=client,
                model=model,
                uploaded_file=uploaded_file,
                prompt=prompt,
                schema_model=VideoTranscript,
                start_seconds=start_seconds,
                end_seconds=end_seconds,
            )
            assert isinstance(result, VideoTranscript)
            return result
        except Exception as error:
            last_error = error
            if not _is_retryable_gemini_error(error) or attempt == 2:
                raise
            time.sleep(min(8.0, 1.5 * (2 ** attempt)) + random.uniform(0.1, 0.5))
    assert last_error is not None
    raise last_error


ACTION_VERB_RE = re.compile(
    r"\b(abrir|acceder|ingresar|iniciar|seleccionar|hacer clic|pulsar|presionar|"
    r"ubicar|buscar|consultar|diligenciar|registrar|escribir|digitar|modificar|"
    r"adjuntar|cargar|descargar|guardar|enviar|aprobar|validar|verificar|confirmar|"
    r"crear|eliminar|cerrar|filtrar|marcar|desmarcar|elegir|visualizar|revisar)\b",
    flags=re.I,
)
TIMESTAMP_RE = re.compile(r"\[(\d{1,2}:\d{2}(?::\d{2})?)\]")


def _derive_actions_from_text(transcript: str, visual_evidence: list[str]) -> list[VideoAction]:
    """Respaldo conservador cuando el modelo omite el arreglo actions."""

    candidates = list(str(transcript or "").splitlines()) + list(visual_evidence)
    actions: list[VideoAction] = []
    last_timestamp = ""
    for raw_line in candidates:
        line = re.sub(r"\s+", " ", str(raw_line)).strip(" -•\t")
        if not line:
            continue
        timestamp_match = TIMESTAMP_RE.search(line)
        if timestamp_match:
            last_timestamp = timestamp_match.group(1)
        cleaned = TIMESTAMP_RE.sub("", line).strip(" :-")
        if len(cleaned) < MIN_ACTION_TEXT_CHARS or not ACTION_VERB_RE.search(cleaned):
            continue
        actions.append(
            VideoAction(
                timestamp_start=last_timestamp,
                action=cleaned,
                evidence=line,
            )
        )
    return actions


def _merge_video_transcript_chunks(
    chunks: list[tuple[float, float, VideoTranscript]],
    duration_seconds: float,
) -> VideoTranscript:
    """Une transcripciones y crea un inventario global de acciones sin duplicados."""

    languages: list[str] = []
    speakers: list[str] = []
    transcript_blocks: list[str] = []
    all_actions: list[VideoAction] = []
    visual_evidence: list[str] = []
    key_facts: list[str] = []
    uncertainties: list[str] = []

    for start_seconds, end_seconds, chunk in chunks:
        interval = f"{_video_clock(start_seconds)}–{_video_clock(end_seconds)}"
        if chunk.detected_language.strip():
            languages.append(chunk.detected_language.strip())
        for speaker in chunk.speakers:
            value = str(speaker).strip()
            if value and value not in speakers:
                speakers.append(value)
        if chunk.full_transcript.strip():
            transcript_blocks.append(f"[TRAMO {interval}]\n{chunk.full_transcript.strip()}")

        chunk_actions = list(chunk.actions)
        if not chunk_actions:
            chunk_actions = _derive_actions_from_text(
                chunk.full_transcript,
                chunk.visual_evidence,
            )
        for action in chunk_actions:
            if not action.timestamp_start.strip():
                action = action.model_copy(update={"timestamp_start": _video_clock(start_seconds)})
            all_actions.append(action)

        for item in chunk.visual_evidence:
            value = str(item).strip()
            if value:
                visual_evidence.append(f"[{interval}] {value}")
        for item in chunk.key_facts:
            value = str(item).strip()
            if value and value not in key_facts:
                key_facts.append(value)
        for item in chunk.uncertainties:
            value = str(item).strip()
            if value:
                uncertainties.append(f"[{interval}] {value}")

    normalized_actions: list[VideoAction] = []
    seen: set[str] = set()
    for action in all_actions:
        action_text = _clean_action_value(action.action)
        if len(action_text) < MIN_ACTION_TEXT_CHARS:
            continue
        key = re.sub(
            r"[^a-záéíóúñ0-9]+",
            " ",
            f"{action.timestamp_start} {action_text} {action.interface_element}".lower(),
        ).strip()
        if not key or key in seen:
            continue
        seen.add(key)
        normalized_actions.append(
            action.model_copy(
                update={
                    "action_id": f"ACC-{len(normalized_actions)+1:04d}",
                    "action": action_text,
                }
            )
        )

    return VideoTranscript(
        detected_language=languages[0] if languages else "",
        duration_estimate=_video_clock(duration_seconds),
        speakers=speakers,
        full_transcript="\n\n".join(transcript_blocks),
        actions=normalized_actions,
        visual_evidence=visual_evidence,
        key_facts=key_facts,
        uncertainties=uncertainties,
    )

def _build_video_file_part(
    uploaded_file,
    start_seconds: float | None = None,
    end_seconds: float | None = None,
) -> types.Part:
    """Construye una referencia File API explícita y opcionalmente recortada."""

    mime_type = str(getattr(uploaded_file, "mime_type", "") or "video/mp4")
    file_uri = str(getattr(uploaded_file, "uri", "") or "")
    if not file_uri:
        raise AppError("Gemini no devolvió la URI del video procesado.")

    kwargs = {
        "file_data": types.FileData(
            file_uri=file_uri,
            mime_type=mime_type,
        )
    }

    if start_seconds is not None and end_seconds is not None:
        kwargs["video_metadata"] = types.VideoMetadata(
            start_offset=f"{max(0, int(start_seconds))}s",
            end_offset=f"{max(1, int(end_seconds))}s",
        )

    return types.Part(**kwargs)


def _response_finish_reason(response) -> str:
    """Devuelve la razón de finalización de la primera alternativa."""

    try:
        candidates = getattr(response, "candidates", None) or []
        if not candidates:
            return ""
        reason = getattr(candidates[0], "finish_reason", None)
        name = getattr(reason, "name", None)
        return str(name or reason or "").upper()
    except Exception:
        return ""


def _tagged_section(text: str, name: str, next_names: list[str]) -> str:
    """Extrae una sección de una respuesta etiquetada sin exigir JSON."""

    end_pattern = "|".join(re.escape(item) for item in next_names)
    pattern = rf"(?is)\[{re.escape(name)}\]\s*(.*?)(?=\n\[(?:{end_pattern})\]\s*|\Z)"
    match = re.search(pattern, text)
    return match.group(1).strip() if match else ""


def _tagged_list(text: str, name: str, next_names: list[str]) -> list[str]:
    block = _tagged_section(text, name, next_names)
    if not block:
        return []
    values: list[str] = []
    for line in block.splitlines():
        cleaned = re.sub(r"^\s*(?:[-•*]|\d+[.)])\s*", "", line).strip()
        if cleaned:
            values.append(cleaned)
    return values


def _parse_tagged_actions(block: str) -> list[VideoAction]:
    actions: list[VideoAction] = []
    for line in str(block or "").splitlines():
        cleaned = re.sub(r"^\s*(?:[-•*]|\d+[.)])\s*", "", line).strip()
        if not cleaned:
            continue
        values = [value.strip() for value in cleaned.split("|")]
        values.extend([""] * (9 - len(values)))
        action = _video_action_from_mapping(
            {
                "timestamp_start": values[0],
                "actor": values[1],
                "system": values[2],
                "location_path": values[3],
                "action": values[4] or cleaned,
                "interface_element": values[5],
                "data_handled": values[6],
                "validation": values[7],
                "result": values[8],
            }
        )
        if action is not None:
            actions.append(action)
    return actions


def _parse_tagged_video_response(
    raw_text: str,
    duration_estimate: str = "",
) -> VideoTranscript:
    """Convierte el formato etiquetado de respaldo en VideoTranscript."""

    cleaned = _clean_json_text(str(raw_text or "")).strip()
    names = [
        "IDIOMA",
        "HABLANTES",
        "TRANSCRIPCION",
        "TRANSCRIPCIÓN",
        "ACCIONES",
        "EVIDENCIA_VISUAL",
        "HECHOS_CLAVE",
        "INCERTIDUMBRES",
    ]

    language = _tagged_section(cleaned, "IDIOMA", names[1:])
    speakers = _tagged_list(cleaned, "HABLANTES", names[2:])
    transcript = _tagged_section(
        cleaned,
        "TRANSCRIPCION",
        ["ACCIONES", "EVIDENCIA_VISUAL", "HECHOS_CLAVE", "INCERTIDUMBRES"],
    )
    if not transcript:
        transcript = _tagged_section(
            cleaned,
            "TRANSCRIPCIÓN",
            ["ACCIONES", "EVIDENCIA_VISUAL", "HECHOS_CLAVE", "INCERTIDUMBRES"],
        )
    action_block = _tagged_section(
        cleaned,
        "ACCIONES",
        ["EVIDENCIA_VISUAL", "HECHOS_CLAVE", "INCERTIDUMBRES"],
    )
    actions = _parse_tagged_actions(action_block)
    visual = _tagged_list(cleaned, "EVIDENCIA_VISUAL", ["HECHOS_CLAVE", "INCERTIDUMBRES"])
    facts = _tagged_list(cleaned, "HECHOS_CLAVE", ["INCERTIDUMBRES"])
    uncertainties = _tagged_list(cleaned, "INCERTIDUMBRES", [])

    if not transcript:
        transcript = cleaned
    if not actions:
        actions = _derive_actions_from_text(transcript, visual)

    return VideoTranscript(
        detected_language=language,
        duration_estimate=duration_estimate,
        speakers=speakers,
        full_transcript=transcript,
        actions=actions,
        visual_evidence=visual,
        key_facts=facts,
        uncertainties=uncertainties,
    )

def _plain_video_fallback_prompt(
    prompt: str,
    start_seconds: float | None,
    end_seconds: float | None,
) -> str:
    interval = ""
    if start_seconds is not None and end_seconds is not None:
        interval = (
            "\nAnaliza exclusivamente el intervalo "
            f"{_video_clock(start_seconds)} a {_video_clock(end_seconds)}. "
            "Usa horas absolutas del video original."
        )

    return (
        prompt
        + interval
        + "\n\nFORMATO DE RESPALDO OBLIGATORIO\n"
          "No uses JSON ni Markdown. Devuelve exactamente estas etiquetas:\n"
          "[IDIOMA]\nIdioma principal\n"
          "[HABLANTES]\n- Hablante 1\n"
          "[TRANSCRIPCION]\nTranscripción cronológica completa del tramo\n"
          "[ACCIONES]\n"
          "- HH:MM:SS | responsable | sistema | ruta | UNA acción atómica | "
          "elemento | dato | validación | resultado\n"
          "Incluye una línea independiente por cada clic, selección, navegación, "
          "registro, carga, validación o resultado. No resumas varias acciones.\n"
          "[EVIDENCIA_VISUAL]\n- Evidencia con hora\n"
          "[HECHOS_CLAVE]\n- Hecho verificable\n"
          "[INCERTIDUMBRES]\n- Duda o fragmento inaudible\n"
          "No agregues texto fuera de esas etiquetas."
    )

def _call_video_generate_content_fallback(
    client,
    model: str,
    uploaded_file,
    prompt: str,
    schema_model: type[BaseModel],
    start_seconds: float | None = None,
    end_seconds: float | None = None,
) -> BaseModel:
    """Analiza un video con esquema nativo y respaldo etiquetado.

    El error anterior se producía porque el JSON solo se pedía por texto: el
    modelo podía inventar campos como ``description`` o cortar el objeto. Esta
    versión aplica ``response_json_schema`` y reduce cada tramo a diez minutos.
    """

    schema = _inline_and_clean_json_schema(schema_model)
    interval_instruction = ""

    if start_seconds is not None and end_seconds is not None:
        interval_instruction = (
            "\n\nINTERVALO OBLIGATORIO\n"
            f"Analiza exclusivamente el tramo {_video_clock(start_seconds)} a "
            f"{_video_clock(end_seconds)} del video completo. Usa marcas de "
            "tiempo absolutas respecto al inicio del video original. "
            "No resumas ni omitas intervenciones relevantes."
        )

    file_part = _build_video_file_part(
        uploaded_file=uploaded_file,
        start_seconds=start_seconds,
        end_seconds=end_seconds,
    )
    prompt_part = types.Part.from_text(text=prompt + interval_instruction)

    structured_error: Exception | None = None
    try:
        response = client.models.generate_content(
            model=model,
            contents=types.Content(
                role="user",
                parts=[file_part, prompt_part],
            ),
            config=types.GenerateContentConfig(
                media_resolution=types.MediaResolution.MEDIA_RESOLUTION_LOW,
                response_mime_type="application/json",
                response_json_schema=schema,
                temperature=0.0,
                max_output_tokens=VIDEO_MAX_OUTPUT_TOKENS,
            ),
        )

        finish_reason = _response_finish_reason(response)
        if "MAX_TOKENS" in finish_reason:
            raise AppError(
                "La respuesta estructurada quedó incompleta por límite de salida."
            )

        parsed = _validate_structured_output(schema_model, response)
        if isinstance(parsed, VideoTranscript) and not parsed.full_transcript.strip():
            raise AppError("Gemini devolvió una transcripción vacía.")
        return parsed

    except Exception as error:
        structured_error = error

    # Respaldo sin JSON: evita que un objeto incompleto bloquee todo el video.
    plain_prompt = _plain_video_fallback_prompt(
        prompt=prompt,
        start_seconds=start_seconds,
        end_seconds=end_seconds,
    )
    plain_response = client.models.generate_content(
        model=model,
        contents=types.Content(
            role="user",
            parts=[
                file_part,
                types.Part.from_text(text=plain_prompt),
            ],
        ),
        config=types.GenerateContentConfig(
            media_resolution=types.MediaResolution.MEDIA_RESOLUTION_LOW,
            temperature=0.0,
            max_output_tokens=VIDEO_MAX_OUTPUT_TOKENS,
        ),
    )

    plain_text = str(getattr(plain_response, "text", None) or "").strip()
    if not plain_text:
        raise AppError(
            "Gemini no devolvió texto para el tramo. Error estructurado previo: "
            f"{str(structured_error)[:350]}"
        )

    finish_reason = _response_finish_reason(plain_response)
    if "MAX_TOKENS" in finish_reason:
        raise AppError(
            "La transcripción del tramo quedó incompleta por límite de salida."
        )

    duration = ""
    if start_seconds is not None and end_seconds is not None:
        duration = _video_clock(max(0, end_seconds - start_seconds))

    parsed_fallback = _parse_tagged_video_response(
        plain_text,
        duration_estimate=duration,
    )
    if not parsed_fallback.full_transcript.strip():
        raise AppError(
            "Gemini devolvió una respuesta sin transcripción utilizable."
        )
    return parsed_fallback


def _merge_action_candidates(
    primary: list[VideoAction],
    audited: list[VideoAction],
) -> list[VideoAction]:
    """Une inventarios conservando la versión más exhaustiva sin duplicados."""

    combined = list(audited) + list(primary)
    result: list[VideoAction] = []
    seen: set[str] = set()
    for action in combined:
        value = _clean_action_value(action.action)
        if len(value) < MIN_ACTION_TEXT_CHARS:
            continue
        key = re.sub(
            r"[^a-záéíóúñ0-9]+",
            " ",
            f"{action.timestamp_start} {value} {action.interface_element}".lower(),
        ).strip()
        if not key or key in seen:
            continue
        seen.add(key)
        result.append(action.model_copy(update={"action": value}))
    return result


def _video_action_audit_prompt(start_seconds: float, end_seconds: float) -> str:
    return f"""
Actúa exclusivamente como auditor de acciones operativas de un video.

INTERVALO
{_video_clock(start_seconds)} a {_video_clock(end_seconds)} del video original.

TAREA
Observa nuevamente TODO el intervalo y registra cada acción operativa como un
elemento independiente en actions. No redactes un resumen ni una explicación
general del tramo.

REGLAS ESTRICTAS
1. Separa cada clic, apertura de menú, selección de pestaña, búsqueda, filtro,
   consulta, diligenciamiento, modificación, carga, descarga, guardado, envío,
   validación, aprobación, confirmación, mensaje o revisión de resultado.
2. Dos verbos operativos diferentes no deben quedar en la misma acción.
3. No omitas acciones por ser rápidas, repetitivas, obvias o silenciosas.
4. Usa marcas de tiempo absolutas HH:MM:SS.
5. Para cada acción completa, cuando sea visible o audible: actor, sistema,
   ruta, acción, elemento, dato, validación, resultado y evidencia.
   Si un campo no tiene información verificable, devuelve una cadena vacía.
   Nunca escribas null, None, N/A, no aplica, por confirmar ni equivalentes.
6. No inventes acciones. Cuando el intervalo no contenga ejecución ni
   instrucciones operativas, devuelve actions vacío.
7. Antes de responder, recorre nuevamente el intervalo y verifica que no haya
   ninguna acción operativa sin registrar.

Devuelve únicamente el objeto estructurado solicitado.
""".strip()


def _extract_actions_for_video_range(
    client,
    model: str,
    uploaded_file,
    start_seconds: float,
    end_seconds: float,
) -> list[VideoAction]:
    """Segunda lectura enfocada solo en acciones para evitar generalizaciones."""

    file_part = _build_video_file_part(
        uploaded_file=uploaded_file,
        start_seconds=start_seconds,
        end_seconds=end_seconds,
    )
    prompt = _video_action_audit_prompt(start_seconds, end_seconds)
    schema = _inline_and_clean_json_schema(VideoActionBatch)

    try:
        response = client.models.generate_content(
            model=model,
            contents=types.Content(
                role="user",
                parts=[file_part, types.Part.from_text(text=prompt)],
            ),
            config=types.GenerateContentConfig(
                media_resolution=types.MediaResolution.MEDIA_RESOLUTION_LOW,
                response_mime_type="application/json",
                response_json_schema=schema,
                temperature=0.0,
                max_output_tokens=VIDEO_MAX_OUTPUT_TOKENS,
            ),
        )
        parsed = _validate_structured_output(VideoActionBatch, response)
        assert isinstance(parsed, VideoActionBatch)
        return parsed.actions
    except Exception:
        fallback_prompt = (
            prompt
            + "\n\nFORMATO DE RESPALDO\n"
              "No uses JSON. Devuelve únicamente [ACCIONES] y una línea por "
              "acción con este formato:\n"
              "- HH:MM:SS | responsable | sistema | ruta | UNA acción | "
              "elemento | dato | validación | resultado"
        )
        response = client.models.generate_content(
            model=model,
            contents=types.Content(
                role="user",
                parts=[file_part, types.Part.from_text(text=fallback_prompt)],
            ),
            config=types.GenerateContentConfig(
                media_resolution=types.MediaResolution.MEDIA_RESOLUTION_LOW,
                temperature=0.0,
                max_output_tokens=VIDEO_MAX_OUTPUT_TOKENS,
            ),
        )
        raw = str(getattr(response, "text", None) or "")
        block = _tagged_section(raw, "ACCIONES", []) or raw
        return _parse_tagged_actions(block)


def _transcribe_video_by_intervals(
    client,
    model: str,
    uploaded_file,
    prompt: str,
    duration_seconds: float,
    progress_callback=None,
    api_key: str = "",
    checkpoint_key: str = "",
) -> VideoTranscript:
    """Procesa el video en tramos reanudables, con auditoría selectiva.

    - Comienza directamente en intervalos de cinco minutos.
    - Reutiliza checkpoints locales y, opcionalmente, de Google Drive.
    - Procesa hasta dos intervalos simultáneamente.
    - Solo repite el análisis audiovisual cuando existen señales de omisión.
    """

    checkpoint_id = _video_checkpoint_id(
        checkpoint_key=checkpoint_key or str(getattr(uploaded_file, "name", "video")),
        model=model,
        duration_seconds=duration_seconds,
    )
    chunks = _load_video_chunk_checkpoints(
        checkpoint_id,
        progress_callback=progress_callback,
    )

    planned: list[tuple[float, float]] = []
    start_seconds = 0.0
    while start_seconds < duration_seconds:
        end_seconds = min(duration_seconds, start_seconds + VIDEO_CHUNK_SECONDS)
        if not _interval_is_covered(start_seconds, end_seconds, chunks):
            planned.append((start_seconds, end_seconds))
        start_seconds = end_seconds

    total_intervals = max(1, int((duration_seconds + VIDEO_CHUNK_SECONDS - 1) // VIDEO_CHUNK_SECONDS))
    recovered_intervals = total_intervals - len(planned)
    if progress_callback is not None:
        if recovered_intervals:
            progress_callback(
                f"Reanudando: {recovered_intervals} de {total_intervals} "
                "intervalos ya estaban completos."
            )
        else:
            progress_callback(
                f"Iniciando {total_intervals} intervalos de cinco minutos."
            )

    def process_range(
        worker_client,
        start_at: float,
        end_at: float,
        depth: int = 0,
    ) -> list[tuple[float, float, VideoTranscript, bool, list[str]]]:
        try:
            result = _generate_video_range_with_retries(
                client=worker_client,
                model=model,
                uploaded_file=uploaded_file,
                prompt=prompt,
                start_seconds=start_at,
                end_seconds=end_at,
            )

            needs_audit, audit_reasons = _chunk_needs_action_audit(result)
            primary_actions = list(result.actions)
            audited = False

            if needs_audit:
                audited_actions = _extract_actions_for_video_range(
                    client=worker_client,
                    model=model,
                    uploaded_file=uploaded_file,
                    start_seconds=start_at,
                    end_seconds=end_at,
                )
                primary_actions = _merge_action_candidates(
                    primary_actions,
                    audited_actions,
                )
                audited = True

            if not primary_actions:
                primary_actions = _derive_actions_from_text(
                    result.full_transcript,
                    result.visual_evidence,
                )
            result = result.model_copy(update={"actions": primary_actions})
            return [(start_at, end_at, result, audited, audit_reasons)]

        except Exception:
            duration = end_at - start_at
            if duration <= VIDEO_MIN_CHUNK_SECONDS or depth >= 4:
                raise
            midpoint = start_at + duration / 2
            left = process_range(worker_client, start_at, midpoint, depth + 1)
            right = process_range(worker_client, midpoint, end_at, depth + 1)
            return left + right

    def worker(start_at: float, end_at: float):
        # Cada hilo usa su propio cliente. Se comparte únicamente la URI de un
        # archivo ya procesado por Gemini Files API.
        worker_client = genai.Client(api_key=api_key) if api_key else client
        owns_client = worker_client is not client
        try:
            return process_range(worker_client, start_at, end_at)
        finally:
            if owns_client:
                try:
                    worker_client.close()
                except Exception:
                    pass

    completed_new = 0
    max_workers = min(_video_parallel_workers(), max(1, len(planned)))

    processing_errors: list[str] = []
    if planned:
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            future_map = {
                executor.submit(worker, start_at, end_at): (start_at, end_at)
                for start_at, end_at in planned
            }
            for future in as_completed(future_map):
                original_start, original_end = future_map[future]
                if progress_callback is not None:
                    progress_callback(
                        "Completando tramo "
                        f"{_video_clock(original_start)}–{_video_clock(original_end)}"
                    )
                try:
                    produced = future.result()
                except Exception as error:
                    processing_errors.append(
                        f"{_video_clock(original_start)}–{_video_clock(original_end)}: "
                        f"{str(error)[:300]}"
                    )
                    continue

                for start_at, end_at, result, audited, reasons in produced:
                    chunks.append((start_at, end_at, result))
                    _save_video_chunk_checkpoint(
                        checkpoint_id=checkpoint_id,
                        model=model,
                        start_seconds=start_at,
                        end_seconds=end_at,
                        transcript=result,
                        audited=audited,
                        audit_reasons=reasons,
                    )
                completed_new += 1
                if progress_callback is not None:
                    done = recovered_intervals + completed_new
                    progress_callback(
                        f"Progreso: {done} de {total_intervals} intervalos. "
                        "Cada tramo terminado quedó guardado."
                    )

    if processing_errors:
        raise AppError(
            "Algunos intervalos no pudieron completarse, pero todos los tramos "
            "exitosos quedaron guardados. Vuelve a ejecutar la transcripción para "
            "reanudar únicamente los intervalos pendientes. Detalle: "
            + " | ".join(processing_errors[:4])
        )

    chunks.sort(key=lambda item: (item[0], item[1]))
    if not _interval_is_covered(0.0, duration_seconds, chunks):
        raise AppError(
            "La transcripción no cubre todavía la duración completa del video. "
            "Vuelve a ejecutar el proceso para reanudar desde el último checkpoint."
        )

    merged = _merge_video_transcript_chunks(chunks, duration_seconds)
    # Se expone el ID para que la interfaz pueda informar y, cuando el usuario
    # solicite una transcripción completamente nueva, eliminar el progreso local.
    try:
        st.session_state["video_checkpoint_id"] = checkpoint_id
    except Exception:
        pass
    return merged

def _transcribe_local_video_path_with_gemini(
    api_key: str,
    model: str,
    local_path: Path,
    display_name: str,
    mime_type: str,
    progress_callback=None,
    checkpoint_key: str = "",
) -> VideoTranscript:
    """Sube el video y lo procesa en intervalos reanudables de cinco minutos."""

    if not api_key:
        raise AppError("No se encontró GEMINI_API_KEY para transcribir el video.")
    if not local_path.exists() or local_path.stat().st_size <= 0:
        raise AppError("El video temporal no existe o está vacío.")

    max_file_api_bytes = 1_950_000_000
    if local_path.stat().st_size > max_file_api_bytes:
        raise AppError(
            "El video supera el límite seguro de 1,95 GB para Gemini Files API. "
            "Comprímelo o divídelo antes de procesarlo."
        )

    client = genai.Client(api_key=api_key)
    uploaded_file = None

    def report(message: str) -> None:
        if progress_callback is not None:
            progress_callback(message)

    try:
        local_duration = _probe_local_video_duration_seconds(local_path)
        if local_duration is not None:
            report(
                "Duración verificada localmente: "
                f"{_video_clock(local_duration)}"
            )

        report("Subiendo el video temporal a Gemini Files API")
        uploaded_file = client.files.upload(file=str(local_path))

        if not getattr(uploaded_file, "name", None):
            raise AppError("Gemini no devolvió un identificador para el video.")

        report("Esperando que Gemini procese el audio y los fotogramas")
        deadline = time.monotonic() + VIDEO_PROCESSING_TIMEOUT_SECONDS

        while True:
            file_info = client.files.get(name=uploaded_file.name)
            state_name = _file_state_name(file_info)

            if "ACTIVE" in state_name:
                uploaded_file = file_info
                break

            if "FAILED" in state_name:
                status = getattr(file_info, "error", None) or getattr(
                    file_info, "status", None
                )
                raise AppError(
                    "Gemini no pudo procesar el video. "
                    f"Estado recibido: {status or state_name}."
                )

            if time.monotonic() >= deadline:
                raise AppError(
                    "El video tardó demasiado en procesarse. "
                    "Prueba con un archivo más corto o de menor tamaño."
                )

            time.sleep(VIDEO_POLL_INTERVAL_SECONDS)

        processed_mime = str(
            getattr(uploaded_file, "mime_type", None) or mime_type
        )
        if not processed_mime.lower().startswith("video/"):
            raise AppError(
                "Gemini recibió el archivo, pero lo identificó con un tipo no "
                f"compatible: {processed_mime}."
            )

        duration_seconds = (
            local_duration
            or _gemini_video_duration_seconds(uploaded_file)
        )

        if duration_seconds is not None:
            report(
                "Duración que se usará para el análisis: "
                f"{_video_clock(duration_seconds)}"
            )

        prompt = video_transcription_prompt(display_name)
        errors: list[str] = []

        for candidate_model in _video_model_candidates(model):
            report(f"Analizando el video con {candidate_model}")

            # Los videos largos se procesan directamente por intervalos para no
            # superar la ventana de contexto de una sola solicitud.
            if (
                duration_seconds is not None
                and duration_seconds > VIDEO_LONG_THRESHOLD_SECONDS
            ):
                try:
                    result = _transcribe_video_by_intervals(
                        client=client,
                        model=candidate_model,
                        uploaded_file=uploaded_file,
                        prompt=prompt,
                        duration_seconds=duration_seconds,
                        progress_callback=progress_callback,
                        api_key=api_key,
                        checkpoint_key=checkpoint_key or display_name,
                    )
                    st.session_state["active_gemini_model"] = candidate_model
                    return result
                except Exception as chunk_error:
                    errors.append(
                        f"{candidate_model} por tramos: "
                        f"{str(chunk_error)[:300]}"
                    )
                    continue

            interaction_error: Exception | None = None

            # Para videos cortos se intenta primero la ruta oficial de
            # Interactions API.
            if hasattr(client, "interactions"):
                try:
                    interaction_result = _call_video_interaction(
                        client=client,
                        model=candidate_model,
                        file_uri=str(getattr(uploaded_file, "uri", "")),
                        mime_type=processed_mime,
                        prompt=prompt,
                        schema_model=VideoTranscript,
                    )
                    assert isinstance(interaction_result, VideoTranscript)
                    if interaction_result.full_transcript.strip():
                        st.session_state["active_gemini_model"] = candidate_model
                        return interaction_result
                except Exception as error:
                    interaction_error = error

            # Respaldo con referencia FileData explícita y resolución baja.
            try:
                generated_result = _call_video_generate_content_fallback(
                    client=client,
                    model=candidate_model,
                    uploaded_file=uploaded_file,
                    prompt=prompt,
                    schema_model=VideoTranscript,
                )
                assert isinstance(generated_result, VideoTranscript)
                if generated_result.full_transcript.strip():
                    st.session_state["active_gemini_model"] = candidate_model
                    return generated_result
            except Exception as generate_error:
                errors.append(
                    f"{candidate_model}: Interactions: "
                    f"{str(interaction_error)[:180] if interaction_error else 'no usado'}"
                    f" | generateContent: {str(generate_error)[:260]}"
                )

                # Si conocemos la duración, el último respaldo divide incluso
                # videos cortos/medianos cuando la solicitud completa falla.
                if duration_seconds is not None:
                    try:
                        chunked_result = _transcribe_video_by_intervals(
                            client=client,
                            model=candidate_model,
                            uploaded_file=uploaded_file,
                            prompt=prompt,
                            duration_seconds=duration_seconds,
                            progress_callback=progress_callback,
                            api_key=api_key,
                            checkpoint_key=checkpoint_key or display_name,
                        )
                        st.session_state["active_gemini_model"] = candidate_model
                        return chunked_result
                    except Exception as chunk_error:
                        errors.append(
                            f"{candidate_model} por tramos: "
                            f"{str(chunk_error)[:300]}"
                        )

        duration_note = (
            f" Duración detectada: {_video_clock(duration_seconds)}."
            if duration_seconds is not None
            else (
                " No fue posible medir la duración porque ffprobe no está "
                "disponible; agrega ffmpeg en packages.txt."
            )
        )
        raise AppError(
            "Gemini procesó el archivo, pero las solicitudes de inferencia "
            "fallaron. El bot ya probó la solicitud completa y el procesamiento "
            "por tramos cortos con división automática."
            + duration_note
            + " Detalles: "
            + " | ".join(errors[-4:])[:1800]
        )

    except AppError:
        raise
    except Exception as error:
        stage_detail = (
            f"Archivo: {display_name}; tamaño: "
            f"{format_file_size(local_path.stat().st_size)}; "
            f"MIME informado: {mime_type}. "
        )
        translated = translate_gemini_error(error)
        raise AppError(stage_detail + str(translated)) from error
    finally:
        if uploaded_file is not None and getattr(uploaded_file, "name", None):
            try:
                client.files.delete(name=uploaded_file.name)
            except Exception:
                pass
        try:
            client.close()
        except Exception:
            pass


def transcribe_video_with_gemini(
    api_key: str,
    model: str,
    video_record: dict,
    progress_callback=None,
) -> VideoTranscript:
    """Transcribe un video cargado directamente en Streamlit."""

    if not video_record.get("is_video"):
        raise AppError("El archivo seleccionado no es un video compatible.")

    suffix = f".{video_record['extension']}"
    with tempfile.TemporaryDirectory() as temp_dir:
        local_path = Path(temp_dir) / f"video_origen{suffix}"
        local_path.write_bytes(video_record["content"])
        return _transcribe_local_video_path_with_gemini(
            api_key=api_key,
            model=model,
            local_path=local_path,
            display_name=video_record["name"],
            mime_type=video_record["mime_type"],
            progress_callback=progress_callback,
            checkpoint_key=str(video_record.get("fingerprint") or video_record["name"]),
        )


def transcribe_drive_video_with_gemini(
    api_key: str,
    model: str,
    drive_reference: str,
    progress_callback=None,
) -> tuple[VideoTranscript, dict]:
    """Descarga un video privado de Drive por bloques y lo transcribe.

    La sesión conserva únicamente metadatos y texto; el video temporal se elimina
    automáticamente al terminar.
    """

    with tempfile.TemporaryDirectory() as temp_dir:
        metadata, local_path = download_drive_video_to_path(
            drive_reference=drive_reference,
            destination_dir=Path(temp_dir),
            progress_callback=progress_callback,
        )
        transcript = _transcribe_local_video_path_with_gemini(
            api_key=api_key,
            model=model,
            local_path=local_path,
            display_name=metadata["name"],
            mime_type=metadata["mime_type"],
            progress_callback=progress_callback,
            checkpoint_key=(
                "drive:"
                + str(metadata.get("file_id") or "")
                + ":"
                + str(metadata.get("modifiedTime") or "")
                + ":"
                + str(metadata.get("size_bytes") or "")
            ),
        )

    source_record = {
        "role": "video de origen desde Google Drive",
        "index": 0,
        "name": metadata["name"],
        "extension": metadata["extension"],
        "mime_type": metadata["mime_type"],
        "content": b"",
        "size_bytes": metadata["size_bytes"],
        "text": "",
        "page_count": None,
        "warnings": [
            "El video se descargó por bloques desde Google Drive y no se guardó "
            "en la memoria de la sesión."
        ],
        "is_video": True,
        "origin_kind": "google_drive",
        "drive_file_id": metadata["file_id"],
        "drive_url": metadata.get("webViewLink") or drive_reference,
        "video_actions": [action.model_dump() for action in transcript.actions],
        "video_action_count": len(transcript.actions),
        "video_checkpoint_id": st.session_state.get("video_checkpoint_id"),
    }
    return transcript, source_record


def translate_gemini_error(error: Exception) -> AppError:
    """Convierte errores técnicos frecuentes en mensajes comprensibles."""

    message = str(error)
    upper_message = message.upper()

    if "API_KEY_INVALID" in upper_message or "INVALID API KEY" in upper_message:
        return AppError(
            "La clave de Gemini no es válida. Revisa .streamlit/secrets.toml."
        )

    if "RESOURCE_EXHAUSTED" in upper_message or "429" in upper_message:
        return AppError(
            "Se alcanzó el límite temporal o gratuito de Gemini. "
            "Espera unos minutos y vuelve a intentar."
        )

    if "NOT_FOUND" in upper_message or "404" in upper_message:
        return AppError(
            "El modelo configurado no está disponible. "
            "Revisa GEMINI_MODEL en secrets.toml."
        )

    if "PERMISSION_DENIED" in upper_message or "403" in upper_message:
        return AppError(
            "La clave no tiene permiso para usar Gemini o el servicio no está "
            "habilitado para ese proyecto."
        )

    if "INVALID_ARGUMENT" in upper_message or "400" in upper_message:
        return AppError(
            "Gemini rechazó la solicitud por un argumento inválido. "
            f"Detalle técnico: {message[:500]}"
        )

    return AppError(f"Gemini no pudo completar el proceso. Detalle: {message[:700]}")


def _inline_and_clean_json_schema(schema_model: type[BaseModel]) -> dict:
    """Genera un JSON Schema compatible con Gemini.

    Pydantic agrega ``additionalProperties: false`` cuando se usa
    ``ConfigDict(extra="forbid")``. Algunas versiones del endpoint antiguo
    ``response_schema`` no aceptan esa propiedad. Esta función expande las
    referencias internas y elimina metadatos incompatibles antes de enviar el
    esquema mediante ``response_json_schema``.
    """

    schema_model.model_rebuild()
    raw_schema = schema_model.model_json_schema()
    definitions = raw_schema.get("$defs", {})

    def resolve(node):
        if isinstance(node, list):
            return [resolve(item) for item in node]

        if not isinstance(node, dict):
            return node

        if "$ref" in node:
            reference = str(node["$ref"])
            prefix = "#/$defs/"
            if reference.startswith(prefix):
                definition_name = reference[len(prefix):]
                definition = definitions.get(definition_name)
                if definition is None:
                    raise AppError(
                        f"No fue posible resolver el esquema interno {reference}."
                    )

                merged = dict(definition)
                merged.update(
                    {
                        key: value
                        for key, value in node.items()
                        if key != "$ref"
                    }
                )
                return resolve(merged)

        cleaned: dict = {}
        for key, value in node.items():
            if key in {
                "$defs",
                "$schema",
                "additionalProperties",
                "examples",
                "default",
            }:
                continue
            cleaned[key] = resolve(value)

        return cleaned

    compatible_schema = resolve(raw_schema)

    if not isinstance(compatible_schema, dict):
        raise AppError("No fue posible construir el esquema de respuesta.")

    return compatible_schema


def _clean_json_text(text: str) -> str:
    """Elimina cercas Markdown que algunos modelos agregan alrededor del JSON."""

    cleaned = text.strip()

    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\s*```$", "", cleaned)

    return cleaned.strip()


def _validate_structured_output(
    schema_model: type[BaseModel],
    response,
) -> BaseModel:
    """Valida la respuesta de Gemini con el modelo Pydantic solicitado."""

    parsed = getattr(response, "parsed", None)

    if isinstance(parsed, schema_model):
        return parsed

    if isinstance(parsed, dict):
        return schema_model.model_validate(parsed)

    output_text = getattr(response, "text", None)
    if not output_text:
        raise AppError(
            "Gemini respondió, pero no devolvió contenido estructurado."
        )

    try:
        return schema_model.model_validate_json(_clean_json_text(output_text))
    except Exception as error:
        raise AppError(
            "Gemini devolvió una respuesta JSON que no cumple la estructura "
            f"requerida. Detalle: {str(error)[:500]}"
        ) from error


def _schema_fallback_instruction(
    schema_model: type[BaseModel],
    compatible_schema: dict,
) -> types.Part:
    """Crea una instrucción de respaldo si el endpoint rechaza el esquema."""

    schema_text = json.dumps(
        compatible_schema,
        ensure_ascii=False,
        separators=(",", ":"),
    )

    return types.Part.from_text(
        text=(
            "\nFORMATO JSON OBLIGATORIO DE RESPUESTA\n"
            f"Devuelve un único objeto JSON válido para {schema_model.__name__}. "
            "No uses Markdown, comentarios ni texto antes o después del JSON. "
            f"Respeta este esquema: {schema_text}"
        )
    )


def _is_schema_endpoint_error(error: Exception) -> bool:
    """Detecta errores de compatibilidad del esquema del endpoint."""

    message = str(error).upper()
    return (
        "RESPONSE_SCHEMA" in message
        or "RESPONSE_JSON_SCHEMA" in message
        or "RESPONSE_FORMAT" in message
        or "ADDITIONAL_PROPERTIES" in message
        or "ADDITIONALPROPERTIES" in message
        or "UNKNOWN NAME" in message
        or "CANNOT FIND FIELD" in message
        # Algunos endpoints devuelven únicamente un 400 genérico cuando el
        # esquema estructurado no es compatible. En ese caso se intenta el
        # modo JSON por prompt y se valida localmente con Pydantic.
        or "INVALID_ARGUMENT" in message
        or "400" in message
    )


def _normalize_model_name(name: str | None) -> str:
    """Normaliza los nombres devueltos por models.list()."""

    normalized = (name or "").strip()
    if normalized.startswith("models/"):
        normalized = normalized[len("models/"):]
    return normalized


def _is_suitable_generation_model(model_info) -> bool:
    """Filtra modelos de texto/multimodales aptos para este bot documental."""

    name = _normalize_model_name(getattr(model_info, "name", ""))
    actions = set(getattr(model_info, "supported_actions", None) or [])

    if not name or "generateContent" not in actions:
        return False

    lowered = name.lower()
    excluded_terms = (
        "embedding",
        "image",
        "live",
        "tts",
        "audio",
        "robotics",
        "computer-use",
        "veo",
        "imagen",
        "lyria",
    )
    return lowered.startswith("gemini") and not any(
        term in lowered for term in excluded_terms
    )


def _model_priority(name: str) -> tuple[int, str]:
    """Ordena primero modelos Flash de menor latencia y mayor disponibilidad."""

    lowered = name.lower()
    preferred = {
        "gemini-flash-lite-latest": 0,
        "gemini-flash-latest": 1,
        "gemini-3.1-flash-lite-preview": 2,
        "gemini-3-flash-preview": 3,
        "gemini-2.0-flash-lite": 4,
        "gemini-2.0-flash-lite-001": 5,
        "gemini-2.0-flash": 6,
        "gemini-2.0-flash-001": 7,
    }
    if lowered in preferred:
        return preferred[lowered], lowered
    if "flash-lite" in lowered and "latest" in lowered:
        return 8, lowered
    if "flash" in lowered and "latest" in lowered:
        return 9, lowered
    if "flash-lite" in lowered:
        return 10, lowered
    if "flash" in lowered and "preview" not in lowered:
        return 11, lowered
    if "flash" in lowered:
        return 12, lowered
    if "pro" in lowered and "preview" not in lowered:
        return 13, lowered
    return 20, lowered


def _discover_candidate_models(client, configured_model: str) -> list[str]:
    """Obtiene modelos aptos y los ordena para evitar quedarse en uno saturado."""

    configured = _normalize_model_name(configured_model)
    preferred_defaults = [
        configured,
        "gemini-flash-lite-latest",
        "gemini-flash-latest",
        "gemini-3.1-flash-lite-preview",
        "gemini-3-flash-preview",
        "gemini-2.0-flash-lite",
        "gemini-2.0-flash",
    ]

    discovered: list[str] = []
    try:
        for model_info in client.models.list():
            if _is_suitable_generation_model(model_info):
                name = _normalize_model_name(getattr(model_info, "name", ""))
                if name and name not in discovered:
                    discovered.append(name)
    except Exception:
        discovered = []

    ordered: list[str] = []

    def add(candidate: str) -> None:
        candidate = _normalize_model_name(candidate)
        if candidate and candidate not in ordered:
            ordered.append(candidate)

    if discovered:
        available = set(discovered)
        for candidate in preferred_defaults:
            if candidate in available:
                add(candidate)
        for candidate in sorted(discovered, key=_model_priority):
            add(candidate)
    else:
        for candidate in preferred_defaults:
            add(candidate)

    # Evita esperas excesivas si el listado contiene demasiados modelos.
    return ordered[:8]


def _is_retryable_gemini_error(error: Exception) -> bool:
    """Detecta errores temporales de capacidad, cuota o red."""

    message = str(error).upper()
    retryable_terms = (
        "503",
        "UNAVAILABLE",
        "HIGH DEMAND",
        "429",
        "RESOURCE_EXHAUSTED",
        "DEADLINE_EXCEEDED",
        "504",
        "GATEWAY_TIMEOUT",
        "CONNECTION RESET",
        "CONNECTION ERROR",
        "TIMED OUT",
        "TIMEOUT",
    )
    return any(term in message for term in retryable_terms)


def _generate_with_retries(
    client,
    model: str,
    contents: list[types.Content],
    config: types.GenerateContentConfig,
    max_attempts: int = 3,
):
    """Reintenta errores temporales con espera exponencial y variación aleatoria."""

    last_error: Exception | None = None

    for attempt in range(max_attempts):
        try:
            return client.models.generate_content(
                model=model,
                contents=contents,
                config=config,
            )
        except Exception as error:
            last_error = error
            if not _is_retryable_gemini_error(error) or attempt == max_attempts - 1:
                raise

            delay = min(8.0, 1.5 * (2 ** attempt)) + random.uniform(0.1, 0.7)
            time.sleep(delay)

    if last_error is not None:
        raise last_error
    raise AppError("Gemini no respondió y no entregó un error identificable.")


def call_gemini_structured(
    api_key: str,
    model: str,
    input_parts: list[types.Part],
    schema_model: type[BaseModel],
) -> BaseModel:
    """Ejecuta Gemini con reintentos y cambio automático de modelo."""

    if not api_key:
        raise AppError(
            "No se encontró GEMINI_API_KEY. Configúrala en "
            ".streamlit/secrets.toml."
        )

    client = genai.Client(api_key=api_key)
    compatible_schema = _inline_and_clean_json_schema(schema_model)
    candidate_models = _discover_candidate_models(client, model)

    if not candidate_models:
        raise AppError(
            "La clave fue aceptada, pero la API no devolvió modelos compatibles "
            "con generateContent para este proyecto."
        )

    attempted_models: list[str] = []
    model_errors: list[str] = []

    try:
        for candidate_model in candidate_models:
            attempted_models.append(candidate_model)
            user_content = types.Content(role="user", parts=list(input_parts))

            try:
                response = _generate_with_retries(
                    client=client,
                    model=candidate_model,
                    contents=[user_content],
                    config=types.GenerateContentConfig(
                        response_mime_type="application/json",
                        response_json_schema=compatible_schema,
                        temperature=0.1,
                    ),
                )
                st.session_state["active_gemini_model"] = candidate_model
                return _validate_structured_output(schema_model, response)

            except AppError:
                raise
            except Exception as error:
                error_message = str(error)
                upper_message = error_message.upper()
                model_errors.append(f"{candidate_model}: {error_message[:180]}")

                # Modelo retirado/no habilitado o temporalmente saturado:
                # se prueba el siguiente modelo compatible.
                if (
                    "404" in upper_message
                    or "NOT_FOUND" in upper_message
                    or _is_retryable_gemini_error(error)
                ):
                    continue

                if not _is_schema_endpoint_error(error):
                    raise translate_gemini_error(error) from error

                fallback_parts = list(input_parts)
                fallback_parts.append(
                    _schema_fallback_instruction(schema_model, compatible_schema)
                )
                fallback_content = types.Content(role="user", parts=fallback_parts)

                try:
                    fallback_response = _generate_with_retries(
                        client=client,
                        model=candidate_model,
                        contents=[fallback_content],
                        config=types.GenerateContentConfig(
                            response_mime_type="application/json",
                            temperature=0.1,
                        ),
                    )
                    st.session_state["active_gemini_model"] = candidate_model
                    return _validate_structured_output(
                        schema_model,
                        fallback_response,
                    )
                except AppError:
                    raise
                except Exception as fallback_error:
                    fallback_message = str(fallback_error)
                    fallback_upper = fallback_message.upper()
                    model_errors.append(
                        f"{candidate_model} (respaldo): {fallback_message[:180]}"
                    )
                    if (
                        "404" in fallback_upper
                        or "NOT_FOUND" in fallback_upper
                        or "400" in fallback_upper
                        or "INVALID_ARGUMENT" in fallback_upper
                        or _is_retryable_gemini_error(fallback_error)
                    ):
                        continue
                    raise translate_gemini_error(fallback_error) from fallback_error

        attempts = ", ".join(attempted_models) or "ninguno"
        detail = " | ".join(model_errors[-4:])
        raise AppError(
            "Gemini está temporalmente saturado o no respondió con los modelos "
            f"disponibles. Modelos probados: {attempts}. "
            "Espera uno o dos minutos y vuelve a presionar Analizar. "
            f"Detalle: {detail[:550]}"
        )

    finally:
        try:
            client.close()
        except Exception:
            pass


def _normalized_question_key(question: str) -> str:
    """Normaliza una pregunta para detectar duplicados semánticos simples."""

    cleaned = re.sub(r"[¿?¡!.,;:()]+", " ", question.lower())
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    stopwords = {
        "el", "la", "los", "las", "un", "una", "de", "del", "en",
        "para", "por", "que", "cuál", "cual", "qué", "como", "cómo",
    }
    tokens = [token for token in cleaned.split() if token not in stopwords]
    return " ".join(tokens)


def _deduplicate_missing_questions(analysis: GuideAnalysis) -> GuideAnalysis:
    """Elimina preguntas repetidas sin perder la primera sección que las requiere."""

    seen: set[str] = set()

    for section in analysis.sections:
        unique_questions: list[MissingQuestion] = []

        for question in section.missing_questions:
            key = _normalized_question_key(question.question)
            if not key or key in seen:
                continue

            seen.add(key)
            question.category = (
                question.category.strip() or "Información del proceso"
            )
            unique_questions.append(question)

        section.missing_questions = unique_questions

    return analysis


def _strip_list_prefix(text: str) -> str:
    """Elimina numeración sin borrar la separación entre título y descripción."""

    cleaned = str(text or "").strip()
    cleaned = re.sub(r"^\s*[•●▪◦\-–—]\s*", "", cleaned)
    cleaned = re.sub(r"^\s*\(?\d+\)?[.)\-:]\s*", "", cleaned)
    lines = [
        re.sub(r"[ \t]+", " ", line).strip()
        for line in cleaned.splitlines()
        if line.strip()
    ]
    return "\n".join(lines).strip()


def normalize_final_document_structure(
    final_document: FinalDocument,
) -> FinalDocument:
    """Normaliza listas, espacios, tablas y orden antes de generar.

    Conserva todas las ocurrencias de los pasos numerados, incluso cuando
    dos acciones tengan el mismo texto, porque pueden corresponder a
    momentos distintos del video.
    """

    normalized_sections: list[FinalSection] = []

    def explicit_number(value: str) -> int | None:
        match = re.match(
            r"^\s*[•●▪◦\-–—]?\s*\(?(\d+)\)?[.)\-:]\s+",
            value,
        )
        return int(match.group(1)) if match else None

    for section in sorted(final_document.sections, key=lambda item: item.order):
        title = re.sub(r"\s+", " ", section.title).strip()

        paragraphs = [
            re.sub(r"[ \t]+", " ", paragraph).strip()
            for paragraph in section.paragraphs
            if paragraph.strip()
        ]

        title_key = title.lower()
        sequential_section = any(
            token in title_key
            for token in (
                "paso a paso",
                "procedimiento",
                "actividades",
                "metodología",
                "metodologia",
                "etapas",
            )
        )

        numbered_candidates: list[tuple[int | None, int, int, str]] = []
        remaining_bullets: list[str] = []

        for item_index, original in enumerate(section.numbered_items):
            cleaned = _strip_list_prefix(original)
            if cleaned:
                numbered_candidates.append(
                    (
                        explicit_number(original),
                        0,
                        item_index,
                        cleaned,
                    )
                )

        for item_index, original in enumerate(section.bullets):
            cleaned = _strip_list_prefix(original)
            if not cleaned:
                continue

            detected_number = explicit_number(original)
            if detected_number is not None or sequential_section:
                numbered_candidates.append(
                    (
                        detected_number,
                        1,
                        item_index,
                        cleaned,
                    )
                )
            else:
                remaining_bullets.append(cleaned)

        if any(item[0] is not None for item in numbered_candidates):
            numbered_candidates.sort(
                key=lambda item: (
                    item[0] is None,
                    item[0] if item[0] is not None else 10**9,
                    item[1],
                    item[2],
                )
            )
        else:
            numbered_candidates.sort(key=lambda item: (item[1], item[2]))

        # IMPORTANTE: no eliminar pasos repetidos.
        # En un procedimiento basado en video, una misma acción puede ejecutarse
        # varias veces en momentos distintos. Cada aparición corresponde a una
        # acción ACC-XXXX independiente y debe conservarse como un paso.
        numbered_items = [
            item[3] for item in numbered_candidates
        ]
        bullets = list(dict.fromkeys(remaining_bullets))

        tables: list[FinalTable] = []
        for table in section.tables:
            headers = [
                re.sub(r"\s+", " ", str(header)).strip()
                for header in table.headers
            ]
            rows = [
                [
                    re.sub(r"\s+", " ", str(value)).strip()
                    for value in row
                ]
                for row in table.rows
            ]
            tables.append(
                table.model_copy(
                    update={
                        "title": re.sub(r"\s+", " ", table.title).strip(),
                        "headers": headers,
                        "rows": rows,
                    }
                )
            )

        normalized_sections.append(
            section.model_copy(
                update={
                    "title": title,
                    "paragraphs": paragraphs,
                    "bullets": bullets,
                    "numbered_items": numbered_items,
                    "tables": tables,
                }
            )
        )

    return final_document.model_copy(
        update={
            "title": re.sub(r"\s+", " ", final_document.title).strip(),
            "subtitle": re.sub(r"\s+", " ", final_document.subtitle).strip(),
            "sections": normalized_sections,
        }
    )


def audit_final_document_with_gemini(
    api_key: str,
    model: str,
    guides: list[dict],
    source: dict,
    analysis: GuideAnalysis,
    final_document: FinalDocument,
) -> AuditReport:
    """Audita el documento editado contra las guías y el origen."""

    relevant_indices = [analysis.selected_guide_index]
    relevant_indices.extend(analysis.supporting_guide_indices)

    parts: list[types.Part] = []
    for index in relevant_indices:
        append_record_to_gemini_input(
            parts,
            guides[index],
            f"GUÍA PARA AUDITORÍA {index}",
        )

    append_record_to_gemini_input(parts, source, "DOCUMENTO DE ORIGEN")
    parts.append(
        types.Part.from_text(
            text=audit_prompt(analysis, final_document)
        )
    )

    result = call_gemini_structured(
        api_key=api_key,
        model=model,
        input_parts=parts,
        schema_model=AuditReport,
    )

    assert isinstance(result, AuditReport)
    return result


def regenerate_section_with_gemini(
    api_key: str,
    model: str,
    guides: list[dict],
    source: dict,
    analysis: GuideAnalysis,
    answers: list[UserAnswer],
    final_document: FinalDocument,
    section_order: int,
) -> FinalSection:
    """Regenera una sección específica sin rehacer el documento completo."""

    section = next(
        (
            item
            for item in final_document.sections
            if item.order == section_order
        ),
        None,
    )
    if section is None:
        raise AppError("No se encontró la sección seleccionada para regenerar.")

    relevant_indices = [analysis.selected_guide_index]
    relevant_indices.extend(analysis.supporting_guide_indices)

    parts: list[types.Part] = []
    for index in relevant_indices:
        append_record_to_gemini_input(
            parts,
            guides[index],
            f"GUÍA PARA REGENERACIÓN {index}",
        )

    append_record_to_gemini_input(parts, source, "DOCUMENTO DE ORIGEN")
    parts.append(
        types.Part.from_text(
            text=render_fidelity_manifest(build_guide_fidelity_manifest(guides))
        )
    )
    parts.append(
        types.Part.from_text(
            text=section_regeneration_prompt(
                analysis,
                answers,
                final_document,
                section,
            )
        )
    )

    result = call_gemini_structured(
        api_key=api_key,
        model=model,
        input_parts=parts,
        schema_model=FinalSection,
    )

    assert isinstance(result, FinalSection)
    result.order = section_order
    return result

def analyze_guides_and_source(
    api_key: str,
    model: str,
    guides: list[dict],
    source: dict,
) -> GuideAnalysis:
    """Selecciona la guía, evalúa el origen y depura preguntas repetidas."""

    parts: list[types.Part] = []

    for index, guide in enumerate(guides):
        append_record_to_gemini_input(parts, guide, f"GUÍA {index}")

    append_record_to_gemini_input(parts, source, "DOCUMENTO DE ORIGEN")
    parts.append(
        types.Part.from_text(
            text=render_fidelity_manifest(build_guide_fidelity_manifest(guides))
        )
    )
    parts.append(
        types.Part.from_text(
            text=(
                "Después de seleccionar la guía principal, conserva literalmente "
                "todos sus títulos y su orden. No fusiones secciones."
            )
        )
    )
    parts.append(
        types.Part.from_text(
            text=analysis_prompt(
                [guide["name"] for guide in guides],
                source["name"],
            )
        )
    )

    result = call_gemini_structured(
        api_key=api_key,
        model=model,
        input_parts=parts,
        schema_model=GuideAnalysis,
    )

    assert isinstance(result, GuideAnalysis)

    if not 0 <= result.selected_guide_index < len(guides):
        raise AppError("Gemini seleccionó un índice de guía inválido.")

    result.selected_guide_name = guides[result.selected_guide_index]["name"]

    valid_supporting: list[int] = []
    for index in result.supporting_guide_indices:
        if (
            0 <= index < len(guides)
            and index != result.selected_guide_index
            and index not in valid_supporting
        ):
            valid_supporting.append(index)
    result.supporting_guide_indices = valid_supporting

    result.sections = sorted(result.sections, key=lambda section: section.order)
    result = lock_analysis_to_selected_guide(result, guides)
    result = _deduplicate_missing_questions(result)

    if not result.sections:
        raise AppError(
            "Gemini no identificó secciones en la guía seleccionada. "
            "Verifica que la guía describa el contenido esperado."
        )

    return result


def generate_final_with_gemini(
    api_key: str,
    model: str,
    guides: list[dict],
    source: dict,
    analysis: GuideAnalysis,
    answers: list[UserAnswer],
) -> FinalDocument:
    """Redacta el borrador y normaliza su estructura documental."""

    relevant_indices = [analysis.selected_guide_index]
    relevant_indices.extend(analysis.supporting_guide_indices)

    parts: list[types.Part] = []

    for index in relevant_indices:
        append_record_to_gemini_input(
            parts,
            guides[index],
            f"GUÍA APLICABLE {index}",
        )

    append_record_to_gemini_input(parts, source, "DOCUMENTO DE ORIGEN")
    video_actions = _source_video_actions(source)
    if video_actions:
        parts.append(
            types.Part.from_text(
                text=(
                    f"CONTROL DE COBERTURA: el origen contiene {len(video_actions)} "
                    "acciones ACC-XXXX. En una sección procedimental dinámica, "
                    "ninguna puede omitirse ni fusionarse."
                )
            )
        )
    parts.append(
        types.Part.from_text(
            text=render_exact_guide_structure(guides, analysis.selected_guide_index)
        )
    )
    parts.append(
        types.Part.from_text(
            text=render_fidelity_manifest(build_guide_fidelity_manifest(guides))
        )
    )
    parts.append(
        types.Part.from_text(
            text=generation_prompt(analysis, answers)
        )
    )

    result = call_gemini_structured(
        api_key=api_key,
        model=model,
        input_parts=parts,
        schema_model=FinalDocument,
    )

    assert isinstance(result, FinalDocument)
    result = normalize_final_document_structure(result)
    result = enforce_fidelity_with_gemini(
        api_key=api_key,
        model=model,
        guides=guides,
        source=source,
        analysis=analysis,
        final_document=result,
    )

    if not result.sections:
        raise AppError("Gemini no generó las secciones del documento final.")

    return result


# =========================================================
# GENERACIÓN DE DOCX Y PDF
# =========================================================


PAGE_TOTAL_PATTERN = re.compile(
    r"(?i)\b(p[áa]gina)(\s*:?\s*)\d+\s*(?:de|/)\s*\d+\b"
)
PAGE_SINGLE_PATTERN = re.compile(
    r"(?i)\b(p[áa]gina)(\s*:?\s*)\d+\b"
)


def _set_update_fields_on_open(document: Document) -> None:
    """Solicita a Word/LibreOffice actualizar PAGE y NUMPAGES."""

    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def _copy_run_format(run, source_rpr) -> None:
    if source_rpr is not None:
        run._r.insert(0, deepcopy(source_rpr))


def _append_word_field(paragraph, field_code: str, source_rpr=None) -> None:
    """Inserta un campo de Word, por ejemplo PAGE o NUMPAGES."""

    run = paragraph.add_run()
    _copy_run_format(run, source_rpr)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = f" {field_code} "

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    display = OxmlElement("w:t")
    display.text = "1"

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.extend([begin, instruction, separate, display, end])


def _clear_paragraph_content(paragraph) -> None:
    """Limpia runs conservando las propiedades del párrafo."""

    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def _replace_page_markers_in_paragraph(paragraph) -> bool:
    """Reemplaza paginación estática por campos PAGE/NUMPAGES."""

    original_text = paragraph.text
    if not original_text.strip():
        return False

    tokenized = PAGE_TOTAL_PATTERN.sub(
        lambda match: (
            f"{match.group(1)}{match.group(2)}"
            "[[PAGE_CURRENT]] de [[PAGE_TOTAL]]"
        ),
        original_text,
    )

    if "[[PAGE_CURRENT]]" not in tokenized:
        tokenized = PAGE_SINGLE_PATTERN.sub(
            lambda match: (
                f"{match.group(1)}{match.group(2)}[[PAGE_CURRENT]]"
            ),
            tokenized,
        )

    if "[[PAGE_CURRENT]]" not in tokenized:
        return False

    source_rpr = None
    if paragraph.runs and paragraph.runs[0]._r.rPr is not None:
        source_rpr = deepcopy(paragraph.runs[0]._r.rPr)

    _clear_paragraph_content(paragraph)

    tokens = re.split(
        r"(\[\[PAGE_CURRENT\]\]|\[\[PAGE_TOTAL\]\])",
        tokenized,
    )
    for token in tokens:
        if not token:
            continue
        if token == "[[PAGE_CURRENT]]":
            _append_word_field(paragraph, "PAGE", source_rpr)
        elif token == "[[PAGE_TOTAL]]":
            _append_word_field(paragraph, "NUMPAGES", source_rpr)
        else:
            run = paragraph.add_run(token)
            _copy_run_format(run, source_rpr)

    return True


def _iter_container_paragraphs(container):
    for paragraph in container.paragraphs:
        yield paragraph

    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_container_paragraphs(cell)


def apply_dynamic_page_numbering(document: Document) -> int:
    """
    Convierte textos como “PÁGINA: 1 de 2” y “Página 1” en campos
    dinámicos que se actualizan al abrir o convertir el archivo.
    """

    updated = 0
    visited_parts: set[int] = set()

    for section in document.sections:
        containers = (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        )

        for container in containers:
            part_id = id(container.part)
            if part_id in visited_parts:
                continue
            visited_parts.add(part_id)

            for paragraph in _iter_container_paragraphs(container):
                if _replace_page_markers_in_paragraph(paragraph):
                    updated += 1

    _set_update_fields_on_open(document)
    return updated


def _set_row_repeat_header(row) -> None:
    """Repite la primera fila de una tabla al pasar de página."""

    tr_pr = row._tr.get_or_add_trPr()
    element = tr_pr.find(qn("w:tblHeader"))
    if element is None:
        element = OxmlElement("w:tblHeader")
        tr_pr.append(element)
    element.set(qn("w:val"), "true")


def _set_row_cant_split(row) -> None:
    """Evita que una fila se parta entre dos páginas."""

    tr_pr = row._tr.get_or_add_trPr()
    element = tr_pr.find(qn("w:cantSplit"))
    if element is None:
        element = OxmlElement("w:cantSplit")
        tr_pr.append(element)


def _normalize_document_text(text: str) -> str:
    return re.sub(r"\s+", " ", str(text)).strip()


def _format_body_paragraph(paragraph, justify: bool = True) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.widow_control = True
    if justify:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _format_heading_paragraph(paragraph) -> None:
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.widow_control = True

def has_docx_style(document: Document, style_name: str) -> bool:
    """Indica si el estilo está materializado dentro del DOCX seleccionado."""

    try:
        document.styles[style_name]
        return True
    except (KeyError, ValueError):
        return False


def configure_docx(document: Document, preserve_guide_layout: bool = False) -> None:
    """Configura tipografía y espaciado base sin romper la guía institucional."""

    if not preserve_guide_layout:
        for section in document.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

    if has_docx_style(document, "Normal"):
        normal_style = document.styles["Normal"]
        if not normal_style.font.name:
            normal_style.font.name = "Arial"
        if not normal_style.font.size:
            normal_style.font.size = Pt(10.5)
        normal_style.paragraph_format.space_after = Pt(6)
        normal_style.paragraph_format.line_spacing = 1.15
        normal_style.paragraph_format.widow_control = True

    for style_name in ("Title", "Heading 1", "Heading 2"):
        if has_docx_style(document, style_name):
            style = document.styles[style_name]
            if not style.font.name:
                style.font.name = "Arial"
            style.paragraph_format.keep_with_next = True
            style.paragraph_format.widow_control = True


def add_safe_heading(document: Document, text: str):
    """Agrega un título unido al primer elemento de su sección."""

    clean_text = _normalize_document_text(text)
    if not clean_text:
        return None

    if has_docx_style(document, "Heading 1"):
        paragraph = document.add_paragraph(clean_text, style="Heading 1")
        _format_heading_paragraph(paragraph)
        return paragraph

    paragraph = document.add_paragraph()
    _format_heading_paragraph(paragraph)
    run = paragraph.add_run(clean_text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(12)
    return paragraph


def add_safe_bullet(document: Document, text: str):
    """Agrega una viñeta limpia sin depender de estilos opcionales."""

    clean_text = _strip_list_prefix(text)
    if not clean_text:
        return None

    if has_docx_style(document, "List Bullet"):
        paragraph = document.add_paragraph(clean_text, style="List Bullet")
    else:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.65)
        paragraph.paragraph_format.first_line_indent = Cm(-0.35)
        bullet_run = paragraph.add_run("• ")
        bullet_run.bold = True
        paragraph.add_run(clean_text)

    _format_body_paragraph(paragraph, justify=False)
    paragraph.paragraph_format.keep_together = True
    return paragraph


def add_safe_numbered(document: Document, text: str, number: int):
    """Agrega cada paso con título en negrita y descripción en línea aparte."""

    clean_text = _strip_list_prefix(text)
    if not clean_text:
        return None

    title, body = _split_step_title_body(clean_text)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.72)
    paragraph.paragraph_format.first_line_indent = Cm(-0.52)

    paragraph.add_run(f"{number}. ").bold = True
    if title:
        title_run = paragraph.add_run(title.rstrip(":") + ":")
        title_run.bold = True
        if body:
            paragraph.add_run().add_break()
            paragraph.add_run(body)
    else:
        paragraph.add_run(body or clean_text)

    _format_body_paragraph(paragraph, justify=False)
    paragraph.paragraph_format.keep_together = True
    return paragraph

def apply_safe_table_borders(table) -> None:
    """Dibuja bordes aunque la guía no contenga el estilo Table Grid."""

    table_properties = table._tbl.tblPr
    borders = table_properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_properties.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        border = borders.find(qn(tag))
        if border is None:
            border = OxmlElement(tag)
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "B7B7B7")


def clear_docx_body_keep_layout(document: Document) -> None:
    """
    Elimina el contenido corporal de una guía DOCX, pero conserva sus secciones,
    encabezados, pies de página, logos, márgenes y configuración de página.

    La propiedad sectPr debe permanecer porque enlaza encabezados y pies.
    """

    body = document._element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def load_visual_guide_document(style_guide: dict | None) -> tuple[Document, bool]:
    """Crea el documento base usando la guía DOCX seleccionada cuando sea posible."""

    if style_guide and style_guide.get("extension") == "docx":
        try:
            document = Document(BytesIO(style_guide["content"]))
            clear_docx_body_keep_layout(document)
            return document, True
        except Exception:
            # Si la guía está dañada o usa una estructura no compatible, se usa
            # un documento limpio para no bloquear la generación del contenido.
            pass

    return Document(), False


def add_docx_table(document: Document, table_data: FinalTable) -> None:
    """Agrega una tabla con alineación, filas repetibles y altura automática."""

    if table_data.title.strip():
        title_paragraph = document.add_paragraph()
        _format_heading_paragraph(title_paragraph)
        title_run = title_paragraph.add_run(
            _normalize_document_text(table_data.title)
        )
        title_run.bold = True

    column_count = max(
        len(table_data.headers),
        max((len(row) for row in table_data.rows), default=0),
    )

    if column_count == 0:
        return

    table = document.add_table(rows=1, cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    if has_docx_style(document, "Table Grid"):
        table.style = "Table Grid"
    else:
        apply_safe_table_borders(table)

    headers = [
        _normalize_document_text(header)
        for header in table_data.headers
    ]
    headers.extend([""] * (column_count - len(headers)))

    header_row = table.rows[0]
    _set_row_repeat_header(header_row)
    _set_row_cant_split(header_row)

    for index, header in enumerate(headers):
        cell = header_row.cells[index]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.keep_together = True
        run = paragraph.add_run(header) if not paragraph.text else None
        if run is None:
            paragraph.text = header
        for existing_run in paragraph.runs:
            existing_run.bold = True
            existing_run.font.name = "Arial"
            existing_run.font.size = Pt(9)

    normalized_headers = [
        re.sub(r"[^a-záéíóúñ]", "", header.lower())
        for header in headers
    ]

    for source_row in table_data.rows:
        row_values = [
            _normalize_document_text(value)
            for value in source_row
        ]
        row_values.extend([""] * (column_count - len(row_values)))

        row = table.add_row()
        _set_row_cant_split(row)

        for index, value in enumerate(row_values):
            cell = row.cells[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.text = value
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.keep_together = True

            header_key = normalized_headers[index] if index < len(normalized_headers) else ""
            if header_key in {"fecha", "versión", "version"}:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            for run in paragraph.runs:
                run.font.name = "Arial"
                run.font.size = Pt(9)

    spacer = document.add_paragraph("")
    spacer.paragraph_format.space_after = Pt(4)


def create_docx(
    final_document: FinalDocument,
    style_guide: dict | None = None,
) -> bytes:
    """
    Crea el Word final y conserva el diseño institucional de la guía DOCX.
    También convierte la paginación estática del encabezado y pie en campos
    PAGE/NUMPAGES dinámicos.
    """

    final_document = normalize_final_document_structure(final_document)
    document, inherited_layout = load_visual_guide_document(style_guide)
    configure_docx(document, preserve_guide_layout=inherited_layout)
    apply_dynamic_page_numbering(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    title.paragraph_format.keep_with_next = True
    title_run = title.add_run(final_document.title.strip())
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(15)

    if final_document.subtitle.strip():
        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_after = Pt(10)
        subtitle_run = subtitle.add_run(final_document.subtitle.strip())
        subtitle_run.italic = True
        subtitle_run.font.name = "Arial"
        subtitle_run.font.size = Pt(10.5)

    if final_document.introductory_note.strip():
        introductory = document.add_paragraph(
            final_document.introductory_note.strip()
        )
        _format_body_paragraph(introductory)

    for section in final_document.sections:
        add_safe_heading(document, section.title)

        first_content_paragraph = None

        for paragraph_text in section.paragraphs:
            if paragraph_text.strip():
                paragraph = document.add_paragraph(paragraph_text.strip())
                _format_body_paragraph(paragraph)
                if first_content_paragraph is None:
                    first_content_paragraph = paragraph

        for number, item in enumerate(section.numbered_items, start=1):
            paragraph = add_safe_numbered(document, item, number)
            if first_content_paragraph is None and paragraph is not None:
                first_content_paragraph = paragraph

        for bullet in section.bullets:
            paragraph = add_safe_bullet(document, bullet)
            if first_content_paragraph is None and paragraph is not None:
                first_content_paragraph = paragraph

        for table_data in section.tables:
            add_docx_table(document, table_data)

    _set_update_fields_on_open(document)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def convert_docx_to_pdf(docx_content: bytes) -> bytes | None:
    """Convierte el DOCX a PDF con LibreOffice para conservar encabezado y pie."""

    libreoffice = shutil.which("libreoffice") or shutil.which("soffice")
    if not libreoffice:
        return None

    try:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            docx_path = temp_path / "documento_generado.docx"
            pdf_path = temp_path / "documento_generado.pdf"
            docx_path.write_bytes(docx_content)

            result = subprocess.run(
                [
                    libreoffice,
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(temp_path),
                    str(docx_path),
                ],
                capture_output=True,
                text=True,
                timeout=120,
                check=False,
            )

            if result.returncode == 0 and pdf_path.exists():
                pdf_bytes = pdf_path.read_bytes()
                if pdf_bytes.startswith(b"%PDF"):
                    return pdf_bytes
    except Exception:
        return None

    return None


def pdf_paragraph(text: str, style: ParagraphStyle) -> Paragraph:
    safe_text = escape(text).replace("\n", "<br/>")
    return Paragraph(safe_text, style)


class NumberedCanvas(pdf_canvas.Canvas):
    """Canvas de respaldo con numeración dinámica Página X de Y."""

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states: list[dict] = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        page_count = len(self._saved_page_states)
        for page_number, state in enumerate(self._saved_page_states, start=1):
            self.__dict__.update(state)
            self.setFont("Helvetica", 8)
            self.drawCentredString(
                A4[0] / 2,
                1.25 * cm,
                f"Página {page_number} de {page_count}",
            )
            super().showPage()
        super().save()

def create_pdf(final_document: FinalDocument) -> bytes:
    """Crea una versión PDF estructurada con numeración dinámica."""

    final_document = normalize_final_document_structure(final_document)
    output = BytesIO()
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "DocumentTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=15,
        leading=18,
        alignment=TA_CENTER,
        spaceAfter=8,
    )
    subtitle_style = ParagraphStyle(
        "DocumentSubtitle",
        parent=styles["Normal"],
        fontName="Helvetica-Oblique",
        fontSize=10,
        leading=13,
        alignment=TA_CENTER,
        spaceAfter=12,
    )
    heading_style = ParagraphStyle(
        "SectionHeading",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=11.5,
        leading=14,
        spaceBefore=10,
        spaceAfter=5,
        keepWithNext=True,
    )
    body_style = ParagraphStyle(
        "BodyTextCustom",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=13,
        spaceAfter=5,
        alignment=0,
        allowWidows=0,
        allowOrphans=0,
    )
    bullet_style = ParagraphStyle(
        "BulletCustom",
        parent=body_style,
        leftIndent=14,
        firstLineIndent=-8,
        bulletIndent=4,
    )
    numbered_style = ParagraphStyle(
        "NumberedCustom",
        parent=body_style,
        leftIndent=18,
        firstLineIndent=-14,
    )
    table_cell_style = ParagraphStyle(
        "TableCell",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=8,
        leading=10,
        alignment=0,
    )
    table_center_style = ParagraphStyle(
        "TableCellCenter",
        parent=table_cell_style,
        alignment=TA_CENTER,
    )
    table_header_style = ParagraphStyle(
        "TableHeader",
        parent=table_cell_style,
        fontName="Helvetica-Bold",
        alignment=TA_CENTER,
    )

    story: list = [
        pdf_paragraph(final_document.title.strip(), title_style),
    ]

    if final_document.subtitle.strip():
        story.append(
            pdf_paragraph(final_document.subtitle.strip(), subtitle_style)
        )

    if final_document.introductory_note.strip():
        story.append(
            pdf_paragraph(final_document.introductory_note.strip(), body_style)
        )
        story.append(Spacer(1, 4))

    available_width = A4[0] - 5 * cm

    for section in final_document.sections:
        section_story: list = [
            pdf_paragraph(section.title.strip(), heading_style),
        ]

        for paragraph_text in section.paragraphs:
            if paragraph_text.strip():
                section_story.append(
                    pdf_paragraph(paragraph_text.strip(), body_style)
                )

        for index, item in enumerate(section.numbered_items, start=1):
            if item.strip():
                step_title, step_body = _split_step_title_body(item)
                if step_title:
                    step_html = (
                        f"<b>{index}. {escape(step_title.rstrip(':'))}:</b>"
                        + (f"<br/>{escape(step_body)}" if step_body else "")
                    )
                else:
                    step_html = f"<b>{index}.</b> {escape(step_body)}"
                section_story.append(Paragraph(step_html, numbered_style))

        for bullet in section.bullets:
            if bullet.strip():
                section_story.append(
                    Paragraph(
                        f"• {escape(_strip_list_prefix(bullet))}",
                        bullet_style,
                    )
                )

        if len(section_story) >= 2:
            story.append(KeepTogether(section_story[:2]))
            story.extend(section_story[2:])
        else:
            story.extend(section_story)

        for table_data in section.tables:
            if table_data.title.strip():
                story.append(
                    pdf_paragraph(table_data.title.strip(), body_style)
                )

            column_count = max(
                len(table_data.headers),
                max((len(row) for row in table_data.rows), default=0),
            )
            if column_count == 0:
                continue

            headers = [
                _normalize_document_text(header)
                for header in table_data.headers
            ]
            headers.extend([""] * (column_count - len(headers)))

            normalized_headers = [
                re.sub(r"[^a-záéíóúñ]", "", header.lower())
                for header in headers
            ]

            rows = [
                headers,
                *[
                    [
                        _normalize_document_text(value)
                        for value in list(row)
                    ]
                    + [""] * (column_count - len(row))
                    for row in table_data.rows
                ],
            ]

            formatted_rows: list[list[Paragraph]] = []
            for row_index, row in enumerate(rows):
                formatted_cells: list[Paragraph] = []
                for column_index, value in enumerate(row):
                    if row_index == 0:
                        cell_style = table_header_style
                    elif normalized_headers[column_index] in {
                        "fecha",
                        "versión",
                        "version",
                    }:
                        cell_style = table_center_style
                    else:
                        cell_style = table_cell_style
                    formatted_cells.append(
                        pdf_paragraph(str(value), cell_style)
                    )
                formatted_rows.append(formatted_cells)

            col_widths = [available_width / column_count] * column_count
            table = Table(
                formatted_rows,
                colWidths=col_widths,
                repeatRows=1,
                hAlign="LEFT",
                splitByRow=1,
            )
            table.setStyle(
                TableStyle(
                    [
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 5),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                        ("TOPPADDING", (0, 0), (-1, -1), 4),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ]
                )
            )
            story.append(table)
            story.append(Spacer(1, 8))

    pdf = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=2.5 * cm,
        leftMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.1 * cm,
        title=final_document.title,
        author="Generador inteligente de documentos",
    )
    pdf.build(story, canvasmaker=NumberedCanvas)
    return output.getvalue()


def output_filename(title: str, suffix: str) -> str:
    """Genera un nombre de archivo seguro."""

    normalized = re.sub(r"[^A-Za-z0-9ÁÉÍÓÚáéíóúÑñ_-]+", "_", title).strip("_")
    normalized = normalized[:80] or "documento_generado"
    return f"{normalized}.{suffix}"


def save_generated_file(filename: str, content: bytes) -> Path:
    """Guarda una copia local como respaldo de descarga en Codespaces."""

    output_dir = Path("output")
    output_dir.mkdir(parents=True, exist_ok=True)
    path = output_dir / Path(filename).name
    path.write_bytes(content)
    return path


def direct_download_link(
    label: str,
    data: bytes,
    filename: str,
    mime_type: str,
) -> str:
    """Crea un enlace de descarga embebido que no depende del endpoint temporal de Streamlit."""

    encoded = base64.b64encode(data).decode("ascii")
    safe_filename = html.escape(Path(filename).name, quote=True)
    safe_label = html.escape(label)

    return (
        '<a download="' + safe_filename + '" '
        'href="data:' + mime_type + ';base64,' + encoded + '" '
        'style="display:block;text-align:center;padding:0.65rem 1rem;'
        'border-radius:0.5rem;border:1px solid rgba(250,250,250,0.25);'
        'text-decoration:none;font-weight:600;color:inherit;">'
        + safe_label + '</a>'
    )


# =========================================================
# ESTADO Y COMPONENTES DE INTERFAZ
# =========================================================


def initialize_state() -> None:
    defaults = {
        "guide_records": None,
        "source_record": None,
        "analysis": None,
        "final_document": None,
        "docx_output": None,
        "pdf_output": None,
        "answers": None,
        "audit_summary": None,
        "upload_key": 0,
        "active_gemini_model": None,
        "video_transcript": None,
        "video_source_text": None,
        "video_source_fingerprint": None,
        "video_checkpoint_id": None,
    }

    for key, default in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = default


def reset_workflow() -> None:
    st.session_state.guide_records = None
    st.session_state.source_record = None
    st.session_state.analysis = None
    st.session_state.final_document = None
    st.session_state.docx_output = None
    st.session_state.pdf_output = None
    st.session_state.answers = None
    st.session_state.audit_summary = None
    st.session_state.active_gemini_model = None
    st.session_state.video_transcript = None
    st.session_state.video_source_text = None
    st.session_state.video_source_fingerprint = None
    st.session_state.video_checkpoint_id = None
    st.session_state.upload_key += 1


def count_required_questions(analysis: GuideAnalysis) -> int:
    return sum(
        1
        for section in analysis.sections
        for question in section.missing_questions
        if question.required
    )


def show_file_record(record: dict, title: str, key: str) -> None:
    """Muestra información y vista previa de un archivo."""

    with st.expander(title):
        columns = st.columns(4)
        columns[0].metric("Formato", record["extension"].upper())
        columns[1].metric("Tamaño", format_file_size(record["size_bytes"]))
        columns[2].metric("Caracteres", len(record["text"]))
        columns[3].metric(
            "Páginas",
            record["page_count"] if record["page_count"] is not None else "N/A",
        )

        for warning in record["warnings"]:
            st.warning(warning)

        preview = record["text"][:6_000]
        st.text_area(
            "Vista previa local",
            value=preview or "No hay texto local para mostrar.",
            height=220,
            disabled=True,
            key=f"preview_{key}",
        )




def show_video_transcript_review(source_record: dict) -> str | None:
    """Permite revisar y corregir el origen textual generado desde el video."""

    st.divider()
    st.subheader("🎬 Revisar transcripción del video")
    st.markdown(
        """<div class="friendly-note"><strong>Revisión obligatoria:</strong>
        confirma nombres, cargos, fechas, cifras, sistemas y fragmentos marcados
        como inaudibles. El texto aprobado será tratado como el documento de
        origen del proceso.</div>""",
        unsafe_allow_html=True,
    )

    with st.expander("Información del video", expanded=False):
        st.write(f"**Archivo:** {source_record.get('name', 'Video')}")
        st.write(
            f"**Tamaño:** {format_file_size(int(source_record.get('size_bytes') or 0))}"
        )
        if source_record.get("origin_kind") == "google_drive":
            st.success(
                "El video provino de Google Drive. Solo se conservan sus "
                "metadatos y la transcripción; el archivo temporal ya fue eliminado."
            )
            drive_url = str(source_record.get("drive_url") or "").strip()
            if drive_url:
                st.link_button(
                    "Abrir el video original en Google Drive",
                    drive_url,
                    width="stretch",
                )
        elif source_record.get("content"):
            st.video(
                source_record["content"],
                format=source_record["mime_type"],
            )
        else:
            st.caption(
                "El video no se conserva en memoria para reducir el consumo de recursos."
            )

    transcript = st.session_state.get("video_transcript")
    if transcript:
        metric_columns = st.columns(3)
        metric_columns[0].metric("Acciones operativas", len(transcript.actions))
        metric_columns[1].metric("Hablantes", len(transcript.speakers))
        metric_columns[2].metric("Alertas", len(transcript.uncertainties))

        if transcript.actions:
            with st.expander("Inventario de acciones detectadas", expanded=True):
                st.caption(
                    "Cada registro debe convertirse en un paso o fila independiente "
                    "cuando la guía solicite PASO A PASO o PROCEDIMIENTO."
                )
                st.dataframe(
                    [
                        {
                            "ID": action.action_id,
                            "Tiempo": action.timestamp_start,
                            "Responsable": action.actor,
                            "Sistema/Ruta": " — ".join(
                                value for value in (action.system, action.location_path)
                                if value.strip()
                            ),
                            "Acción": action.action,
                            "Validación/Resultado": action.validation or action.result,
                        }
                        for action in transcript.actions
                    ],
                    width="stretch",
                    hide_index=True,
                )
        else:
            st.error(
                "No se detectaron acciones operativas. Vuelve a transcribir el video "
                "antes de generar un procedimiento."
            )

    if transcript and transcript.uncertainties:
        with st.expander("Fragmentos que requieren confirmación", expanded=True):
            for uncertainty in transcript.uncertainties:
                st.warning(uncertainty)

    editor_key = f"video_source_editor_{st.session_state.upload_key}"
    edited_text = st.text_area(
        "Transcripción y evidencia visual",
        value=st.session_state.get("video_source_text") or source_record.get("text", ""),
        height=520,
        key=editor_key,
        help=(
            "Puedes corregir errores de transcripción. No elimines hechos ni "
            "agregues información que no aparezca en el video."
        ),
    )

    columns = st.columns(2)
    approved = columns[0].button(
        "Aprobar transcripción y analizar guías",
        type="primary",
        width="stretch",
        disabled=not edited_text.strip(),
    )
    retranscribe = columns[1].button(
        "Volver a transcribir el video",
        width="stretch",
    )

    if retranscribe:
        checkpoint_id = str(
            source_record.get("video_checkpoint_id")
            or st.session_state.get("video_checkpoint_id")
            or ""
        ).strip()
        if checkpoint_id:
            _clear_video_checkpoints(checkpoint_id)
        st.session_state.video_transcript = None
        st.session_state.video_source_text = None
        st.session_state.video_source_fingerprint = None
        st.session_state.video_checkpoint_id = None
        st.session_state.source_record = None
        st.session_state.analysis = None
        st.rerun()

    if approved:
        return edited_text.strip()

    return None


def _clear_outputs_after_new_analysis() -> None:
    st.session_state.final_document = None
    st.session_state.docx_output = None
    st.session_state.pdf_output = None
    st.session_state.answers = None
    st.session_state.audit_summary = None


def show_analysis(analysis: GuideAnalysis, guides: list[dict]) -> None:
    st.subheader("Resultado del análisis")

    columns = st.columns(4)
    columns[0].metric("Proceso", analysis.detected_process)
    columns[1].metric("Guía principal", analysis.selected_guide_name)
    columns[2].metric("Secciones", len(analysis.sections))
    columns[3].metric(
        "Preguntas obligatorias",
        count_required_questions(analysis),
    )

    status_counts = {
        status: sum(
            1 for section in analysis.sections if section.status == status
        )
        for status in STATUS_LABELS
    }
    completed_ratio = (
        status_counts["completo"] / len(analysis.sections)
        if analysis.sections
        else 0
    )
    st.progress(
        completed_ratio,
        text=(
            f"{status_counts['completo']} de {len(analysis.sections)} "
            "secciones tienen información completa"
        ),
    )

    st.info(analysis.selection_reason)

    structure_manifest = build_guide_structure_manifest(guides)
    selected_structure = next(
        (
            item for item in structure_manifest
            if item["guide_index"] == analysis.selected_guide_index
        ),
        None,
    )
    if selected_structure and selected_structure["sections"]:
        with st.expander("Estructura exacta bloqueada por la guía", expanded=True):
            st.info(
                "El bot conservará literalmente estos títulos, en este orden, "
                "y no permitirá reducir la cantidad de pasos."
            )
            for section in selected_structure["sections"]:
                step_count = len(section["numbered_items"])
                if step_count:
                    st.success(
                        f"{section['order']}. {section['title']} — "
                        f"{step_count} pasos obligatorios"
                    )
                    for item_index, item in enumerate(section["numbered_items"], start=1):
                        st.write(f"{item_index}. {item}")
                else:
                    st.write(f"**{section['order']}. {section['title']}**")

    if analysis.supporting_guide_indices:
        supporting_names = [
            guides[index]["name"]
            for index in analysis.supporting_guide_indices
        ]
        st.write("**Guías complementarias:** " + ", ".join(supporting_names))

    if analysis.general_requirements:
        with st.expander("Requisitos generales identificados"):
            for requirement in analysis.general_requirements:
                st.write(f"- {requirement}")

    for warning in analysis.warnings:
        st.warning(warning)

    comparison_rows = []
    for section in analysis.sections:
        comparison_rows.append(
            {
                "Sección": section.title,
                "Estado": STATUS_LABELS[section.status],
                "Criterios": len(section.criteria),
                "Evidencias": len(section.evidence),
                "Preguntas": len(section.missing_questions),
            }
        )

    st.markdown("### Comparación guía vs. información encontrada")
    st.dataframe(
        comparison_rows,
        width="stretch",
        hide_index=True,
        column_config={
            "Sección": st.column_config.TextColumn(width="large"),
            "Estado": st.column_config.TextColumn(width="small"),
        },
    )

    st.markdown("### Evaluación detallada por sección")

    for section in analysis.sections:
        icon = {
            "completo": "✅",
            "parcial": "🟠",
            "faltante": "🔴",
        }[section.status]

        with st.expander(
            f"{icon} {section.order}. {section.title} — "
            f"{STATUS_LABELS[section.status]}"
        ):
            st.write("**Instrucción interpretada de la guía**")
            st.write(section.guide_instruction)

            st.write("**Criterios**")
            for criterion in section.criteria:
                st.write(f"- {criterion}")

            if section.evidence:
                st.write("**Sustento encontrado en el documento de origen**")
                for evidence in section.evidence:
                    st.write(f"- {evidence}")

            if section.draft_content.strip():
                st.write("**Borrador sustentado**")
                st.write(section.draft_content)

            if section.missing_questions:
                st.write("**Información que debe completarse**")
                for question in section.missing_questions:
                    required = "Obligatoria" if question.required else "Opcional"
                    st.write(
                        f"- **{question.question}** "
                        f"({question.category} · {required})"
                    )
                    st.caption(question.why_needed)

            st.caption(f"Formato previsto: {section.output_format}")


def collect_answers_form(analysis: GuideAnalysis) -> list[UserAnswer] | None:
    """Agrupa preguntas críticas por categoría y valida las obligatorias."""

    all_questions = [
        (section, question_index, question)
        for section in analysis.sections
        for question_index, question in enumerate(section.missing_questions)
    ]

    if not all_questions:
        st.success(
            "La información disponible es suficiente. "
            "Puedes generar el borrador directamente."
        )

    grouped_questions: dict[str, list[tuple]] = {}
    for item in all_questions:
        category = item[2].category.strip() or "Información del proceso"
        grouped_questions.setdefault(category, []).append(item)

    answers_by_key: dict[str, str] = {}

    with st.form("missing_information_form", clear_on_submit=False):
        st.markdown("### Completar información faltante")
        st.caption(
            "Responde únicamente con información verificable. "
            "Los campos marcados con * son obligatorios."
        )

        item_index = 0
        for category, category_items in grouped_questions.items():
            st.markdown(f"#### {category}")

            for section, question_index, question in category_items:
                label = question.question + (" *" if question.required else "")
                help_text = f"Sección: {section.title}. {question.why_needed}"
                key = (
                    f"answer_{section.order}_{question_index}_{item_index}"
                )

                answers_by_key[key] = st.text_area(
                    label,
                    height=90,
                    help=help_text,
                    key=key,
                    placeholder="Escribe una respuesta concreta y verificable.",
                )
                item_index += 1

        submitted = st.form_submit_button(
            "Crear borrador para revisión",
            type="primary",
            width="stretch",
        )

    if not submitted:
        return None

    missing_required: list[str] = []
    answers: list[UserAnswer] = []
    item_index = 0

    for category_items in grouped_questions.values():
        for section, question_index, question in category_items:
            key = f"answer_{section.order}_{question_index}_{item_index}"
            answer_text = answers_by_key.get(key, "").strip()

            if question.required and not answer_text:
                missing_required.append(question.question)

            if answer_text:
                answers.append(
                    UserAnswer(
                        section_title=section.title,
                        question=question.question,
                        answer=answer_text,
                    )
                )
            item_index += 1

    if missing_required:
        st.error(
            "Debes responder los siguientes datos obligatorios:\n\n- "
            + "\n- ".join(missing_required)
        )
        return None

    return answers


def _split_paragraph_editor(value: str) -> list[str]:
    """Convierte bloques separados por línea en blanco en párrafos."""

    return [
        re.sub(r"[ \t]+", " ", item).strip()
        for item in re.split(r"\n\s*\n", value.strip())
        if item.strip()
    ]


def _split_list_editor(value: str) -> list[str]:
    """Convierte una línea por elemento en una lista simple."""

    return [
        _strip_list_prefix(item)
        for item in value.splitlines()
        if _strip_list_prefix(item)
    ]


def _split_numbered_editor(value: str) -> list[str]:
    """Conserva el título y la descripción de cada paso como una sola unidad."""

    blocks = [
        _strip_list_prefix(block)
        for block in re.split(r"\n\s*\n+", str(value or "").strip())
        if _strip_list_prefix(block)
    ]
    if blocks:
        return blocks

    return _split_list_editor(value)


def _records_from_data_editor(edited_data) -> list[dict]:
    if hasattr(edited_data, "to_dict"):
        return edited_data.to_dict(orient="records")
    if isinstance(edited_data, list):
        return [dict(item) for item in edited_data]
    return []


def replace_section_in_document(
    final_document: FinalDocument,
    regenerated_section: FinalSection,
) -> FinalDocument:
    sections = [
        regenerated_section
        if section.order == regenerated_section.order
        else section
        for section in final_document.sections
    ]
    return normalize_final_document_structure(
        final_document.model_copy(update={"sections": sections})
    )


def edit_final_document_form(
    final_document: FinalDocument,
) -> FinalDocument | None:
    """Permite editar títulos, texto, listas y celdas antes de generar."""

    final_document = normalize_final_document_structure(final_document)
    section_values: list[dict] = []

    with st.form("final_document_editor", clear_on_submit=False):
        st.markdown("### Vista previa editable")
        st.caption(
            "Revisa el contenido antes de crear los archivos. "
            "Los cambios guardados reemplazarán el borrador de la IA."
        )

        edited_title = st.text_input(
            "Título del documento",
            value=final_document.title,
        )
        edited_subtitle = st.text_input(
            "Subtítulo",
            value=final_document.subtitle,
        )
        edited_intro = st.text_area(
            "Nota introductoria",
            value=final_document.introductory_note,
            height=100,
        )

        for section_index, section in enumerate(final_document.sections):
            with st.expander(
                f"{section.order}. {section.title}",
                expanded=section_index == 0,
            ):
                st.text_input(
                    "Título de la sección — bloqueado por la guía",
                    value=section.title,
                    key=f"edit_title_{section.order}",
                    disabled=True,
                    help="El título y el orden provienen literalmente de la guía principal.",
                )
                title_value = section.title
                paragraphs_value = st.text_area(
                    "Párrafos",
                    value="\n\n".join(section.paragraphs),
                    height=max(120, min(320, 80 + 35 * len(section.paragraphs))),
                    key=f"edit_paragraphs_{section.order}",
                    help="Separa los párrafos con una línea en blanco.",
                )
                numbered_value = st.text_area(
                    "Pasos numerados",
                    value="\n\n".join(section.numbered_items),
                    height=max(90, min(250, 70 + 25 * len(section.numbered_items))),
                    key=f"edit_numbered_{section.order}",
                    help=(
                        "Cada paso contiene un título y una descripción. Separa un paso "
                        "de otro con una línea en blanco; el bot aplicará la numeración."
                    ),
                )
                bullets_value = st.text_area(
                    "Viñetas",
                    value="\n".join(section.bullets),
                    height=max(90, min(250, 70 + 25 * len(section.bullets))),
                    key=f"edit_bullets_{section.order}",
                    help="Escribe un elemento por línea, sin el símbolo de viñeta.",
                )

                edited_tables: list[tuple[str, list[str], object]] = []
                for table_index, table_data in enumerate(section.tables):
                    st.markdown(f"**Tabla {table_index + 1}**")
                    table_title = st.text_input(
                        "Título de la tabla",
                        value=table_data.title,
                        key=f"table_title_{section.order}_{table_index}",
                    )

                    headers = list(table_data.headers)
                    if headers:
                        records = []
                        for row in table_data.rows:
                            padded = list(row) + [""] * (len(headers) - len(row))
                            records.append(
                                {
                                    header: padded[column_index]
                                    for column_index, header in enumerate(headers)
                                }
                            )

                        editor_source = (
                            records
                            if records
                            else {header: [] for header in headers}
                        )
                        edited_data = st.data_editor(
                            editor_source,
                            num_rows="dynamic",
                            width="stretch",
                            hide_index=True,
                            key=f"table_editor_{section.order}_{table_index}",
                        )
                        edited_tables.append(
                            (table_title, headers, edited_data)
                        )
                    else:
                        st.caption("La tabla no tiene encabezados editables.")

                section_values.append(
                    {
                        "section": section,
                        "title": title_value,
                        "paragraphs": paragraphs_value,
                        "numbered": numbered_value,
                        "bullets": bullets_value,
                        "tables": edited_tables,
                    }
                )

        submitted = st.form_submit_button(
            "Guardar cambios, auditar y generar archivos",
            type="primary",
            width="stretch",
        )

    if not submitted:
        return None

    rebuilt_sections: list[FinalSection] = []

    for values in section_values:
        original_section: FinalSection = values["section"]
        rebuilt_tables: list[FinalTable] = []

        for table_title, headers, edited_data in values["tables"]:
            records = _records_from_data_editor(edited_data)
            rows = [
                [
                    _normalize_document_text(record.get(header, ""))
                    for header in headers
                ]
                for record in records
            ]
            rebuilt_tables.append(
                FinalTable(
                    title=table_title.strip(),
                    headers=headers,
                    rows=rows,
                )
            )

        if not values["tables"]:
            rebuilt_tables = original_section.tables

        rebuilt_sections.append(
            original_section.model_copy(
                update={
                    "title": values["title"].strip(),
                    "paragraphs": _split_paragraph_editor(values["paragraphs"]),
                    "numbered_items": _split_numbered_editor(values["numbered"]),
                    "bullets": _split_list_editor(values["bullets"]),
                    "tables": rebuilt_tables,
                }
            )
        )

    edited_document = final_document.model_copy(
        update={
            "title": edited_title.strip() or final_document.title,
            "subtitle": edited_subtitle.strip(),
            "introductory_note": edited_intro.strip(),
            "sections": rebuilt_sections,
        }
    )

    return normalize_final_document_structure(edited_document)


def show_error_panel(
    user_message: str,
    error: Exception,
    technical_label: str = "Ver detalle técnico",
) -> None:
    """Muestra un mensaje comprensible y oculta el detalle técnico."""

    st.error(user_message)
    with st.expander(technical_label):
        st.code(str(error)[:2000], language=None)

def show_final_document(final_document: FinalDocument) -> None:
    st.subheader("Documento final")
    source_record = st.session_state.get("source_record") or {}
    action_count = int(source_record.get("video_action_count") or 0)
    if action_count:
        st.success(
            f"Cobertura procedimental aplicada: {action_count} acciones "
            "operativas identificadas en el video."
        )
    st.markdown(f"## {final_document.title}")

    if final_document.subtitle.strip():
        st.caption(final_document.subtitle)

    if final_document.introductory_note.strip():
        st.write(final_document.introductory_note)

    for section in final_document.sections:
        with st.expander(f"{section.order}. {section.title}", expanded=False):
            for paragraph_text in section.paragraphs:
                st.write(paragraph_text)

            for index, item in enumerate(section.numbered_items, start=1):
                step_title, step_body = _split_step_title_body(item)
                if step_title:
                    st.markdown(f"**{index}. {step_title.rstrip(':')}:**")
                    if step_body:
                        st.write(step_body)
                else:
                    st.write(f"{index}. {step_body}")

            for bullet in section.bullets:
                st.write(f"- {bullet}")

            for table_data in section.tables:
                if table_data.title.strip():
                    st.write(f"**{table_data.title}**")

                if table_data.headers and table_data.rows:
                    normalized_rows = []
                    for row in table_data.rows:
                        padded = list(row) + [""] * (
                            len(table_data.headers) - len(row)
                        )
                        normalized_rows.append(
                            padded[: len(table_data.headers)]
                        )

                    st.dataframe(
                        {
                            header: [row[index] for row in normalized_rows]
                            for index, header in enumerate(table_data.headers)
                        },
                        width="stretch",
                        hide_index=True,
                    )

            if section.source_basis:
                st.caption("Sustento: " + " | ".join(section.source_basis))

    st.markdown("### Validación contra la guía")

    compliance_counts = {"cumple": 0, "parcial": 0, "no_aplica": 0}
    for item in final_document.validation:
        compliance_counts[item.status] += 1

    total_checks = sum(compliance_counts.values())
    compliance_ratio = (
        compliance_counts["cumple"] / total_checks
        if total_checks
        else 0
    )
    st.progress(
        compliance_ratio,
        text=(
            f"{compliance_counts['cumple']} de {total_checks} "
            "criterios cumplen completamente"
        ),
    )

    columns = st.columns(3)
    columns[0].metric("Cumple", compliance_counts["cumple"])
    columns[1].metric("Parcial", compliance_counts["parcial"])
    columns[2].metric("No aplica", compliance_counts["no_aplica"])

    with st.expander("Ver validación detallada"):
        for item in final_document.validation:
            icon = {
                "cumple": "✅",
                "parcial": "🟠",
                "no_aplica": "⚪",
            }[item.status]
            st.write(
                f"{icon} **{item.section_title}** — "
                f"{item.criterion}: {item.note}"
            )

    audit_summary = st.session_state.get("audit_summary")
    if audit_summary:
        st.info(f"**Resumen editorial:** {audit_summary}")

    for warning in final_document.warnings:
        st.warning(warning)


# =========================================================
# DISEÑO VISUAL DE LA INTERFAZ
# =========================================================


def inject_app_styles() -> None:
    """Aplica un diseño más moderno, legible y amigable."""

    st.markdown(
        """
        <style>
        :root {
            --brand: #ff4b4b;
            --brand-dark: #d9363e;
            --panel: rgba(255,255,255,0.055);
            --line: rgba(255,255,255,0.12);
            --muted: rgba(255,255,255,0.72);
        }

        .block-container {
            max-width: 1180px;
            padding-top: 2rem;
            padding-bottom: 4rem;
        }

        [data-testid="stSidebar"] {
            border-right: 1px solid var(--line);
        }

        .app-hero {
            padding: 1.6rem 1.8rem;
            border: 1px solid rgba(255,75,75,0.30);
            border-radius: 22px;
            background:
                radial-gradient(circle at 90% 20%, rgba(255,75,75,0.22), transparent 34%),
                linear-gradient(135deg, rgba(255,75,75,0.12), rgba(255,255,255,0.035));
            box-shadow: 0 16px 42px rgba(0,0,0,0.18);
            margin-bottom: 1.4rem;
        }

        .app-hero h1 {
            margin: 0 0 .45rem 0;
            font-size: clamp(2rem, 4vw, 3.2rem);
            line-height: 1.05;
        }

        .app-hero p {
            margin: 0;
            max-width: 830px;
            color: var(--muted);
            font-size: 1.04rem;
        }

        .step-grid {
            display: grid;
            grid-template-columns: repeat(4, minmax(0, 1fr));
            gap: .75rem;
            margin: .5rem 0 1.5rem;
        }

        .step-card {
            padding: .9rem 1rem;
            border-radius: 15px;
            border: 1px solid var(--line);
            background: var(--panel);
            min-height: 82px;
        }

        .step-card.active {
            border-color: rgba(255,75,75,0.68);
            background: rgba(255,75,75,0.13);
        }

        .step-number {
            display: inline-flex;
            align-items: center;
            justify-content: center;
            width: 25px;
            height: 25px;
            border-radius: 50%;
            background: rgba(255,255,255,.10);
            font-size: .8rem;
            font-weight: 700;
            margin-bottom: .35rem;
        }

        .step-card.active .step-number { background: var(--brand); color: white; }
        .step-card strong { display: block; font-size: .92rem; }
        .step-card span { color: var(--muted); font-size: .78rem; }

        [data-testid="stFileUploader"] {
            background: var(--panel);
            border: 1px dashed rgba(255,255,255,.22);
            border-radius: 16px;
            padding: .55rem .7rem;
        }

        [data-testid="stMetric"] {
            border: 1px solid var(--line);
            background: var(--panel);
            padding: .8rem 1rem;
            border-radius: 14px;
        }

        .friendly-note {
            border-left: 4px solid var(--brand);
            padding: .8rem 1rem;
            border-radius: 0 12px 12px 0;
            background: rgba(255,75,75,.08);
            margin: .75rem 0 1rem;
        }

        .sidebar-brand {
            padding: 1rem;
            border-radius: 16px;
            background: linear-gradient(135deg, rgba(255,75,75,.20), rgba(255,255,255,.05));
            border: 1px solid rgba(255,75,75,.25);
            margin-bottom: 1rem;
        }

        .sidebar-brand strong { font-size: 1.05rem; }
        .sidebar-brand small { color: var(--muted); }

        .stButton > button, .stDownloadButton > button {
            border-radius: 12px;
            min-height: 46px;
            font-weight: 650;
        }

        @media (max-width: 850px) {
            .step-grid { grid-template-columns: 1fr 1fr; }
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def render_hero() -> None:
    st.markdown(
        """
        <section class="app-hero">
          <h1>Generador documental inteligente</h1>
          <p>
            Carga las guías institucionales y un documento o video de origen. El sistema
            interpreta las reglas, transcribe el video cuando corresponda,
            identifica información faltante y genera un Word y un PDF listos para revisión.
          </p>
        </section>
        """,
        unsafe_allow_html=True,
    )


def render_workflow_steps(current_step: int) -> None:
    steps = [
        (1, "Cargar archivos", "Guías y documento o video"),
        (2, "Analizar", "Comparación con Gemini"),
        (3, "Completar", "Solo datos críticos faltantes"),
        (4, "Generar", "Word y PDF institucionales"),
    ]
    cards = []
    for number, title, subtitle in steps:
        active = " active" if number <= current_step else ""
        cards.append(
            f'<div class="step-card{active}"><div class="step-number">{number}</div>'
            f'<strong>{title}</strong><span>{subtitle}</span></div>'
        )
    st.markdown('<div class="step-grid">' + ''.join(cards) + '</div>', unsafe_allow_html=True)


def current_workflow_step() -> int:
    if st.session_state.get("docx_output") is not None:
        return 4
    if st.session_state.get("final_document") is not None:
        return 3
    if st.session_state.get("analysis") is not None:
        return 3
    if st.session_state.get("guide_records"):
        return 2
    return 1

# =========================================================
# APLICACIÓN PRINCIPAL
# =========================================================


def main() -> None:
    st.set_page_config(
        page_title="Generador documental con guías",
        page_icon="📄",
        layout="wide",
    )

    initialize_state()
    inject_app_styles()

    api_key = read_secret("GEMINI_API_KEY")
    configured_model = read_secret("GEMINI_MODEL", MODEL_DEFAULT)
    if configured_model in {"gemini-2.5-flash", "gemini-2.5-flash-lite"}:
        configured_model = MODEL_DEFAULT

    render_hero()
    render_workflow_steps(current_workflow_step())

    with st.expander("¿Cómo funciona este proceso?", expanded=False):
        st.markdown(
            """
            1. **Carga las guías institucionales** y elige un documento, un video pequeño o un video de Google Drive.
            2. Los videos de Drive se **descargan por bloques al disco temporal**, sin guardarlos completos en la memoria de Streamlit.
            3. Gemini procesa el video en **tramos reanudables de cinco minutos**, con auditoría selectiva y hasta dos tramos simultáneos.
            4. Cada tramo terminado se guarda como checkpoint local y, cuando se configura, también en **Google Drive**.
            5. El sistema **selecciona la guía aplicable y respeta exactamente sus títulos, orden y cantidad de pasos**.
            6. Solo solicita **datos críticos que realmente estén ausentes**.
            7. Presenta un **borrador editable** y crea el Word y el PDF institucionales.
            """
        )

    with st.sidebar:
        st.markdown(
            """<div class="sidebar-brand"><strong>Centro de control</strong><br>
            <small>Estado del servicio y opciones del proceso</small></div>""",
            unsafe_allow_html=True,
        )
        st.subheader("Configuración")

        if api_key:
            st.success("Clave de Gemini detectada")
        else:
            st.error("Falta GEMINI_API_KEY")

        drive_email = drive_service_account_email()
        if drive_email:
            st.success("Cuenta de Google Drive detectada")
            st.caption(f"Correo de acceso: {drive_email}")
        else:
            st.warning("Google Drive aún no está configurado")

        checkpoint_folder = _drive_checkpoint_folder_id()
        if checkpoint_folder:
            st.success("Checkpoints reanudables en Google Drive habilitados")
            st.caption(
                "La carpeta debe estar compartida como Editor con la cuenta de servicio."
            )
        else:
            st.info(
                "Los checkpoints se guardarán localmente. Para conservarlos incluso "
                "después de reiniciar el servidor, configura DRIVE_CHECKPOINT_FOLDER_ID."
            )

        model = st.text_input(
            "Modelo Gemini preferido",
            value=configured_model,
            help=(
                "El bot verifica los modelos habilitados para tu clave y cambia "
                "automáticamente a uno disponible si este nombre ya fue retirado."
            ),
        ).strip()

        active_model = st.session_state.get("active_gemini_model")
        if active_model:
            st.caption(f"Modelo usado en el último proceso: {active_model}")
        else:
            st.caption("Selección automática de modelo habilitada.")

        st.caption(
            f"Hasta {MAX_GUIDES} guías. Carga directa: {MAX_VIDEO_SIZE_MB} MB. "
            f"Google Drive: {MAX_DRIVE_VIDEO_SIZE_MB} MB recomendados."
        )

        if st.button("Reiniciar proceso", width="stretch"):
            reset_workflow()
            st.rerun()

    st.subheader("📁 1. Cargar guías e información de origen")

    st.markdown(
        """<div class="friendly-note"><strong>Consejo:</strong> para conservar
        exactamente el encabezado, pie de página, logos y márgenes, carga al
        menos una guía en formato <strong>DOCX</strong>.</div>""",
        unsafe_allow_html=True,
    )

    guide_uploads = st.file_uploader(
        "Guías del proceso",
        type=GUIDE_EXTENSIONS,
        accept_multiple_files=True,
        key=f"guides_{st.session_state.upload_key}",
        help=(
            "Carga las guías que describen los apartados, criterios y reglas "
            "del documento final."
        ),
    )

    upload_mode_label = "Subir documento o video pequeño"
    drive_mode_label = "Usar video desde Google Drive"
    source_mode = st.radio(
        "¿De dónde viene la información de origen?",
        options=[upload_mode_label, drive_mode_label],
        horizontal=True,
        key="source_origin_mode",
        help=(
            "Usa Google Drive para videos grandes. Así el navegador no envía el "
            "archivo completo a la memoria de Streamlit."
        ),
    )

    source_upload = None
    drive_reference = ""
    source_from_drive = source_mode == drive_mode_label

    if source_from_drive:
        drive_reference = st.text_input(
            "Enlace o ID del video en Google Drive",
            key="drive_video_reference",
            placeholder="https://drive.google.com/file/d/ID/view",
            help=(
                "Comparte el video o la carpeta como Lector con el correo de la "
                "cuenta de servicio mostrado debajo. No necesitas hacerlo público."
            ),
        ).strip()

        service_email = drive_service_account_email()
        if service_email:
            st.success(
                "Google Drive configurado. Comparte el archivo o la carpeta como "
                f"Lector con: {service_email}"
            )
        else:
            st.error(
                "Falta DRIVE_SERVICE_ACCOUNT_JSON en los Secrets de Streamlit. "
                "Sin esa credencial el bot no puede leer archivos privados de Drive."
            )

        st.caption(
            f"Drive: máximo recomendado {MAX_DRIVE_VIDEO_SIZE_MB} MB. "
            "El video se descarga por bloques al disco temporal y se elimina al terminar."
        )
    else:
        source_upload = st.file_uploader(
            "Documento o video de origen",
            type=SOURCE_EXTENSIONS,
            accept_multiple_files=False,
            key=f"source_{st.session_state.upload_key}",
            help=(
                "Puedes cargar DOCX, PDF, TXT o un video pequeño. Para videos "
                "grandes selecciona Google Drive."
            ),
        )
        st.caption(
            f"Carga directa: máximo {MAX_VIDEO_SIZE_MB} MB por video y "
            f"{MAX_TOTAL_UPLOAD_MB} MB en total."
        )

    source_extension = (
        file_extension(source_upload.name)
        if source_upload is not None
        else ""
    )
    source_is_video = source_from_drive or is_video_extension(source_extension)

    current_video_fingerprint = None
    if source_from_drive and drive_reference:
        try:
            current_video_fingerprint = (
                "drive:" + extract_drive_file_id(drive_reference)
            )
        except AppError:
            current_video_fingerprint = "drive:invalid:" + drive_reference
    elif source_upload is not None and source_is_video:
        declared_size = int(getattr(source_upload, "size", 0) or 0)
        current_video_fingerprint = (
            f"upload:{safe_filename(source_upload.name)}:{declared_size}"
        )

    stored_fingerprint = st.session_state.get("video_source_fingerprint")
    if stored_fingerprint and stored_fingerprint != current_video_fingerprint:
        st.session_state.video_transcript = None
        st.session_state.video_source_text = None
        st.session_state.video_source_fingerprint = None
        st.session_state.source_record = None
        st.session_state.analysis = None
        _clear_outputs_after_new_analysis()

    def upload_size(upload) -> int:
        if upload is None:
            return 0
        declared = getattr(upload, "size", None)
        if declared is not None:
            return int(declared)
        return len(upload.getvalue())

    def validate_current_inputs() -> None:
        if not guide_uploads:
            raise AppError("Debes cargar al menos una guía.")

        if len(guide_uploads) > MAX_GUIDES:
            raise AppError(f"Solo se permiten hasta {MAX_GUIDES} guías.")

        names = [
            safe_filename(upload.name).lower()
            for upload in guide_uploads
        ]
        if len(names) != len(set(names)):
            raise AppError("Hay guías con nombres duplicados.")

        if source_from_drive:
            extract_drive_file_id(drive_reference)
            info = read_drive_service_account_info()
            if not info.get("client_email") or not info.get("private_key"):
                raise AppError(
                    "Configura DRIVE_SERVICE_ACCOUNT_JSON en los Secrets de "
                    "Streamlit antes de usar Google Drive."
                )
            total_upload_bytes = sum(upload_size(upload) for upload in guide_uploads)
        else:
            if source_upload is None:
                raise AppError("Debes cargar el documento o video de origen.")
            total_upload_bytes = sum(
                upload_size(upload) for upload in guide_uploads
            ) + upload_size(source_upload)

        if total_upload_bytes > MAX_TOTAL_UPLOAD_BYTES:
            raise AppError(
                f"Los archivos enviados por el navegador superan "
                f"{MAX_TOTAL_UPLOAD_MB} MB. Usa Google Drive para el video grande."
            )

    def build_current_guides() -> list[dict]:
        return [
            build_file_record(
                upload,
                role="guía normativa",
                index=index,
                allowed_extensions=GUIDE_EXTENSIONS,
            )
            for index, upload in enumerate(guide_uploads or [])
        ]

    if source_is_video:
        st.info(
            "El video se procesa directamente en intervalos de cinco minutos, "
            "con hasta dos intervalos simultáneos y auditoría selectiva. Cada "
            "tramo completado se guarda para poder reanudar después de un reinicio."
        )
        if source_from_drive:
            st.caption(
                "El video no pasa por el cargador del navegador. Se descarga por "
                "fragmentos, se envía temporalmente a Gemini y luego se elimina."
            )
        else:
            st.caption(
                f"Límite de carga directa: {MAX_VIDEO_SIZE_MB} MB. Para archivos "
                "mayores utiliza la opción de Google Drive."
            )

        transcript_ready = (
            st.session_state.get("video_transcript") is not None
            and st.session_state.get("video_source_fingerprint")
            == current_video_fingerprint
            and st.session_state.get("source_record") is not None
        )

        if not transcript_ready:
            drive_ready = bool(
                drive_reference and drive_service_account_email()
            ) if source_from_drive else True
            transcribe_clicked = st.button(
                "Transcribir o reanudar video y extraer evidencia visual",
                type="primary",
                width="stretch",
                disabled=not api_key or not drive_ready,
            )

            if transcribe_clicked:
                try:
                    validate_current_inputs()

                    with st.status(
                        "Procesando el video...",
                        expanded=True,
                    ) as status:
                        st.write("1/4 · Validando el origen y las guías")
                        guide_records = build_current_guides()
                        progress_slot = st.empty()

                        def video_progress(message: str) -> None:
                            progress_slot.write(f"2/4 · {message}")

                        if source_from_drive:
                            transcript, source_record = (
                                transcribe_drive_video_with_gemini(
                                    api_key=api_key,
                                    model=model or MODEL_DEFAULT,
                                    drive_reference=drive_reference,
                                    progress_callback=video_progress,
                                )
                            )
                        else:
                            source_record = build_file_record(
                                source_upload,
                                role="video de origen",
                                index=0,
                                allowed_extensions=SOURCE_EXTENSIONS,
                            )
                            transcript = transcribe_video_with_gemini(
                                api_key=api_key,
                                model=model or MODEL_DEFAULT,
                                video_record=source_record,
                                progress_callback=video_progress,
                            )
                            # No conservar otra copia del video dentro del estado.
                            source_record["content"] = b""
                            source_record["origin_kind"] = "direct_upload"
                            source_record["video_checkpoint_id"] = (
                                st.session_state.get("video_checkpoint_id")
                            )

                        st.write("3/4 · Preparando el origen textual editable")
                        source_text = render_video_transcript_as_source(
                            transcript,
                            source_record["name"],
                        )
                        source_record["text"] = source_text
                        source_record["warnings"] = [
                            warning
                            for warning in source_record.get("warnings", [])
                            if "Debe transcribirse" not in warning
                        ]
                        source_record["warnings"].append(
                            "Origen generado desde video. Revisa la transcripción "
                            "antes de continuar con el análisis documental."
                        )
                        source_record["video_transcript"] = transcript.model_dump()
                        source_record["video_actions"] = [
                            action.model_dump() for action in transcript.actions
                        ]
                        source_record["video_action_count"] = len(transcript.actions)

                        st.write("4/4 · Video listo para revisión")
                        status.update(
                            label="Transcripción lista para revisión",
                            state="complete",
                            expanded=False,
                        )

                    st.session_state.guide_records = guide_records
                    st.session_state.source_record = source_record
                    st.session_state.video_transcript = transcript
                    st.session_state.video_source_text = source_text
                    st.session_state.video_source_fingerprint = (
                        current_video_fingerprint
                    )
                    st.session_state.analysis = None
                    _clear_outputs_after_new_analysis()
                    st.success(
                        "El video fue transcrito. Revisa el texto antes de "
                        "compararlo con las guías."
                    )
                    st.rerun()

                except AppError as error:
                    show_error_panel(
                        "No fue posible transcribir el video.",
                        error,
                    )
                except Exception as error:
                    show_error_panel(
                        "Ocurrió un error inesperado al procesar el video.",
                        error,
                    )

        else:
            source_record = st.session_state.source_record
            approved_source_text = show_video_transcript_review(source_record)

            if approved_source_text is not None:
                try:
                    validate_current_inputs()
                    guide_records = build_current_guides()
                    source_record = deepcopy(source_record)
                    source_record["text"] = approved_source_text
                    source_record["role"] = (
                        "transcripción revisada del video de origen"
                    )
                    edited_actions = _parse_action_inventory_from_source_text(
                        approved_source_text
                    )
                    if edited_actions:
                        source_record["video_actions"] = [
                            action.model_dump() for action in edited_actions
                        ]
                        source_record["video_action_count"] = len(edited_actions)
                    source_record["warnings"] = [
                        warning
                        for warning in source_record.get("warnings", [])
                        if "Revisa la transcripción" not in warning
                    ]
                    source_record["warnings"].append(
                        "La transcripción fue revisada y aprobada por el usuario."
                    )

                    with st.status(
                        "Analizando las guías con la transcripción aprobada...",
                        expanded=True,
                    ) as status:
                        st.write("1/4 · Validando las guías")
                        total_text = sum(
                            len(record["text"]) for record in guide_records
                        ) + len(source_record["text"])
                        if total_text > MAX_TOTAL_TEXT_CHARS:
                            st.warning(
                                "Los archivos contienen mucho texto. Las guías "
                                "se limitarán cuando sea necesario, pero la "
                                "transcripción y el inventario del video se "
                                "conservarán para la generación."
                            )

                        st.write("2/4 · Seleccionando la guía aplicable")
                        st.write("3/4 · Comparando criterios y evidencia")
                        analysis = analyze_guides_and_source(
                            api_key=api_key,
                            model=model or MODEL_DEFAULT,
                            guides=guide_records,
                            source=source_record,
                        )
                        st.write("4/4 · Identificando brechas y preguntas críticas")
                        status.update(
                            label="Análisis completado",
                            state="complete",
                            expanded=False,
                        )

                    st.session_state.guide_records = guide_records
                    st.session_state.source_record = source_record
                    st.session_state.video_source_text = approved_source_text
                    st.session_state.analysis = analysis
                    _clear_outputs_after_new_analysis()
                    st.success(
                        "La transcripción fue aprobada y analizada correctamente."
                    )
                    st.rerun()

                except AppError as error:
                    show_error_panel(
                        "No fue posible analizar la transcripción con las guías.",
                        error,
                    )
                except Exception as error:
                    show_error_panel(
                        "Ocurrió un error inesperado durante el análisis.",
                        error,
                    )

    else:
        analyze_clicked = st.button(
            "Analizar guías y documento de origen",
            type="primary",
            width="stretch",
            disabled=not api_key,
        )

        if analyze_clicked:
            try:
                validate_current_inputs()

                with st.status(
                    "Analizando el proceso documental...",
                    expanded=True,
                ) as status:
                    st.write("1/4 · Leyendo y validando los archivos")
                    guide_records = build_current_guides()
                    source_record = build_file_record(
                        source_upload,
                        role="documento de origen",
                        index=0,
                        allowed_extensions=SOURCE_EXTENSIONS,
                    )

                    total_text = sum(
                        len(record["text"]) for record in guide_records
                    )
                    total_text += len(source_record["text"])

                    if total_text > MAX_TOTAL_TEXT_CHARS:
                        st.warning(
                            "Los archivos contienen mucho texto. Se limitará el "
                            "contenido textual por archivo; los PDF se enviarán "
                            "completos de forma nativa."
                        )

                    st.write("2/4 · Interpretando las reglas de las guías")
                    st.write("3/4 · Extrayendo evidencia del documento de origen")
                    analysis = analyze_guides_and_source(
                        api_key=api_key,
                        model=model or MODEL_DEFAULT,
                        guides=guide_records,
                        source=source_record,
                    )
                    st.write("4/4 · Identificando brechas y preguntas críticas")
                    status.update(
                        label="Análisis completado",
                        state="complete",
                        expanded=False,
                    )

                st.session_state.guide_records = guide_records
                st.session_state.source_record = source_record
                st.session_state.analysis = analysis
                st.session_state.video_transcript = None
                st.session_state.video_source_text = None
                st.session_state.video_source_fingerprint = None
                _clear_outputs_after_new_analysis()

                st.success("Análisis completado correctamente.")

            except AppError as error:
                show_error_panel(
                    "No fue posible completar el análisis. "
                    "Revisa el mensaje técnico y vuelve a intentarlo.",
                    error,
                )
            except Exception as error:
                show_error_panel(
                    "Ocurrió un error inesperado durante el análisis.",
                    error,
                )

    guide_records = st.session_state.guide_records
    source_record = st.session_state.source_record
    analysis = st.session_state.analysis
    final_document = st.session_state.final_document

    if guide_records and source_record:
        st.divider()
        st.subheader("Archivos procesados")

        for index, guide in enumerate(guide_records):
            show_file_record(
                guide,
                title=f"Guía {index}: {guide['name']}",
                key=f"guide_{index}",
            )

        show_file_record(
            source_record,
            title=f"Documento de origen: {source_record['name']}",
            key="source",
        )

    if analysis and guide_records and source_record:
        st.divider()
        show_analysis(analysis, guide_records)

        if final_document is None:
            st.divider()
            answers = collect_answers_form(analysis)

            if answers is not None:
                try:
                    with st.status(
                        "Construyendo el borrador profesional...",
                        expanded=True,
                    ) as status:
                        st.write("1/3 · Redactando las secciones")
                        draft_document = generate_final_with_gemini(
                            api_key=api_key,
                            model=model or MODEL_DEFAULT,
                            guides=guide_records,
                            source=source_record,
                            analysis=analysis,
                            answers=answers,
                        )
                        st.write("2/3 · Normalizando listas, tablas y redacción")
                        draft_document = normalize_final_document_structure(
                            draft_document
                        )
                        st.write("3/3 · Preparando la vista previa editable")
                        status.update(
                            label="Borrador listo para revisión",
                            state="complete",
                            expanded=False,
                        )

                    st.session_state.answers = answers
                    st.session_state.final_document = draft_document
                    st.session_state.docx_output = None
                    st.session_state.pdf_output = None
                    st.session_state.audit_summary = None
                    st.success(
                        "Borrador creado. Revísalo y edítalo antes de generar "
                        "los archivos."
                    )
                    st.rerun()

                except AppError as error:
                    show_error_panel(
                        "Gemini no pudo construir el borrador en este intento.",
                        error,
                    )
                except Exception as error:
                    show_error_panel(
                        "Ocurrió un error inesperado al construir el borrador.",
                        error,
                    )

    final_document = st.session_state.final_document

    if (
        final_document
        and analysis
        and guide_records
        and source_record
    ):
        st.divider()
        st.subheader("✍️ 3. Revisar y editar el borrador")

        st.markdown(
            """<div class="friendly-note"><strong>Control del usuario:</strong>
            puedes editar directamente cada sección o regenerar únicamente una
            sección sin rehacer todo el documento.</div>""",
            unsafe_allow_html=True,
        )

        section_options = {
            f"{section.order}. {section.title}": section.order
            for section in final_document.sections
        }
        selected_section_label = st.selectbox(
            "Sección para regenerar",
            options=list(section_options),
            help=(
                "La IA regenerará solo la sección seleccionada usando las "
                "mismas guías, el origen y tus respuestas."
            ),
        )

        if st.button(
            "Regenerar únicamente esta sección",
            width="stretch",
        ):
            try:
                with st.spinner("Regenerando la sección seleccionada..."):
                    regenerated = regenerate_section_with_gemini(
                        api_key=api_key,
                        model=model or MODEL_DEFAULT,
                        guides=guide_records,
                        source=source_record,
                        analysis=analysis,
                        answers=st.session_state.answers or [],
                        final_document=final_document,
                        section_order=section_options[selected_section_label],
                    )

                regenerated_document = replace_section_in_document(
                    final_document,
                    regenerated,
                )
                st.session_state.final_document = enforce_fidelity_with_gemini(
                    api_key=api_key,
                    model=model or MODEL_DEFAULT,
                    guides=guide_records,
                    source=source_record,
                    analysis=analysis,
                    final_document=regenerated_document,
                )
                st.session_state.docx_output = None
                st.session_state.pdf_output = None
                st.session_state.audit_summary = None
                st.success("Sección regenerada correctamente.")
                st.rerun()

            except AppError as error:
                show_error_panel(
                    "No fue posible regenerar la sección.",
                    error,
                )
            except Exception as error:
                show_error_panel(
                    "Ocurrió un error inesperado al regenerar la sección.",
                    error,
                )

        edited_document = edit_final_document_form(
            st.session_state.final_document
        )

        if edited_document is not None:
            selected_guide = guide_records[analysis.selected_guide_index]

            try:
                with st.status(
                    "Auditando y generando los archivos...",
                    expanded=True,
                ) as status:
                    st.write("1/4 · Guardando tus modificaciones")
                    edited_document = normalize_final_document_structure(
                        edited_document
                    )

                    st.write("2/5 · Verificando que no falte ningún paso de la guía")
                    edited_document = enforce_fidelity_with_gemini(
                        api_key=api_key,
                        model=model or MODEL_DEFAULT,
                        guides=guide_records,
                        source=source_record,
                        analysis=analysis,
                        final_document=edited_document,
                    )

                    st.write("3/5 · Auditando el contenido contra la guía")
                    try:
                        audit_report = audit_final_document_with_gemini(
                            api_key=api_key,
                            model=model or MODEL_DEFAULT,
                            guides=guide_records,
                            source=source_record,
                            analysis=analysis,
                            final_document=edited_document,
                        )
                        combined_warnings = list(
                            dict.fromkeys(
                                edited_document.warnings
                                + audit_report.warnings
                            )
                        )
                        edited_document = edited_document.model_copy(
                            update={
                                "validation": audit_report.validation,
                                "warnings": combined_warnings,
                            }
                        )
                        st.session_state.audit_summary = (
                            audit_report.editorial_summary
                        )
                    except Exception as audit_error:
                        st.warning(
                            "No fue posible completar la revalidación con "
                            "Gemini. El documento se generará con la validación "
                            "del borrador."
                        )
                        with st.expander("Ver detalle de la revalidación"):
                            st.code(str(audit_error)[:1600], language=None)

                    st.write("4/5 · Aplicando formato institucional y paginación")
                    docx_output = create_docx(
                        edited_document,
                        style_guide=selected_guide,
                    )

                    st.write("5/5 · Creando la versión PDF")
                    pdf_output = convert_docx_to_pdf(docx_output)
                    if pdf_output is None:
                        pdf_output = create_pdf(edited_document)

                    status.update(
                        label="Documento auditado y generado",
                        state="complete",
                        expanded=False,
                    )

                st.session_state.final_document = edited_document
                st.session_state.docx_output = docx_output
                st.session_state.pdf_output = pdf_output

                if selected_guide.get("extension") == "docx":
                    st.success(
                        "Documento generado con encabezado, pie, márgenes, "
                        "paginación dinámica y formato de la guía seleccionada."
                    )
                else:
                    st.success("Documento final generado y validado.")
                    st.info(
                        "La guía principal no es DOCX. Para reproducir exactamente "
                        "encabezado y pie, carga su versión Word."
                    )
                st.rerun()

            except AppError as error:
                show_error_panel(
                    "No fue posible generar el documento final.",
                    error,
                )
            except Exception as error:
                show_error_panel(
                    "Ocurrió un error inesperado durante la generación.",
                    error,
                )

    final_document = st.session_state.final_document
    docx_output = st.session_state.get("docx_output")
    pdf_output = st.session_state.get("pdf_output")

    if final_document and docx_output and pdf_output:
        st.divider()
        st.subheader("✅ 4. Documento listo")
        show_final_document(final_document)

        st.markdown("### Descargar archivos")
        files_are_valid = True

        if not isinstance(docx_output, (bytes, bytearray)) or not docx_output:
            st.error(
                "El archivo Word no está disponible en memoria. "
                "Vuelve a generar el documento."
            )
            files_are_valid = False

        if not isinstance(pdf_output, (bytes, bytearray)) or not pdf_output:
            st.error(
                "El archivo PDF no está disponible en memoria. "
                "Vuelve a generar el documento."
            )
            files_are_valid = False

        if files_are_valid:
            docx_bytes = bytes(docx_output)
            pdf_bytes = bytes(pdf_output)

            if not docx_bytes.startswith(b"PK"):
                st.error(
                    "El Word generado no tiene una estructura DOCX válida."
                )
                files_are_valid = False

            if not pdf_bytes.startswith(b"%PDF"):
                st.error(
                    "El PDF generado no tiene una estructura PDF válida."
                )
                files_are_valid = False

        if files_are_valid:
            docx_name = output_filename(final_document.title, "docx")
            pdf_name = output_filename(final_document.title, "pdf")

            docx_path = save_generated_file(docx_name, docx_bytes)
            pdf_path = save_generated_file(pdf_name, pdf_bytes)

            st.markdown("#### Descarga directa")
            direct_columns = st.columns(2)
            direct_columns[0].markdown(
                direct_download_link(
                    "⬇️ Descargar Word",
                    docx_bytes,
                    docx_name,
                    MIME_TYPES["docx"],
                ),
                unsafe_allow_html=True,
            )
            direct_columns[1].markdown(
                direct_download_link(
                    "⬇️ Descargar PDF",
                    pdf_bytes,
                    pdf_name,
                    MIME_TYPES["pdf"],
                ),
                unsafe_allow_html=True,
            )

            with st.expander("Descarga alternativa de Streamlit"):
                download_columns = st.columns(2)
                download_columns[0].download_button(
                    "Descargar Word con Streamlit",
                    data=docx_bytes,
                    file_name=docx_name,
                    mime=MIME_TYPES["docx"],
                    width="stretch",
                    key="download_final_docx",
                    on_click="ignore",
                )
                download_columns[1].download_button(
                    "Descargar PDF con Streamlit",
                    data=pdf_bytes,
                    file_name=pdf_name,
                    mime=MIME_TYPES["pdf"],
                    width="stretch",
                    key="download_final_pdf",
                    on_click="ignore",
                )

            st.caption(
                f"Word: {format_file_size(len(docx_bytes))} · "
                f"PDF: {format_file_size(len(pdf_bytes))}"
            )
            st.info(
                "También se guardaron copias en "
                f"`{docx_path.as_posix()}` y `{pdf_path.as_posix()}`."
            )

    st.divider()
    st.caption(
        "La aplicación utiliza las guías como reglas de construcción, conserva "
        "la trazabilidad del origen y evita inventar información."
    )


if __name__ == "__main__":
    main()
